"""
File Content Reader Utility
Created: January 29, 2026
Last Updated: January 29, 2026

Extracts text and data from various file types for AI analysis.
Supports: PDF, DOCX, XLSX, CSV, TXT, PNG, JPG (OCR)

USAGE:
    from file_content_reader import extract_file_content
    content = extract_file_content('/path/to/file.pdf')
    print(content['text'])

Author: Jim @ Shiftwork Solutions LLC
"""

import os
import mimetypes
from pathlib import Path


def extract_file_content(file_path):
    """
    Extract content from any supported file type.
    
    Returns dict with:
        - success: bool
        - text: str (extracted text content)
        - data: dict (structured data for spreadsheets)
        - file_type: str
        - error: str (if failed)
    """
    
    if not os.path.exists(file_path):
        return {
            'success': False,
            'error': 'File not found',
            'text': '',
            'data': None,
            'file_type': 'unknown'
        }
    
    # Determine file type
    file_ext = Path(file_path).suffix.lower()
    
    try:
        # PDF files
        if file_ext == '.pdf':
            return extract_pdf(file_path)
        
        # Word documents
        elif file_ext in ['.docx', '.doc']:
            return extract_docx(file_path)
        
        # Excel spreadsheets
        elif file_ext in ['.xlsx', '.xls']:
            return extract_xlsx(file_path)
        
        # CSV files
        elif file_ext == '.csv':
            return extract_csv(file_path)
        
        # Text files
        elif file_ext in ['.txt', '.md', '.log']:
            return extract_txt(file_path)
        
        # Images (OCR)
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            return extract_image(file_path)
        
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_ext}',
                'text': '',
                'data': None,
                'file_type': file_ext[1:] if file_ext else 'unknown'
            }
    
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'text': '',
            'data': None,
            'file_type': file_ext[1:] if file_ext else 'unknown'
        }


def extract_pdf(file_path):
    """Extract text from PDF files using PyPDF2"""
    try:
        import PyPDF2
        
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
        
        full_text = "\n\n".join(text_content)
        
        return {
            'success': True,
            'text': full_text,
            'data': {'num_pages': num_pages},
            'file_type': 'pdf',
            'error': None
        }
    
    except ImportError:
        return {
            'success': False,
            'error': 'PyPDF2 not installed. Install with: pip install PyPDF2',
            'text': '',
            'data': None,
            'file_type': 'pdf'
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'PDF extraction failed: {str(e)}',
            'text': '',
            'data': None,
            'file_type': 'pdf'
        }


def extract_docx(file_path):
    """Extract text from Word documents"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract tables
        tables_text = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(" | ".join(row_data))
            
            if table_data:
                tables_text.append(f"\n--- Table {table_idx + 1} ---\n" + "\n".join(table_data))
        
        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n" + "\n".join(tables_text)
        
        return {
            'success': True,
            'text': full_text,
            'data': {
                'num_paragraphs': len(paragraphs),
                'num_tables': len(doc.tables)
            },
            'file_type': 'docx',
            'error': None
        }
    
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx not installed. Install with: pip install python-docx',
            'text': '',
            'data': None,
            'file_type': 'docx'
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'DOCX extraction failed: {str(e)}',
            'text': '',
            'data': None,
            'file_type': 'docx'
        }


def extract_xlsx(file_path):
    """Extract data from Excel spreadsheets"""
    try:
        import openpyxl
        
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        sheets_text = []
        sheets_data = {}
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # Convert sheet to text
            rows = []
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                # Filter out completely empty rows
                if any(cell is not None for cell in row):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    rows.append(row_text)
                    sheet_data.append(list(row))
            
            if rows:
                sheets_text.append(f"\n--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
                sheets_data[sheet_name] = sheet_data
        
        full_text = "\n".join(sheets_text)
        
        return {
            'success': True,
            'text': full_text,
            'data': {
                'sheets': sheets_data,
                'num_sheets': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames
            },
            'file_type': 'xlsx',
            'error': None
        }
    
    except ImportError:
        return {
            'success': False,
            'error': 'openpyxl not installed. Install with: pip install openpyxl',
            'text': '',
            'data': None,
            'file_type': 'xlsx'
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'XLSX extraction failed: {str(e)}',
            'text': '',
            'data': None,
            'file_type': 'xlsx'
        }


def extract_csv(file_path):
    """Extract data from CSV files"""
    try:
        import csv
        
        rows = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            csv_reader = csv.reader(file)
            for row in csv_reader:
                if any(cell.strip() for cell in row):  # Skip empty rows
                    rows.append(" | ".join(row))
        
        full_text = "\n".join(rows)
        
        return {
            'success': True,
            'text': full_text,
            'data': {'num_rows': len(rows)},
            'file_type': 'csv',
            'error': None
        }
    
    except Exception as e:
        return {
            'success': False,
            'error': f'CSV extraction failed: {str(e)}',
            'text': '',
            'data': None,
            'file_type': 'csv'
        }


def extract_txt(file_path):
    """Extract text from plain text files"""
    try:
        # Try multiple encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        
        text = None
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            # Fall back to binary mode with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
        
        return {
            'success': True,
            'text': text,
            'data': {
                'num_lines': len(text.split('\n')),
                'num_chars': len(text)
            },
            'file_type': 'txt',
            'error': None
        }
    
    except Exception as e:
        return {
            'success': False,
            'error': f'Text extraction failed: {str(e)}',
            'text': '',
            'data': None,
            'file_type': 'txt'
        }


def extract_image(file_path):
    """Extract text from images using OCR (Tesseract)"""
    try:
        from PIL import Image
        import pytesseract
        
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        
        return {
            'success': True,
            'text': text,
            'data': {
                'image_size': image.size,
                'image_mode': image.mode
            },
            'file_type': 'image',
            'error': None
        }
    
    except ImportError:
        return {
            'success': False,
            'error': 'PIL or pytesseract not installed. OCR not available.',
            'text': '[Image file - OCR not available. Content cannot be extracted.]',
            'data': None,
            'file_type': 'image'
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'Image OCR failed: {str(e)}',
            'text': '[Image file - OCR failed. Content cannot be extracted.]',
            'data': None,
            'file_type': 'image'
        }


def get_file_summary(file_path, max_chars=500):
    """
    Get a brief summary of file content for preview.
    Useful for showing user what was extracted.
    """
    result = extract_file_content(file_path)
    
    if not result['success']:
        return f"⚠️ Could not read file: {result['error']}"
    
    text = result['text']
    
    if len(text) <= max_chars:
        return text
    
    return text[:max_chars] + f"\n\n... (truncated, {len(text)} total characters)"


def extract_multiple_files(file_paths):
    """
    Extract content from multiple files and combine.
    
    Returns dict with:
        - success: bool
        - files: list of dicts with individual results
        - combined_text: str (all files combined)
        - total_chars: int
    """
    results = []
    combined_text = []
    total_chars = 0
    all_success = True
    
    for file_path in file_paths:
        result = extract_file_content(file_path)
        results.append({
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'result': result
        })
        
        if result['success']:
            file_header = f"\n{'='*60}\nFILE: {os.path.basename(file_path)}\n{'='*60}\n"
            combined_text.append(file_header + result['text'])
            total_chars += len(result['text'])
        else:
            all_success = False
    
    return {
        'success': all_success,
        'files': results,
        'combined_text': "\n\n".join(combined_text),
        'total_chars': total_chars,
        'num_files': len(file_paths)
    }


# I did no harm and this file is not truncated
