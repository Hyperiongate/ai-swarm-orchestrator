"""
File Analysis Agent
Created: January 29, 2026
Last Updated: January 29, 2026

Specialized AI agent for file-related tasks:
- Analyze uploaded files
- Reformat documents professionally
- Extract insights from data files
- Create new documents in specific folders

USAGE:
    from file_analysis_agent import get_file_analysis_agent
    agent = get_file_analysis_agent()
    result = agent.analyze_file(file_path, user_request)

Author: Jim @ Shiftwork Solutions LLC
"""

import os
import json
from datetime import datetime
from file_content_reader import extract_file_content, extract_multiple_files


class FileAnalysisAgent:
    """
    Specialized agent for file analysis and document creation tasks.
    """
    
    def __init__(self):
        self.name = "FileAnalysisAgent"
        self.capabilities = [
            "analyze_uploaded_files",
            "reformat_documents",
            "extract_data_insights",
            "create_professional_documents",
            "save_to_specific_folders"
        ]
    
    def analyze_file(self, file_path, user_request, ai_function):
        """
        Analyze a single file based on user request.
        
        Args:
            file_path: Path to the file
            user_request: What the user wants to know
            ai_function: AI function to call (e.g., call_claude_sonnet)
        
        Returns:
            dict with analysis results
        """
        print(f"📄 FileAnalysisAgent: Analyzing {os.path.basename(file_path)}")
        
        # Extract file content
        content_result = extract_file_content(file_path)
        
        if not content_result['success']:
            return {
                'success': False,
                'error': f"Could not read file: {content_result['error']}",
                'analysis': None
            }
        
        file_text = content_result['text']
        file_type = content_result['file_type']
        
        # Truncate if too long (keep within token limits)
        max_chars = 15000
        if len(file_text) > max_chars:
            file_text = file_text[:max_chars] + f"\n\n... [File truncated. Total length: {len(file_text)} characters]"
        
        # Build analysis prompt
        prompt = f"""You are analyzing a {file_type.upper()} file for the user.

FILE NAME: {os.path.basename(file_path)}
FILE TYPE: {file_type.upper()}

FILE CONTENT:
{file_text}

USER REQUEST: {user_request}

Please provide a comprehensive analysis addressing the user's request. Be specific, reference actual content from the file, and provide actionable insights."""
        
        # Call AI
        try:
            response = ai_function(prompt, max_tokens=4000)
            
            if isinstance(response, dict):
                analysis_text = response.get('content', str(response))
            else:
                analysis_text = str(response)
            
            return {
                'success': True,
                'analysis': analysis_text,
                'file_name': os.path.basename(file_path),
                'file_type': file_type,
                'file_size': len(file_text),
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"AI analysis failed: {str(e)}",
                'analysis': None
            }
    
    def analyze_multiple_files(self, file_paths, user_request, ai_function):
        """
        Analyze multiple files together.
        """
        print(f"📚 FileAnalysisAgent: Analyzing {len(file_paths)} files")
        
        # Extract all files
        multi_result = extract_multiple_files(file_paths)
        
        if not multi_result['success']:
            return {
                'success': False,
                'error': 'Some files could not be read',
                'analysis': None,
                'files': multi_result['files']
            }
        
        combined_text = multi_result['combined_text']
        
        # Truncate if too long
        max_chars = 15000
        if len(combined_text) > max_chars:
            combined_text = combined_text[:max_chars] + f"\n\n... [Content truncated. Total length: {multi_result['total_chars']} characters]"
        
        # Build analysis prompt
        file_list = ", ".join([os.path.basename(fp) for fp in file_paths])
        
        prompt = f"""You are analyzing {len(file_paths)} files together for the user.

FILES: {file_list}

COMBINED CONTENT:
{combined_text}

USER REQUEST: {user_request}

Please provide a comprehensive analysis addressing the user's request. Compare and contrast information across files where relevant, identify patterns, and provide actionable insights."""
        
        # Call AI
        try:
            response = ai_function(prompt, max_tokens=4000)
            
            if isinstance(response, dict):
                analysis_text = response.get('content', str(response))
            else:
                analysis_text = str(response)
            
            return {
                'success': True,
                'analysis': analysis_text,
                'files_analyzed': [os.path.basename(fp) for fp in file_paths],
                'num_files': len(file_paths),
                'total_chars': multi_result['total_chars'],
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"AI analysis failed: {str(e)}",
                'analysis': None
            }
    
    def reformat_document(self, source_file_path, user_request, ai_function, output_format='docx'):
        """
        Reformat a document professionally.
        
        Args:
            source_file_path: Original document
            user_request: How to reformat (e.g., "make more professional", "executive summary style")
            ai_function: AI function to call
            output_format: Output file format (docx, pdf, txt)
        
        Returns:
            dict with reformatted content and file path
        """
        print(f"✨ FileAnalysisAgent: Reformatting {os.path.basename(source_file_path)}")
        
        # Extract source content
        content_result = extract_file_content(source_file_path)
        
        if not content_result['success']:
            return {
                'success': False,
                'error': f"Could not read source file: {content_result['error']}",
                'output_path': None
            }
        
        source_text = content_result['text']
        
        # Build reformatting prompt
        prompt = f"""You are reformatting a document professionally.

ORIGINAL FILE: {os.path.basename(source_file_path)}

ORIGINAL CONTENT:
{source_text}

REFORMATTING REQUEST: {user_request}

Please reformat this document according to the request. Maintain all important information but improve:
- Professional tone and structure
- Clarity and readability
- Formatting and organization
- Grammar and style

Provide the complete reformatted document."""
        
        # Call AI
        try:
            response = ai_function(prompt, max_tokens=4000)
            
            if isinstance(response, dict):
                reformatted_text = response.get('content', str(response))
            else:
                reformatted_text = str(response)
            
            # Create output file
            output_result = self._create_document_file(
                content=reformatted_text,
                filename=f"reformatted_{os.path.basename(source_file_path)}",
                file_format=output_format,
                title=f"Reformatted: {os.path.basename(source_file_path)}"
            )
            
            if not output_result['success']:
                return {
                    'success': False,
                    'error': f"Could not create output file: {output_result['error']}",
                    'output_path': None,
                    'content': reformatted_text
                }
            
            return {
                'success': True,
                'output_path': output_result['file_path'],
                'output_filename': output_result['filename'],
                'content': reformatted_text,
                'file_format': output_format,
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"Reformatting failed: {str(e)}",
                'output_path': None
            }
    
    def create_document_from_request(self, user_request, ai_function, project_id=None, 
                                    file_format='docx', save_to_folder=None):
        """
        Create a new document from scratch based on user request.
        Can save to a specific project folder.
        
        Args:
            user_request: What document to create
            ai_function: AI function to call
            project_id: Optional project ID to save to
            file_format: Output format (docx, pdf, txt, xlsx)
            save_to_folder: Optional specific folder path
        
        Returns:
            dict with created document info
        """
        print(f"📝 FileAnalysisAgent: Creating document from request")
        
        # Build document creation prompt
        prompt = f"""Create a professional document based on this request:

REQUEST: {user_request}

Please create a complete, well-structured document that fully addresses this request.
Include all relevant sections, professional formatting, and comprehensive content.
This document will be saved as a {file_format.upper()} file."""
        
        # Call AI
        try:
            response = ai_function(prompt, max_tokens=4000)
            
            if isinstance(response, dict):
                document_text = response.get('content', str(response))
            else:
                document_text = str(response)
            
            # Generate filename from request
            filename = self._generate_filename_from_request(user_request, file_format)
            
            # Determine save location
            if save_to_folder:
                output_dir = save_to_folder
            elif project_id:
                output_dir = f"/tmp/projects/{project_id}"
            else:
                output_dir = "/tmp"
            
            os.makedirs(output_dir, exist_ok=True)
            
            # Create the document file
            output_result = self._create_document_file(
                content=document_text,
                filename=filename,
                file_format=file_format,
                output_dir=output_dir,
                title=self._generate_title_from_request(user_request)
            )
            
            if not output_result['success']:
                return {
                    'success': False,
                    'error': f"Could not create document: {output_result['error']}",
                    'output_path': None,
                    'content': document_text
                }
            
            return {
                'success': True,
                'output_path': output_result['file_path'],
                'output_filename': output_result['filename'],
                'output_dir': output_dir,
                'content': document_text,
                'file_format': file_format,
                'project_id': project_id,
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"Document creation failed: {str(e)}",
                'output_path': None
            }
    
    def _create_document_file(self, content, filename, file_format='docx', 
                            output_dir='/tmp', title=None):
        """
        Internal method to create actual document files.
        """
        try:
            # Clean filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            base_name = os.path.splitext(filename)[0]
            filename = f"{base_name}_{timestamp}.{file_format}"
            
            output_path = os.path.join(output_dir, filename)
            
            if file_format == 'docx':
                return self._create_docx(content, output_path, title)
            elif file_format == 'txt':
                return self._create_txt(content, output_path)
            elif file_format == 'xlsx':
                return self._create_xlsx(content, output_path, title)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported format: {file_format}',
                    'file_path': None
                }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'file_path': None
            }
    
    def _create_docx(self, content, output_path, title=None):
        """Create a Word document"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            
            doc = Document()
            
            # Add title
            if title:
                heading = doc.add_heading(title, 0)
                heading.alignment = 1  # Center
            
            # Add Shiftwork Solutions header
            company_para = doc.add_paragraph('Shiftwork Solutions LLC')
            company_para.alignment = 1
            company_para.runs[0].font.size = Pt(10)
            company_para.runs[0].font.color.rgb = RGBColor(102, 126, 234)
            
            doc.add_paragraph()  # Spacing
            
            # Parse content and add to document
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check for headers (markdown style)
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    para = doc.add_paragraph(line)
                    para.style.font.size = Pt(11)
            
            # Save
            doc.save(output_path)
            
            return {
                'success': True,
                'file_path': output_path,
                'filename': os.path.basename(output_path),
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f'DOCX creation failed: {str(e)}',
                'file_path': None
            }
    
    def _create_txt(self, content, output_path):
        """Create a text file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                'success': True,
                'file_path': output_path,
                'filename': os.path.basename(output_path),
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f'TXT creation failed: {str(e)}',
                'file_path': None
            }
    
    def _create_xlsx(self, content, output_path, title=None):
        """Create an Excel spreadsheet from content"""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill
            
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = title[:31] if title else "Data"
            
            # Parse content into rows
            lines = content.split('\n')
            row_num = 1
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Try to split by common delimiters
                if '|' in line:
                    cells = [cell.strip() for cell in line.split('|')]
                elif '\t' in line:
                    cells = [cell.strip() for cell in line.split('\t')]
                else:
                    cells = [line]
                
                for col_num, cell_value in enumerate(cells, start=1):
                    sheet.cell(row=row_num, column=col_num, value=cell_value)
                
                row_num += 1
            
            # Style header row
            if row_num > 1:
                for cell in sheet[1]:
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="667EEA", end_color="667EEA", fill_type="solid")
            
            workbook.save(output_path)
            
            return {
                'success': True,
                'file_path': output_path,
                'filename': os.path.basename(output_path),
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f'XLSX creation failed: {str(e)}',
                'file_path': None
            }
    
    def _generate_filename_from_request(self, request, file_format):
        """Generate a sensible filename from user request"""
        # Take first 5 words
        words = request.split()[:5]
        base = "_".join(words)
        
        # Clean up
        base = "".join(c if c.isalnum() or c == '_' else '_' for c in base)
        base = base.lower()
        
        return f"{base}.{file_format}"
    
    def _generate_title_from_request(self, request):
        """Generate a document title from user request"""
        # Capitalize first letter of each word
        words = request.split()[:8]
        title = " ".join(words)
        
        if len(title) > 60:
            title = title[:57] + "..."
        
        return title.title()


# Singleton instance
_file_analysis_agent = None

def get_file_analysis_agent():
    """Get the FileAnalysisAgent singleton instance"""
    global _file_analysis_agent
    if _file_analysis_agent is None:
        _file_analysis_agent = FileAnalysisAgent()
    return _file_analysis_agent


# I did no harm and this file is not truncated
