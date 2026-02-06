"""
Orchestration Handler - Main AI Task Processing
Created: January 31, 2026
Last Updated: February 5, 2026 - FIXED MEMORY CRASH ON LARGE EXCEL FILES

This file handles the main /api/orchestrate endpoint that processes user requests.
Separated from core.py to make it easier to fix and maintain.

CRITICAL KNOWN ISSUE - February 5, 2026:
⚠️ FILE SELECTION vs FILE UPLOAD behavior difference:
- FILE UPLOAD (drag & drop): Triggers progressive analysis, uses consulting prompts ✅
- FILE SELECTION (from project): Bypasses progressive analysis, limited extraction ❌
  * Only reads first worksheet
  * Limited to 10K chars per file  
  * Needs fix in database_file_management.py
  
WORKAROUND: For large Excel files (5MB+), use FILE UPLOAD instead of selecting from project!

UPDATES:
- February 5, 2026 (v6): CRITICAL FIX - Memory crash on large Excel cloud downloads
  * Removed redundant pd.ExcelFile() call in handle_large_excel_initial() that loaded
    entire 26MB+ file into memory AGAIN after chunk extraction already succeeded
  * Now uses sheet_names from chunk_result (provided by progressive_file_analyzer v5)
  * This was causing 500 errors after successful Google Drive download
  * Also fixed: cloud-downloaded files now use correct user_request (not the URL wrapper)
- February 5, 2026: INCREASED PROGRESSIVE ANALYSIS from 100 to 500 rows for initial analysis
  * Large files now get 5x more data in first pass (100 → 500 rows)
  * Provides much deeper operational insights from the start
  * Still allows user to request more rows if needed
  * Fixes the "compacting" issue where analysis was too shallow
- February 5, 2026: FIXED GPT-4 token limit (was 8000, max is 4096)
- February 5, 2026: RUTHLESS ANALYSIS ENHANCEMENT - No more shallow summaries!
  * Prompt now demands specific numbers, percentages, and dollar amounts
  * Forces operational insights (shift patterns, coverage gaps, overtime risks)
  * Requires actionable recommendations with quantified impact
  * Checks ALL worksheets in Excel files (not just first sheet)
  * Increased tokens: 4000→6000 (progressive), 8000 (full file)
  * Analysis tone: "You are Jim Goodwin" - consulting engagement mindset
- February 5, 2026: MAJOR ENHANCEMENT - Replaced weak generic analysis prompts with powerful 
  consulting-grade prompts that leverage Shiftwork Solutions' 30+ years of expertise
  * Large Excel files now get 7-point deep analysis (not surface observations)
  * Increased max_tokens from 4000 to 8000 for detailed insights
  * Analysis now includes: statistical depth, operational patterns, cost insights,
    red flags, and actionable recommendations with specific numbers
- February 5, 2026: FIXED syntax error - completed try/except block around Excel size detection
- February 5, 2026: FIXED syntax error (duplicate code blocks, bracket mismatch) 
- February 4, 2026: Added auto-learning integration at all conversation endpoints
  * GPT-4 file handler now learns from file analysis conversations
  * Main orchestration endpoint learns from all AI conversations
  * System builds cumulative intelligence from every interaction
- February 1, 2026: FIXED file_contents UnboundLocalError in file browser
  * Initialize file_contents early to prevent crashes
  * Added complete file_ids handling for project file selection
  * File browser now fully functional!
- January 31, 2026: File upload contents now properly passed to Claude
- January 31, 2026: Added progressive analysis for large Excel files (hybrid approach)
  * Files under 5MB: Full analysis
  * Files 5-25MB: Analyze first 500 rows, ask user if they want more
  * Files over 25MB: Rejected with helpful message
  * User can request: "next 500", "next 1000", "analyze all"

Author: Jim @ Shiftwork Solutions LLC
"""

from flask import Blueprint, request, jsonify, session
import time
import json
import os
from datetime import datetime
from werkzeug.utils import secure_filename

# Import all the utilities we need
from database import (
    get_db, create_conversation, get_conversation, add_message,
    get_conversation_context, save_generated_document,
    get_schedule_context, save_schedule_context,
    get_client_profile_context, add_avoidance_pattern, get_avoidance_context,
    update_client_profile
)
from orchestration import (
    analyze_task_with_sonnet,
    handle_with_opus,
    execute_specialist_task,
    validate_with_consensus
)
from database_file_management import get_files_for_ai_context, get_file_stats_by_project
from file_content_reader import extract_multiple_files
from code_assistant_agent import get_code_assistant
from orchestration.proactive_agent import ProactiveAgent
from schedule_request_handler_combined import get_combined_schedule_handler
from progressive_file_analyzer import get_progressive_analyzer
from conversation_learning import learn_from_conversation  # Auto-learning from conversations
from cloud_file_handler import get_cloud_handler

# ============================================================================
from background_file_processor import get_background_processor
import uuid
# LEARNING LOOP ENHANCEMENT - February 4, 2026
# Import learning systems to close the loop
# ============================================================================
from orchestration.task_analysis import get_learning_context
from enhanced_intelligence import EnhancedIntelligence
from specialized_knowledge import get_specialized_knowledge
from proactive_suggestions import get_proactive_suggestions
from conversation_summarizer import get_conversation_summarizer
from proactive_curiosity_engine import get_curiosity_engine  # Phase 1: Proactive Curiosity

# Create blueprint
orchestration_bp = Blueprint('orchestration', __name__)


def convert_markdown_to_html(text):
    """Convert markdown text to styled HTML"""
    if not text:
        return text
    import markdown
    html = markdown.markdown(text, extensions=['extra', 'nl2br'])
    return f'<div style="line-height: 1.8; color: #333;">{html}</div>'

def is_cloud_link(text):
    """Check if text contains a cloud storage link"""
    if not text:
        return False
    text_lower = text.lower()
    cloud_indicators = [
        'drive.google.com',
        'docs.google.com',
        'dropbox.com',
        'onedrive',
        '1drv.ms'
    ]
    return any(indicator in text_lower for indicator in cloud_indicators)

def should_create_document(user_request):
    """Determine if we should create a downloadable document"""
    doc_keywords = ['create', 'generate', 'make', 'write', 'draft', 'prepare',
                    'document', 'report', 'proposal', 'presentation', 'schedule',
                    'contract', 'agreement', 'summary', 'analysis']
    request_lower = user_request.lower()
    for keyword in doc_keywords:
        if keyword in request_lower:
            if 'presentation' in request_lower or 'powerpoint' in request_lower or 'slides' in request_lower:
                return True, 'pptx'
            elif 'spreadsheet' in request_lower or 'excel' in request_lower:
                return True, 'xlsx'
            elif 'pdf' in request_lower:
                return True, 'pdf'
            else:
                return True, 'docx'
    return False, None


def get_knowledge_context_for_prompt(knowledge_base, user_request, max_context=3000):
    """Get knowledge context to include in completion prompts"""
    if not knowledge_base:
        return ""
    try:
        context = knowledge_base.get_context_for_task(user_request, max_context=max_context)
        if context:
            return f"\n\n=== SHIFTWORK SOLUTIONS KNOWLEDGE BASE ===\nUse this information from our 30+ years of expertise:\n\n{context}\n\n=== END KNOWLEDGE BASE ===\n\n"
        return ""
    except Exception as e:
        print(f"Knowledge context retrieval failed: {e}")
        return ""


def generate_document_title(user_request, doc_type):
    """Generate a human-readable title from the user request"""
    title = user_request.strip()
    if title:
        title = title[0].upper() + title[1:]
    if len(title) > 60:
        title = title[:57] + '...'
    return title


def categorize_document(user_request, doc_type):
    """Determine document category based on request content"""
    request_lower = user_request.lower()
    if any(word in request_lower for word in ['schedule', 'dupont', 'panama', 'pitman', '2-2-3']):
        return 'schedule'
    elif any(word in request_lower for word in ['proposal', 'bid', 'quote']):
        return 'proposal'
    elif any(word in request_lower for word in ['report', 'analysis', 'summary']):
        return 'report'
    elif any(word in request_lower for word in ['contract', 'agreement', 'legal']):
        return 'contract'
    elif any(word in request_lower for word in ['checklist', 'list', 'todo']):
        return 'checklist'
    elif any(word in request_lower for word in ['email', 'letter', 'memo']):
        return 'communication'
    elif any(word in request_lower for word in ['presentation', 'slides', 'deck']):
        return 'presentation'
    else:
        return 'general'


@orchestration_bp.route('/api/orchestrate', methods=['POST'])
def orchestrate():
    """
    Main orchestration endpoint with proactive intelligence and conversation memory.
    
    UPDATED February 5, 2026 (v6): Fixed memory crash on large Excel cloud downloads
    UPDATED February 5, 2026: Increased progressive analysis from 100 to 500 rows
    UPDATED February 5, 2026: Fixed syntax error (incomplete try block for Excel detection)
    UPDATED February 4, 2026: Added auto-learning integration
    UPDATED January 31, 2026: Extracted to separate file for better maintainability
    FIXED January 31, 2026: File contents now properly passed to Claude AI
    """
    try:
        # Parse request data (JSON or FormData)
        if request.is_json:
            data = request.json
            user_request = data.get('request')
            enable_consensus = data.get('enable_consensus', True)
            project_id = data.get('project_id')
            conversation_id = data.get('conversation_id')
            mode = data.get('mode', 'quick')
            file_paths = []
        else:
            user_request = request.form.get('request')
            enable_consensus = request.form.get('enable_consensus', 'true').lower() == 'true'
            project_id = request.form.get('project_id')
            conversation_id = request.form.get('conversation_id')
            mode = request.form.get('mode', 'quick')
            
            # Handle file uploads
            file_paths = []
            if 'files' in request.files:
                files = request.files.getlist('files')
                
                if project_id:
                    upload_dir = f'/tmp/projects/{project_id}'
                else:
                    upload_dir = f'/tmp/uploads'
                os.makedirs(upload_dir, exist_ok=True)
                
                for file in files:
                    if file and file.filename:
                        filename = secure_filename(file.filename)
                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                        name, ext = os.path.splitext(filename)
                        filename = f"{name}_{timestamp}{ext}"
                        file_path = os.path.join(upload_dir, filename)
                        file.save(file_path)
                        file_paths.append(file_path)
                        print(f"📎 Saved uploaded file: {filename}")
        
        # ====================================================================
        # CRITICAL FIX February 1, 2026: Initialize file_contents early
        # This prevents UnboundLocalError when file_ids are used
        # ====================================================================
        file_contents = ""
        
        # ====================================================================
        # CLOUD LINK DETECTION - February 5, 2026
        # Check if user provided a cloud storage link instead of file upload
        # Downloads file using streaming to prevent RAM crashes
        # ====================================================================
        if is_cloud_link(user_request):
            print(f"🔗 Cloud storage link detected in request")
            
            # Extract URL from request
            import re
            url_pattern = r'https?://[^\s]+'
            urls = re.findall(url_pattern, user_request)
            
            if urls:
                cloud_url = urls[0]  # Use first URL found
                print(f"🔗 Processing cloud link: {cloud_url[:50]}...")
                
                # Create conversation if needed
                if not conversation_id:
                    conversation_id = create_conversation(mode=mode, project_id=project_id)
                
                # Download file from cloud using STREAMING (prevents RAM crash!)
                handler = get_cloud_handler()
                local_filepath, metadata = handler.process_cloud_link(cloud_url)
                
                if not metadata['success']:
                    return jsonify({
                        'success': False,
                        'error': f"Could not download file from {metadata.get('service', 'cloud storage')}: {metadata.get('error', 'Unknown error')}",
                        'conversation_id': conversation_id
                    }), 400
                
                # File downloaded successfully - add to file_paths for processing
                file_paths = [local_filepath]
                print(f"✅ Downloaded {metadata['size_bytes'] / (1024*1024):.1f}MB from {metadata['service']}")
                
                # ================================================================
                # FIX February 5, 2026 (v6): Clean user_request for cloud links
                # The frontend wraps the URL in "Please analyze this file: <URL>"
                # Strip the URL so the AI gets a clean analysis request
                # ================================================================
                clean_request = re.sub(url_pattern, '', user_request).strip()
                if clean_request and clean_request.lower() not in ['please analyze this file:', 'please analyze this file', 'analyze this file:']:
                    user_request = clean_request
                else:
                    # Default request when user just pasted a link
                    file_basename = os.path.basename(local_filepath)
                    user_request = f"Please analyze this file: {metadata.get('original_filename', file_basename)}"
                print(f"📝 Clean user request: {user_request}")
                # ================================================================
                
                # Continue to normal file processing below...
        # ====================================================================
        # END CLOUD LINK DETECTION
        # ====================================================================
        
        # ====================================================================
        # FILE BROWSER SUPPORT - Handle file_ids from project file selection
        # Added: January 31, 2026
        # Fixed: February 1, 2026 - file_contents initialized above
        # ====================================================================
        
        # Check if user selected files from project (file_ids parameter)
        file_ids_param = None
        if request.is_json:
            file_ids_param = data.get('file_ids')
        else:
            file_ids_param = request.form.get('file_ids')
        
        # 🔍 DIAGNOSTIC STEP 1: Log initial file_ids detection
        print(f"🔍 DIAGNOSTIC STEP 1: file_ids_param = {file_ids_param}")
        print(f"🔍 DIAGNOSTIC STEP 1: file_ids_param type = {type(file_ids_param)}")
        print(f"🔍 DIAGNOSTIC STEP 1: project_id = {project_id}")
        print(f"🔍 DIAGNOSTIC STEP 1: request.is_json = {request.is_json}")
        
        if file_ids_param and project_id:
            try:
                # Parse file_ids (comes as JSON string)
                if isinstance(file_ids_param, str):
                    import json
                    file_ids = json.loads(file_ids_param)
                else:
                    file_ids = file_ids_param
                
                # 🔍 DIAGNOSTIC STEP 2: Log parsed file_ids
                print(f"🔍 DIAGNOSTIC STEP 2: Parsed file_ids = {file_ids}")
                print(f"🔍 DIAGNOSTIC STEP 2: Number of files = {len(file_ids) if file_ids else 0}")
                
                if file_ids and len(file_ids) > 0:
                    print(f"📎 User selected {len(file_ids)} file(s) from project {project_id}")
                    
                    # Fetch file context from database
                    from database_file_management import get_files_for_ai_context
                    
                    # Get detailed file info with actual file contents
                    selected_file_context = get_files_for_ai_context(
                        project_id=project_id, 
                        file_ids=file_ids,  # Pass specific file IDs
                        max_files=len(file_ids),  # Get all selected files
                        max_chars_per_file=10000  # Allow up to 10k chars per file
                    )
                    
                    # 🔍 DIAGNOSTIC STEP 3: Log retrieved context
                    print(f"🔍 DIAGNOSTIC STEP 3: selected_file_context returned")
                    print(f"🔍 DIAGNOSTIC STEP 3: Context length = {len(selected_file_context) if selected_file_context else 0}")
                    if selected_file_context:
                        print(f"🔍 DIAGNOSTIC STEP 3: Context preview (first 200 chars):")
                        print(f"{selected_file_context[:200]}")
                    else:
                        print(f"🔍 DIAGNOSTIC STEP 3: Context is EMPTY/NONE")
                    
                    if selected_file_context:
                        print(f"✅ Retrieved context for {len(file_ids)} selected file(s)")
                        
                        # Add to file_contents for AI processing
                        file_contents += "\n\n" + selected_file_context
                        
                        # 🔍 DIAGNOSTIC STEP 3 (continued): Log file_contents after adding
                        print(f"🔍 DIAGNOSTIC STEP 3: file_contents now has {len(file_contents)} chars")
                        print(f"🔍 DIAGNOSTIC STEP 3: file_contents preview (first 200 chars):")
                        print(f"{file_contents[:200]}")
                    else:
                        print(f"⚠️ No file context retrieved for file_ids: {file_ids}")
                        
            except Exception as file_ids_error:
                print(f"⚠️ Error processing file_ids: {file_ids_error}")
                import traceback
                traceback.print_exc()
        
        # ====================================================================
        # END FILE BROWSER SUPPORT
        # ====================================================================
        
        if not user_request:
            return jsonify({'success': False, 'error': 'Request text required'}), 400
        
        if file_paths:
            print(f"📎 {len(file_paths)} file(s) attached to request")
        
        overall_start = time.time()
        
        # ====================================================================
        # PROGRESSIVE FILE ANALYSIS - New Feature (January 31, 2026)
        # ====================================================================
        
        # Check if this is a continuation request (user asking for more rows)
        analyzer = get_progressive_analyzer()
        continuation_request = analyzer.parse_user_continuation_request(user_request)
        
        if continuation_request and conversation_id:
            # User is requesting more data from a previous file upload
            # Retrieve file path from session
            file_analysis_state = session.get(f'file_analysis_{conversation_id}')
            
            if file_analysis_state:
                return handle_progressive_continuation(
                    conversation_id=conversation_id,
                    user_request=user_request,
                    continuation_request=continuation_request,
                    file_analysis_state=file_analysis_state,
                    overall_start=overall_start
                )
        
        # Check for large Excel files that need progressive analysis
        if file_paths:
            for file_path in file_paths:
                file_info = analyzer.get_file_info(file_path)
                
                print(f"📊 FILE SIZE DETECTION:")
                print(f"   - File: {os.path.basename(file_path)}")
                print(f"   - Size: {file_info.get('file_size_mb', 0):.1f}MB")
                print(f"   - is_large: {file_info.get('is_large')}")
                print(f"   - is_too_large: {file_info.get('is_too_large')}")
                print(f"   - file_type: {file_info.get('file_type')}")
                
                # Reject files that are too large
                if file_info.get('is_too_large'):
                    file_size_mb = file_info.get('file_size_mb', 0)
                    return jsonify({
                        'success': False,
                        'error': f'File too large ({file_size_mb:.1f}MB). Maximum file size is 25MB.',
                        'suggestion': 'Please reduce file size by:\n- Filtering to specific date range\n- Removing unnecessary columns\n- Splitting into multiple smaller files\n- Exporting only relevant sheets'
                    }), 413
                
                # Handle VERY large files with background processing (>50MB)
                if file_info.get('file_size_mb', 0) > 50 and file_info.get('file_type') in ['.xlsx', '.xls']:
                    print(f"🔄 TRIGGERING BACKGROUND PROCESSING for {os.path.basename(file_path)}")
                    
                    # Create conversation and task
                    if not conversation_id:
                        conversation_id = create_conversation(mode=mode, project_id=project_id)
                    
                    db = get_db()
                    cursor = db.execute('INSERT INTO tasks (user_request, status, conversation_id) VALUES (?, ?, ?)',
                                       (user_request, 'processing_background', conversation_id))
                    task_id = cursor.lastrowid
                    db.commit()
                    db.close()
                    
                    # Submit to background processor
                    processor = get_background_processor()
                    result = processor.submit_job(
                        job_id=f"JOB_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}",
                        file_path=file_path,
                        user_request=user_request,
                        conversation_id=conversation_id,
                        task_id=task_id,
                        user_name="User"
                    )
                    
                    if result['success']:
                        # Post initial message
                        initial_msg = f"""🔄 **Processing large file in background**

**File:** {os.path.basename(file_path)} ({file_info['file_size_mb']}MB)
**Estimated time:** ~{result['estimated_minutes']} minutes

I'm analyzing all rows in the background. I'll post the complete analysis here when finished.

You can continue using the app while I work on this!"""
                        
                        add_message(conversation_id, 'assistant', initial_msg, task_id,
                                   {'orchestrator': 'background_processor', 'job_id': result['job_id']})
                        
                        return jsonify({
                            'success': True,
                            'task_id': task_id,
                            'conversation_id': conversation_id,
                            'result': convert_markdown_to_html(initial_msg),
                            'orchestrator': 'background_processor',
                            'job_id': result['job_id'],
                            'background_processing': True,
                            'estimated_minutes': result['estimated_minutes']
                        })
                    else:
                        # Background submission failed - fall back to progressive
                        print(f"⚠️ Background processing failed: {result.get('error')}, falling back to progressive")
                        return handle_large_excel_initial(
                            file_path=file_path,
                            user_request=user_request,
                            conversation_id=conversation_id,
                            project_id=project_id,
                            mode=mode,
                            file_info=file_info,
                            overall_start=overall_start
                        )
                
                # Handle moderately large files with progressive analysis (5-50MB)
                elif file_info.get('is_large') and file_info.get('file_type') in ['.xlsx', '.xls']:
                    print(f"🔄 TRIGGERING PROGRESSIVE ANALYSIS for {os.path.basename(file_path)}")
                    return handle_large_excel_initial(
                        file_path=file_path,
                        user_request=user_request,
                        conversation_id=conversation_id,
                        project_id=project_id,
                        mode=mode,
                        file_info=file_info,
                        overall_start=overall_start
                    )
        
        # ====================================================================
        # STANDARD FILE HANDLING (Small files - under 5MB)
        # ====================================================================
        
        # Extract file contents from uploaded files (APPEND to file_contents)
        # NOTE: file_contents already initialized above and may contain file_ids context
        if file_paths:
            try:
                extracted = extract_multiple_files(file_paths)
                if extracted['success'] and extracted.get('combined_text'):
                    extracted_text = extracted['combined_text']
                    
                    # Check if extracted content is too long (over 50,000 chars)
                    # UPDATED February 5, 2026: Increased from 50,000 to 200,000 chars for large files
                    if len(extracted_text) > 200000:
                        print(f"⚠️ File content very large ({len(extracted_text)} chars) - truncating")
                        extracted_text = extracted_text[:200000] + f"\n\n... (truncated {len(extracted_text) - 200000} characters for performance)"
                                    
                    # APPEND to file_contents (don't overwrite file_ids context!)
                    if file_contents:
                        file_contents += "\n\n" + extracted_text
                    else:
                        file_contents = extracted_text
                else:
                    print(f"⚠️ File extraction returned no content")
            except Exception as extract_error:
                print(f"⚠️ Could not extract file contents: {extract_error}")
        
        # 🔍 DIAGNOSTIC STEP 4: Check if file_contents will route to GPT-4
        print(f"🔍 DIAGNOSTIC STEP 4: Checking file_contents for GPT-4 routing")
        print(f"🔍 DIAGNOSTIC STEP 4: file_contents is truthy = {bool(file_contents)}")
        print(f"🔍 DIAGNOSTIC STEP 4: file_contents length = {len(file_contents) if file_contents else 0}")
        
        # Route file analysis to GPT-4 (better for file handling)
        if file_contents:
            print(f"🔍 DIAGNOSTIC STEP 4: ROUTING TO GPT-4! file_contents has {len(file_contents)} chars")
            print(f"📎 File content detected ({len(file_contents)} chars) - routing to GPT-4")
            
            if not conversation_id:
                conversation_id = create_conversation(mode=mode, project_id=project_id)
                print(f"Created new conversation: {conversation_id}")
            
            add_message(conversation_id, 'user', user_request, file_contents=file_contents if file_contents else None)
            
            db = get_db()
            cursor = db.execute('INSERT INTO tasks (user_request, status, conversation_id) VALUES (?, ?, ?)',
                               (user_request, 'processing', conversation_id))
            task_id = cursor.lastrowid
            db.commit()

            # ================================================================
            # LEARNING LOOP ENHANCEMENT - Initialize Intelligence Engine
            # February 4, 2026 - FOR ALL CONVERSATIONS
            # ================================================================
            intelligence = None
            try:
                intelligence = EnhancedIntelligence()
                print("🧠 EnhancedIntelligence initialized")
            except Exception as intel_error:
                print(f"⚠️ EnhancedIntelligence init failed (non-critical): {intel_error}")
                import traceback
                print(traceback.format_exc())
                intelligence = None  # Explicitly set to None on failure
            
            from orchestration.ai_clients import call_gpt4
            
            # ================================================================
            # FIXED February 5, 2026: Complete try/except block for Excel size detection
            # ================================================================
            # Determine if this is a large Excel file
            is_large_excel = False
            row_count_estimate = 0
            try:
                if file_paths:
                    for fp in file_paths:
                        if fp.endswith(('.xlsx', '.xls')):
                            # Estimate rows from file size (rough: 1MB ≈ 1000-2000 rows)
                            file_size_mb = os.path.getsize(fp) / (1024 * 1024)
                            row_count_estimate = int(file_size_mb * 1500)  # Conservative estimate
                            if file_size_mb > 5:
                                is_large_excel = True
                                break
            except Exception as excel_detect_error:
                print(f"⚠️ Excel size detection failed (non-critical): {excel_detect_error}")
                is_large_excel = False
                row_count_estimate = 0
            # ================================================================
            # END FIX
            # ================================================================

            if is_large_excel and row_count_estimate > 0:
                # Check all worksheets
                import pandas as pd
                try:
                    if file_paths and file_paths[0].endswith(('.xlsx', '.xls')):
                        excel_file = pd.ExcelFile(file_paths[0])
                        sheet_names = excel_file.sheet_names
                        num_sheets = len(sheet_names)
                        sheets_info = f"\n📋 **FILE CONTAINS {num_sheets} WORKSHEET(S):** {', '.join(sheet_names)}\nAnalyze ALL worksheets that contain relevant data.\n"
                    else:
                        sheets_info = ""
                except:
                    sheets_info = ""
                
                file_analysis_prompt = f"""You are Jim Goodwin, owner of Shiftwork Solutions LLC with 30+ years analyzing workforce operations for hundreds of clients.

This is a CONSULTING ENGAGEMENT worth $16,500/week. The client expects DEEP, ACTIONABLE ANALYSIS - not surface observations.

The user uploaded a LARGE Excel file (approximately {row_count_estimate:,} rows) and asked: {user_request}
{sheets_info}

CRITICAL INSTRUCTIONS - THIS IS A CONSULTING ENGAGEMENT:
Your analysis must be ACTIONABLE and SPECIFIC - not generic observations. Treat this like a $16,500/week consulting project.

REQUIRED DEPTH OF ANALYSIS:

1. DATA STRUCTURE & QUALITY:
   - Exact row/column counts and date ranges covered
   - Data quality issues (missing values, anomalies, data entry errors)
   - Granularity (hourly? daily? weekly?)
   - ALL WORKSHEETS reviewed (not just first sheet)

2. STATISTICAL ANALYSIS (with actual numbers):
   - Mean, median, std deviation for key metrics
   - Min/max values and when they occurred
   - Distribution patterns (normal? skewed? bimodal?)
   - Outliers with specific values and dates

3. OPERATIONAL PATTERNS (shiftwork expertise):
   - Day-of-week patterns: Which days are peaks? Troughs?
   - Time-based trends: Are hours increasing/decreasing over time?
   - Shift pattern detection: Do the patterns suggest 8hr, 10hr, 12hr shifts?
   - Coverage gaps: Where are the understaffing risks?
   - Overtime indicators: Where is excessive variability?

4. DEPARTMENT/CATEGORY ANALYSIS:
   - Rank order by total volume (with actual numbers)
   - Variance analysis: Which areas have the most volatility?
   - Workload balance: Are some areas overloaded vs others?
   - Correlation analysis: Do departments move together?

5. COST & EFFICIENCY INSIGHTS:
   - Total hours and estimated labor costs (assume $30/hr average)
   - Departments with highest cost exposure
   - Efficiency opportunities (high variance = scheduling problems)

6. RED FLAGS & RISKS:
   - Sustained high overtime indicators
   - Erratic patterns suggesting poor scheduling
   - Coverage gaps on critical days
   - Departments showing burnout risk patterns

7. ACTIONABLE RECOMMENDATIONS:
   - Top 3 specific improvements with expected impact
   - Which departments need schedule redesign first
   - Estimated cost savings from optimization
   - Next steps for implementation

FILE CONTENTS (representative sample from {len(file_contents):,} characters of data):

{file_contents}

DELIVER A CONSULTING-GRADE ANALYSIS WITH SPECIFIC NUMBERS, TRENDS, AND RECOMMENDATIONS.
This should read like a professional report you'd deliver to a client paying $16,500/week for your expertise."""
            else:
                file_analysis_prompt = f"""You are analyzing data for Shiftwork Solutions LLC, a consulting firm specializing in 24/7 operations optimization.

The user asked: {user_request}

Here are the file contents:

{file_contents}

Provide a DETAILED, SPECIFIC analysis with:
- Actual numbers and statistics (not vague descriptions)
- Operational patterns and trends
- Data quality observations
- Specific recommendations based on Shiftwork Solutions' 30+ years of expertise
- Actionable next steps

Be concrete and consulting-grade in your analysis."""

            try:
                gpt_response = call_gpt4(file_analysis_prompt, max_tokens=4000)
                
                if not gpt_response.get('error') and gpt_response.get('content'):
                    actual_output = gpt_response.get('content', '')
                    formatted_output = convert_markdown_to_html(actual_output)
                    
                    # ================================================================
                    # CREATE PROFESSIONAL DOCUMENT - February 1, 2026
                    # Use docx-js for professional formatting
                    # ================================================================
                    document_created = False
                    document_url = None
                    document_id = None
                    
                    try:
                        from document_creation_helper import create_analysis_document
                        
                        # Get analyzed file names
                        analyzed_files = []
                        if file_ids_param:
                            # Parse file_ids if needed
                            if isinstance(file_ids_param, str):
                                import json
                                file_ids = json.loads(file_ids_param)
                            else:
                                file_ids = file_ids_param
                            
                            from database_file_management import get_project_manager
                            pm = get_project_manager()
                            for fid in file_ids:
                                file_info = pm.get_file(fid)
                                if file_info:
                                    analyzed_files.append(file_info.get('original_filename', 'Unknown'))
                        
                        file_names_str = ", ".join(analyzed_files) if analyzed_files else "Uploaded Files"
                        
                        # Create document using Python (not Node.js!)
                        print(f"📄 Creating professional analysis document...")
                        doc_result = create_analysis_document(
                            analysis_text=actual_output,
                            file_names_str=file_names_str,
                            output_path='/tmp/file_analysis.docx'
                        )
                        
                        if doc_result['success'] and os.path.exists('/tmp/file_analysis.docx'):
                            # Save to database
                            file_size = os.path.getsize('/tmp/file_analysis.docx')
                            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                            filename = f'file_analysis_{timestamp}.docx'
                            
                            # Move to permanent location
                            final_path = f'/tmp/{filename}'
                            os.rename('/tmp/file_analysis.docx', final_path)
                            
                            document_id = save_generated_document(
                                filename=filename,
                                original_name=f"File Analysis - {file_names_str[:50]}",
                                document_type='docx',
                                file_path=final_path,
                                file_size=file_size,
                                task_id=task_id,
                                conversation_id=conversation_id,
                                project_id=project_id,
                                title=f"File Analysis Report",
                                description=f"AI analysis of {file_names_str}",
                                category='report'
                            )
                            
                            document_created = True
                            document_url = f'/api/generated-documents/{document_id}/download'
                            print(f"✅ Created professional analysis document: {filename}")
                        else:
                            print(f"⚠️ Document creation failed: {doc_result.get('error', 'Unknown error')}")
                    
                    except Exception as doc_error:
                        print(f"⚠️ Could not create analysis document: {doc_error}")
                        import traceback
                        traceback.print_exc()  
                    
                    # ================================================================
                    # END DOCUMENT CREATION
                    # ================================================================
                    
                    total_time = time.time() - overall_start
                    db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                              ('completed', 'gpt4_file_handler', total_time, task_id))
                    db.commit()
                    db.close()
                    
                    add_message(conversation_id, 'assistant', actual_output, task_id,
                               {'orchestrator': 'gpt4_file_handler', 'file_analysis': True, 'execution_time': total_time})
                 
                    # Auto-learn from this conversation
                    try:
                        learn_from_conversation(user_request, actual_output if actual_output else '')
                    except Exception as learn_error:
                        print(f"⚠️ Auto-learning failed (non-critical): {learn_error}")
                    
                    # ================================================================
                    # PHASE 1: PROACTIVE CURIOSITY ENGINE - February 5, 2026
                    # Ask curious follow-up questions after completing tasks
                    # ================================================================
                    curious_question = None
                    try:
                        curiosity_engine = get_curiosity_engine()
                        curiosity_check = curiosity_engine.should_be_curious(
                            conversation_id,
                            {
                                'user_request': user_request,
                                'ai_response': actual_output if actual_output else '',
                                'task_completed': True
                            }
                        )
                        
                        if curiosity_check['should_ask']:
                            curious_question = curiosity_check['question']
                            print(f"🤔 Curious follow-up: {curious_question}")
                    except Exception as curiosity_error:
                        print(f"⚠️ Curiosity engine failed (non-critical): {curiosity_error}")
                    # ================================================================
                    
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'conversation_id': conversation_id,
                        'result': formatted_output,
                        'orchestrator': 'gpt4_file_handler',
                        'execution_time': total_time,
                        'message': '📎 File analyzed by GPT-4',
                        'document_created': document_created,
                        'document_url': document_url,
                        'document_id': document_id,
                        'document_type': 'docx' if document_created else None,
                        'curious_question': curious_question
                    })
                   
                else:
                    error_msg = gpt_response.get('content', 'Unknown error')
                    print(f"⚠️ GPT-4 analysis failed: {error_msg}")
                    
                    # FALLBACK: Use Claude Sonnet for file analysis instead
                    print(f"🔄 Falling back to Claude Sonnet for file analysis")
                    
                    from orchestration.ai_clients import call_claude_sonnet
                    
                    # Use same consulting-grade prompt but with Claude
                    sonnet_response = call_claude_sonnet(file_analysis_prompt, max_tokens=4000)
                    
                    if isinstance(sonnet_response, dict) and not sonnet_response.get('error'):
                        actual_output = sonnet_response.get('content', '')
                        formatted_output = convert_markdown_to_html(actual_output)
                        
                        total_time = time.time() - overall_start
                        db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                                  ('completed', 'claude_sonnet_file_handler', total_time, task_id))
                        db.commit()
                        db.close()
                        
                        add_message(conversation_id, 'assistant', actual_output, task_id,
                                   {'orchestrator': 'claude_sonnet_file_handler', 'file_analysis': True, 
                                    'execution_time': total_time, 'fallback_from': 'gpt4'})
                        
                        return jsonify({
                            'success': True,
                            'task_id': task_id,
                            'conversation_id': conversation_id,
                            'result': formatted_output,
                            'orchestrator': 'claude_sonnet_file_handler',
                            'execution_time': total_time,
                            'message': '📎 File analyzed by Claude Sonnet (GPT-4 unavailable)',
                            'fallback': True
                        })
                    else:
                        # Both GPT-4 and Claude failed
                        db.execute('UPDATE tasks SET status = ? WHERE id = ?', ('failed', task_id))
                        db.commit()
                        db.close()
                        
                        return jsonify({
                            'success': False,
                            'error': f'File analysis failed: {error_msg}',
                            'task_id': task_id,
                            'conversation_id': conversation_id
                        }), 500
                    
            except Exception as gpt_error:
                print(f"GPT-4 file analysis error: {gpt_error}")
                
                # FALLBACK: Try Claude Sonnet
                try:
                    print(f"🔄 Exception fallback to Claude Sonnet")
                    from orchestration.ai_clients import call_claude_sonnet
                    
                    sonnet_response = call_claude_sonnet(file_analysis_prompt, max_tokens=4000)
                    
                    if isinstance(sonnet_response, dict) and not sonnet_response.get('error'):
                        actual_output = sonnet_response.get('content', '')
                        formatted_output = convert_markdown_to_html(actual_output)
                        
                        total_time = time.time() - overall_start
                        db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                                  ('completed', 'claude_sonnet_file_handler', total_time, task_id))
                        db.commit()
                        db.close()
                        
                        add_message(conversation_id, 'assistant', actual_output, task_id,
                                   {'orchestrator': 'claude_sonnet_file_handler', 'file_analysis': True, 
                                    'execution_time': total_time, 'fallback_from': 'gpt4_exception'})
                        
                        return jsonify({
                            'success': True,
                            'task_id': task_id,
                            'conversation_id': conversation_id,
                            'result': formatted_output,
                            'orchestrator': 'claude_sonnet_file_handler',
                            'execution_time': total_time,
                            'message': '📎 File analyzed by Claude Sonnet',
                            'fallback': True
                        })
                except:
                    pass
                    
                db.close()
                # Don't return here - let it fall through to check conversation history

        # Check for file contents in conversation history OR progressive file session
        if not file_contents and conversation_id:
            # First check if there's a progressive file analysis in session
            file_analysis_state = session.get(f'file_analysis_{conversation_id}')
            
            if file_analysis_state and file_analysis_state.get('file_path'):
                file_path = file_analysis_state['file_path']
                
                # File exists in session - extract a chunk for context
                if os.path.exists(file_path):
                    print(f"📎 Using file from progressive analysis session")
                    
                    try:
                        analyzer = get_progressive_analyzer()
                        # Get first 100 rows for context
                        chunk_result = analyzer.extract_excel_chunk(file_path, start_row=0, num_rows=100)
                        
                        if chunk_result['success']:
                            file_contents = f"{chunk_result['summary']}\n\n{chunk_result['text_preview']}"
                            print(f"✅ Retrieved file context from session ({len(file_contents)} chars)")
                    except Exception as extract_error:
                        print(f"⚠️ Could not extract from session file: {extract_error}")
            
            # If still no file, check conversation history
            if not file_contents:
                from database import get_conversation_file_contents
                file_contents = get_conversation_file_contents(conversation_id)
            if file_contents:
                print(f"📎 Retrieved file contents from conversation history")
                
                db = get_db()
                cursor = db.execute('INSERT INTO tasks (user_request, status, conversation_id) VALUES (?, ?, ?)',
                               (user_request, 'processing', conversation_id))
                task_id = cursor.lastrowid
                db.commit()
                
                from orchestration.ai_clients import call_gpt4
                
                file_analysis_prompt = f"""The user previously uploaded files and is now asking a follow-up question: {user_request}

Here are the file contents from the previous upload:

{file_contents}

Please respond to the user's follow-up question based on these files."""

                try:
                    gpt_response = call_gpt4(file_analysis_prompt, max_tokens=4000)
                    
                    if not gpt_response.get('error') and gpt_response.get('content'):
                        actual_output = gpt_response.get('content', '')
                        formatted_output = convert_markdown_to_html(actual_output)
                        
                        total_time = time.time() - overall_start
                        db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                                  ('completed', 'gpt4_file_handler', total_time, task_id))
                        db.commit()
                        db.close()
                        
                        add_message(conversation_id, 'assistant', actual_output, task_id,
                               {'orchestrator': 'gpt4_file_handler', 'file_analysis': True, 'execution_time': total_time})
                        
                        return jsonify({
                            'success': True,
                            'task_id': task_id,
                            'conversation_id': conversation_id,
                            'result': formatted_output,
                            'orchestrator': 'gpt4_file_handler',
                            'execution_time': total_time,
                            'message': '📎 Follow-up answered using uploaded files'
                        })
                    else:
                        db.close()
                        
                except Exception as gpt_error:
                    print(f"GPT-4 follow-up error: {gpt_error}")
                    db.close()
        
        # Parse clarification answers
        clarification_answers = None
        if request.is_json:
            clarification_answers_raw = request.json.get('clarification_answers')
            if clarification_answers_raw:
                if isinstance(clarification_answers_raw, str):
                    try:
                        clarification_answers = json.loads(clarification_answers_raw)
                    except:
                        clarification_answers = None
                else:
                    clarification_answers = clarification_answers_raw
        else:
            clarification_answers_str = request.form.get('clarification_answers')
            if clarification_answers_str:
                try:
                    clarification_answers = json.loads(clarification_answers_str)
                except:
                    clarification_answers = None
        
        if clarification_answers:
            print(f"✅ Clarification answers received: {clarification_answers}")
            context_additions = []
            for field, value in clarification_answers.items():
                context_additions.append(f"{field}: {value}")
            if context_additions:
                user_request = f"{user_request}\n\nAdditional context:\n" + "\n".join(context_additions)
        
        # Create conversation if needed
        if not conversation_id:
            conversation_id = create_conversation(mode=mode, project_id=project_id)
            print(f"Created new conversation: {conversation_id}")
        
        add_message(conversation_id, 'user', user_request)
        
        # Initialize proactive agent
        proactive = None
        try:
            proactive = ProactiveAgent()
        except Exception as proactive_init_error:
            print(f"Proactive agent init failed: {proactive_init_error}")
        
        # Proactive pre-check
        if proactive and not clarification_answers:
            try:
                pre_check = proactive.pre_process_request(user_request)
                if pre_check['action'] == 'ask_questions':
                    db = get_db()
                    cursor = db.execute('INSERT INTO tasks (user_request, status, conversation_id) VALUES (?, ?, ?)',
                                       (user_request, 'needs_clarification', conversation_id))
                    task_id = cursor.lastrowid
                    db.commit()
                    db.close()
                    return jsonify({'success': True, 'needs_clarification': True, 'clarification_data': pre_check['data'],
                                   'task_id': task_id, 'conversation_id': conversation_id})
                if pre_check['action'] == 'detect_project':
                    return jsonify({'success': True, 'project_detected': True, 'project_data': pre_check['data'],
                                   'conversation_id': conversation_id})
            except Exception as proactive_error:
                print(f"Proactive check failed: {proactive_error}")
        
        # Get knowledge base
        import sys
        app_module = sys.modules.get('app')
        knowledge_base = getattr(app_module, 'knowledge_base', None) if app_module else None
        conversation_context = get_conversation_context(conversation_id, max_messages=10)

        # Get project file context
        file_context = ""
        if project_id:
            try:
                file_context = get_files_for_ai_context(project_id, max_files=5, max_chars_per_file=2000)
                if file_context:
                    print(f"✅ Added file context from project {project_id}")
            except Exception as file_ctx_error:
                print(f"⚠️ Could not load file context: {file_ctx_error}")
        
        db = get_db()
        cursor = db.execute('INSERT INTO tasks (user_request, status, conversation_id) VALUES (?, ?, ?)',
                           (user_request, 'processing', conversation_id))
        task_id = cursor.lastrowid
        db.commit()
        
        # Code assistant check
        try:
            code_assistant_agent = get_code_assistant(knowledge_base=knowledge_base)
            feedback_check = code_assistant_agent.detect_code_feedback(user_request)
            
            if feedback_check['is_code_feedback']:
                print(f"🔧 Code feedback detected for: {feedback_check['target_file']}")
                
                from orchestration.ai_clients import call_claude_sonnet
                
                code_result = code_assistant_agent.process_code_feedback(user_request, call_claude_sonnet)
                
                if code_result['success']:
                    deployment_pkg = code_result['deployment_package']
                    doc_id = None
                    document_url = None
                    
                    if code_result.get('output_path'):
                        try:
                            file_size = os.path.getsize(code_result['output_path'])
                            doc_id = save_generated_document(
                                filename=os.path.basename(code_result['output_path']),
                                original_name=f"Fixed: {code_result['target_file']}",
                                document_type='py',
                                file_path=code_result['output_path'],
                                file_size=file_size,
                                task_id=task_id,
                                conversation_id=conversation_id,
                                project_id=project_id,
                                title=f"Code Fix: {code_result['target_file']}",
                                description=deployment_pkg['change_summary'],
                                category='code'
                            )
                            document_url = f'/api/generated-documents/{doc_id}/download'
                        except Exception as doc_error:
                            print(f"Could not save code file: {doc_error}")
                    
                    response_html = convert_markdown_to_html(code_result['message'])
                    
                    db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                              ('completed', 'code_assistant', time.time() - overall_start, task_id))
                    db.commit()
                    
                    add_message(conversation_id, 'assistant', code_result['message'], task_id,
                               {'document_created': True, 'document_type': 'py', 'document_id': doc_id,
                                'orchestrator': 'code_assistant', 'target_file': code_result['target_file']})
                    
                    suggestions = []
                    if proactive:
                        try:
                            suggestions = proactive.post_process_result(task_id, user_request, code_result['message'])
                        except:
                            pass
                    
                    db.close()
                    
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'conversation_id': conversation_id,
                        'result': response_html,
                        'document_url': document_url,
                        'document_id': doc_id,
                        'document_created': True,
                        'document_type': 'py',
                        'execution_time': time.time() - overall_start,
                        'orchestrator': 'code_assistant',
                        'target_file': code_result['target_file'],
                        'deployment_instructions': deployment_pkg['deployment_instructions'],
                        'suggestions': suggestions
                    })
        
        except Exception as code_assistant_error:
            print(f"Code Assistant failed: {code_assistant_error}")
            
        # Schedule handler check
        schedule_handler = get_combined_schedule_handler()
        schedule_context = get_schedule_context(conversation_id)
        schedule_result = schedule_handler.process_request(user_request, schedule_context)
        
        if schedule_result['action'] != 'not_schedule_request':
            if 'context' in schedule_result:
                save_schedule_context(conversation_id, schedule_result['context'])
            
            if schedule_result['action'] == 'generate_schedule':
                source_filepath = schedule_result['filepath']
                filename = os.path.basename(source_filepath)
                file_path = source_filepath
                file_size = os.path.getsize(file_path)
                shift_length = schedule_result['shift_length']
                pattern_key = schedule_result['pattern_key']
                doc_title = f"{shift_length}-Hour {pattern_key.upper().replace('_', ' ')} Schedule Pattern"
                
                doc_id = save_generated_document(
                    filename=filename,
                    original_name=doc_title,
                    document_type='xlsx',
                    file_path=file_path,
                    file_size=file_size,
                    task_id=task_id,
                    conversation_id=conversation_id,
                    project_id=project_id,
                    title=doc_title,
                    description=f"Visual {shift_length}-hour {pattern_key} schedule pattern",
                    category='schedule'
                )
                
                response_html = convert_markdown_to_html(schedule_result['message'])
                
                db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                          ('completed', 'pattern_schedule_generator', time.time() - overall_start, task_id))
                db.commit()
                
                add_message(conversation_id, 'assistant', schedule_result['message'], task_id,
                           {'document_created': True, 'document_type': 'xlsx', 'document_id': doc_id, 
                            'orchestrator': 'pattern_schedule_generator', 'shift_length': shift_length,
                            'pattern': pattern_key})
                
                session.pop('schedule_context', None)
                
                suggestions = []
                if proactive:
                    try:
                        suggestions = proactive.post_process_result(task_id, user_request, schedule_result['message'])
                    except:
                        pass
                
                db.close()
                
                return jsonify({
                    'success': True,
                    'task_id': task_id,
                    'conversation_id': conversation_id,
                    'result': response_html,
                    'document_url': f'/api/generated-documents/{doc_id}/download',
                    'document_id': doc_id,
                    'document_created': True,
                    'document_type': 'xlsx',
                    'execution_time': time.time() - overall_start,
                    'orchestrator': 'pattern_schedule_generator',
                    'knowledge_applied': False,
                    'formatting_applied': True,
                    'specialists_used': [],
                    'consensus': None,
                    'suggestions': suggestions,
                    'shift_length': shift_length,
                    'pattern': pattern_key
                })
            
            else:
                response_html = convert_markdown_to_html(schedule_result['message'])
                db.execute('UPDATE tasks SET status = ? WHERE id = ?', ('in_progress', task_id))
                db.commit()
                
                add_message(conversation_id, 'assistant', schedule_result['message'], task_id,
                           {'waiting_for_input': True, 'orchestrator': 'pattern_schedule_generator',
                            'waiting_for': schedule_result.get('waiting_for')})
                
                db.close()
                
                return jsonify({
                    'success': True,
                    'task_id': task_id,
                    'conversation_id': conversation_id,
                    'result': response_html,
                    'needs_input': True,
                    'waiting_for': schedule_result.get('waiting_for'),
                    'orchestrator': 'pattern_schedule_generator',
                    'execution_time': time.time() - overall_start
                })

        # Introspection check
        try:
            from introspection import is_introspection_request, get_introspection_engine
            
            introspection_check = is_introspection_request(user_request)
            
            if introspection_check['is_introspection']:
                print(f"🔍 Introspection request detected")
                
                intro_engine = get_introspection_engine()
                action = introspection_check['action']
                
                if action == 'run':
                    days = 7
                    if 'monthly' in user_request.lower() or '30' in user_request:
                        days = 30
                    
                    report = intro_engine.run_introspection(days=days, is_monthly=(days >= 28))
                    response_html = convert_markdown_to_html(report.get('reflection', ''))
                    
                    db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                              ('completed', 'introspection_engine', time.time() - overall_start, task_id))
                    db.commit()
                    
                    add_message(conversation_id, 'assistant', report.get('reflection', ''), task_id,
                               {'orchestrator': 'introspection_engine', 'introspection_id': report.get('insight_id')})
                    
                    db.close()
                    
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'conversation_id': conversation_id,
                        'result': response_html,
                        'orchestrator': 'introspection_engine',
                        'execution_time': time.time() - overall_start,
                        'introspection_report': {
                            'id': report.get('insight_id'),
                            'health_score': report.get('summary', {}).get('health_score', 0),
                            'period_days': report.get('period_days', days),
                            'type': report.get('introspection_type', 'weekly')
                        }
                    })
                
                elif action == 'show_latest':
                    latest = intro_engine.get_latest_introspection()
                    
                    if not latest or not latest.get('full_report'):
                        response_text = "No introspection has been run yet."
                    else:
                        response_text = latest.get('full_report', {}).get('reflection', 'No reflection available.')
                    
                    response_html = convert_markdown_to_html(response_text)
                    
                    db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ? WHERE id = ?',
                              ('completed', 'introspection_engine', task_id))
                    db.commit()
                    
                    add_message(conversation_id, 'assistant', response_text, task_id,
                               {'orchestrator': 'introspection_engine'})
                    
                    db.close()
                    
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'conversation_id': conversation_id,
                        'result': response_html,
                        'orchestrator': 'introspection_engine',
                        'execution_time': time.time() - overall_start
                    })
        
        except ImportError:
            print("ℹ️  Introspection not available")
        except Exception as intro_error:
            print(f"⚠️  Introspection detection failed: {intro_error}")
        
        # Regular AI orchestration
        try:
            print(f"Analyzing task: {user_request[:100]}...")
            analysis = analyze_task_with_sonnet(user_request, knowledge_base=knowledge_base, file_paths=file_paths, file_contents=file_contents)
            task_type = analysis.get('task_type', 'general')
            confidence = analysis.get('confidence', 0.5)
            escalate = analysis.get('escalate_to_opus', False)
            specialists_needed = analysis.get('specialists_needed', [])
            knowledge_applied = analysis.get('knowledge_applied', False)
            knowledge_sources = analysis.get('knowledge_sources', [])
            
            if specialists_needed:
                specialists_needed = [s for s in specialists_needed if s and s.lower() != 'none']
            
            orchestrator = 'sonnet'
            opus_guidance = None
            
            if escalate:
                print("Escalating to Opus...")
                orchestrator = 'opus'
                try:
                    opus_result = handle_with_opus(user_request, analysis, knowledge_base=knowledge_base, file_paths=file_paths, file_contents=file_contents)
                    opus_guidance = opus_result.get('strategic_analysis', '')
                    if opus_result.get('specialist_assignments'):
                        for assignment in opus_result.get('specialist_assignments', []):
                            specialist = assignment.get('ai') or assignment.get('specialist')
                            if specialist and specialist.lower() != 'none':
                                specialists_needed.append(specialist)
                except Exception as opus_error:
                    print(f"Opus guidance failed: {opus_error}")
            
            specialist_results = []
            specialist_output = None
            if specialists_needed:
                for specialist_info in specialists_needed:
                    if isinstance(specialist_info, dict):
                        specialist = specialist_info.get('specialist') or specialist_info.get('ai')
                        specialist_task = specialist_info.get('task', user_request)
                    else:
                        specialist = specialist_info
                        specialist_task = user_request
                    if specialist and specialist.lower() != 'none':
                        result = execute_specialist_task(specialist, specialist_task, file_paths=file_paths, file_contents=file_contents)
                        specialist_results.append(result)
                        if result.get('success') and result.get('output'):
                            specialist_output = result.get('output')
      
            from orchestration.ai_clients import call_claude_opus, call_claude_sonnet
            knowledge_context = get_knowledge_context_for_prompt(knowledge_base, user_request)

            # ================================================================
            # LEARNING LOOP ENHANCEMENT - Get Past Learnings
            # February 4, 2026: Retrieve learned patterns from database
            # ================================================================
            learning_context = ""
            try:
                learning_context = get_learning_context()
                if learning_context:
                    print(f"🧠 Retrieved learning context ({len(learning_context)} chars)")
            except Exception as learn_ctx_error:
                print(f"⚠️ Could not get learning context (non-critical): {learn_ctx_error}")
            
            # ================================================================
            # FIX #2: CLIENT PROFILE CONTEXT
            # February 4, 2026: Inject accumulated client knowledge
            # ================================================================
            client_profile_context = ""
            if project_id:
                try:
                    # Get project to find client name
                    db_temp = get_db()
                    project = db_temp.execute('SELECT client_name FROM projects WHERE project_id = ?', (project_id,)).fetchone()
                    db_temp.close()
                    
                    if project and project['client_name']:
                        client_profile_context = get_client_profile_context(project['client_name'])
                        if client_profile_context:
                            print(f"👤 Retrieved client profile for {project['client_name']}")
                except Exception as profile_error:
                    print(f"⚠️ Could not get client profile (non-critical): {profile_error}")
            
            # ================================================================
            # FIX #3: AVOIDANCE PATTERNS CONTEXT
            # February 4, 2026: Inject patterns to avoid
            # ================================================================
            avoidance_context = ""
            try:
                avoidance_context = get_avoidance_context(days=30, limit=5)
                if avoidance_context:
                    print(f"🚫 Retrieved {avoidance_context.count('⚠️') + avoidance_context.count('🚫')} avoidance patterns")
            except Exception as avoid_error:
                print(f"⚠️ Could not get avoidance context (non-critical): {avoid_error}")

            # ================================================================
            # FIX #6: SPECIALIZED KNOWLEDGE INJECTION
            # February 4, 2026: Industry expertise and normative benchmarks
            # ================================================================
            specialized_context = ""
            try:
                specialist = get_specialized_knowledge()
                
                # Get industry from project if available
                industry = None
                if project_id:
                    db_temp = get_db()
                    proj = db_temp.execute('SELECT industry FROM projects WHERE project_id = ?', (project_id,)).fetchone()
                    db_temp.close()
                    if proj:
                        industry = proj['industry']
                
                specialized_context = specialist.build_expertise_context(user_request, industry)
                if specialized_context:
                    print(f"🎓 Injected specialized knowledge for {industry or 'general'}")
            except Exception as spec_error:
                print(f"⚠️ Could not get specialized knowledge: {spec_error}")

            # ================================================================
            # FIX #5: MULTI-TURN CONTEXT AWARENESS
            # February 4, 2026: Conversation summarization
            # ================================================================
            summary_context = ""
            try:
                summarizer = get_conversation_summarizer()
                
                # Check if we should summarize
                if summarizer.should_summarize(conversation_id):
                    from orchestration.ai_clients import call_claude_sonnet
                    summarizer.summarize_conversation(conversation_id, call_claude_sonnet)
                
                # Get existing summaries
                summary_context = summarizer.get_conversation_context(conversation_id)
                if summary_context:
                    print(f"📝 Retrieved conversation summary")
            except Exception as summary_error:
                print(f"⚠️ Could not get conversation summary: {summary_error}")

            # ================================================================
            # LEARNING LOOP ENHANCEMENT - Initialize Intelligence Engine
            # February 4, 2026
            # ================================================================
            intelligence = None
            try:
                intelligence = EnhancedIntelligence()
                print("🧠 EnhancedIntelligence initialized")
            except Exception as intel_error:
                print(f"⚠️ EnhancedIntelligence init failed (non-critical): {intel_error}")
                intelligence = None
              
            # Get smart defaults from user history
            smart_defaults = {}
            if intelligence:
                try:
                    smart_defaults = intelligence.get_smart_defaults('general')
                    if smart_defaults and any(v for v in smart_defaults.values() if v):
                        print(f"🎯 Retrieved smart defaults: {list(smart_defaults.keys())}")
                except Exception as defaults_error:
                    print(f"⚠️ Could not get smart defaults (non-critical): {defaults_error}")

            # Build project context
            project_context = ""
            if project_id:
                try:
                    db_temp = get_db()
                    project = db_temp.execute('SELECT * FROM projects WHERE project_id = ?', (project_id,)).fetchone()
                    db_temp.close()
                    
                    if project:
                        file_stats = get_file_stats_by_project(project_id)
                        
                        project_context = f"""

=== CURRENT PROJECT CONTEXT ===
You are working inside the "{project['client_name']}" PROJECT FOLDER.
- Industry: {project['industry']}
- Facility Type: {project['facility_type']}
- Project Phase: {project['project_phase']}

This project folder contains: {file_stats.get('total_files', 0)} files
- Uploaded: {file_stats.get('uploaded_files', 0)}
- Generated: {file_stats.get('generated_files', 0)}
===

"""
                except Exception as proj_ctx_error:
                    print(f"⚠️ Could not load project context: {proj_ctx_error}")
            
            # Build conversation history
            conversation_history = ""
            if conversation_context and len(conversation_context) > 1:
                conversation_history = "\n\n=== CONVERSATION HISTORY ===\n"
                for msg in conversation_context[:-1]:
                    role_label = "User" if msg['role'] == 'user' else "Assistant"
                    content_preview = msg['content'][:500] + '...' if len(msg['content']) > 500 else msg['content']
                    conversation_history += f"{role_label}: {content_preview}\n"
                conversation_history += "=== END CONVERSATION HISTORY ===\n\n"
            
            if specialist_output:
                actual_output = specialist_output
            else:
                # 🔧 CRITICAL FIX: Add file contents to completion prompt
                if file_contents:
                    file_section = f"""

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📎 ATTACHED FILES - READ THESE CAREFULLY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{file_contents}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
                else:
                    file_section = ""

                completion_prompt = f"""{knowledge_context}{project_context}{file_context}{conversation_history}{learning_context}{client_profile_context}{avoidance_context}{specialized_context}{summary_context}{file_section}

USER REQUEST: {user_request}

Please complete this request fully. Provide the actual deliverable.
Be comprehensive and professional."""

                if opus_guidance:
                    completion_prompt += f"\n\nSTRATEGIC GUIDANCE:\n{opus_guidance}"

                if file_contents:
                    print(f"🔍 Completion prompt contains {len(file_contents)} chars of file content")
                              
                if orchestrator == 'opus':
                    response = call_claude_opus(completion_prompt, conversation_history=conversation_context, files_attached=bool(file_contents))
                else:
                    response = call_claude_sonnet(completion_prompt, conversation_history=conversation_context, files_attached=bool(file_contents))
                
                if isinstance(response, dict):
                    if response.get('error'):
                        actual_output = f"Error: {response.get('content', 'Unknown error')}"
                    else:
                        actual_output = response.get('content', '')
                else:
                    actual_output = str(response)
            
            print(f"Task completed. Output length: {len(actual_output) if actual_output else 0} chars")
            
            # Document creation
            document_created = False
            document_url = None
            document_type = None
            document_id = None
            
            should_create, doc_type = should_create_document(user_request)
            if should_create and actual_output and not actual_output.startswith('Error'):
                try:
                    from docx import Document
                    from docx.shared import Pt
                    doc = Document()
                    title = doc.add_heading('Shiftwork Solutions LLC', 0)
                    title.alignment = 1
                    lines = actual_output.split('\n')
                    for line in lines:
                        if line.strip():
                            if line.startswith('#'):
                                level = len(line) - len(line.lstrip('#'))
                                text = line.lstrip('#').strip()
                                doc.add_heading(text, level=min(level, 3))
                            else:
                                p = doc.add_paragraph(line)
                                p.style.font.size = Pt(11)
                    
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = f'shiftwork_document_{timestamp}.docx'
                    save_paths = [f'/tmp/{filename}', f'./{filename}']
                    
                    saved = False
                    saved_path = None
                    for filepath in save_paths:
                        try:
                            os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else '.', exist_ok=True)
                            doc.save(filepath)
                            saved = True
                            saved_path = filepath
                            print(f"Document created: {filepath}")
                            break
                        except Exception as save_error:
                            print(f"Could not save to {filepath}: {save_error}")
                    
                    if saved and saved_path:
                        file_size = os.path.getsize(saved_path)
                        doc_title = generate_document_title(user_request, 'docx')
                        doc_category = categorize_document(user_request, 'docx')
                        document_id = save_generated_document(
                            filename=filename, original_name=doc_title, document_type='docx',
                            file_path=saved_path, file_size=file_size, task_id=task_id,
                            conversation_id=conversation_id, project_id=project_id,
                            title=doc_title, description=f"Generated from: {user_request[:100]}",
                            category=doc_category
                        )
                        document_created = True
                        document_url = f'/api/generated-documents/{document_id}/download'
                        document_type = 'docx'
                except Exception as doc_error:
                    print(f"Document creation failed: {doc_error}")
            
            formatted_output = convert_markdown_to_html(actual_output)
            
            consensus_result = None
            if enable_consensus and actual_output and not actual_output.startswith('Error'):
                try:
                    consensus_result = validate_with_consensus(actual_output)
                except Exception as consensus_error:
                    print(f"Consensus validation failed: {consensus_error}")
            
            total_time = time.time() - overall_start
            db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                      ('completed', orchestrator, total_time, task_id))
            db.commit()
            db.close()
            
            add_message(conversation_id, 'assistant', actual_output, task_id,
                       {'orchestrator': orchestrator, 'knowledge_applied': knowledge_applied,
                        'document_created': document_created, 'document_type': document_type,
                        'document_id': document_id, 'execution_time': total_time})
            
            suggestions = []
            if proactive:
                try:
                    suggestions = proactive.post_process_result(task_id, user_request, actual_output if actual_output else '')
                except Exception as suggest_error:
                    print(f"Suggestion generation failed: {suggest_error}")

            # ================================================================
            # LEARNING LOOP ENHANCEMENT - Store What We Learned
            # February 4, 2026: Complete the loop by storing new patterns
            # ================================================================
            if intelligence and actual_output and not actual_output.startswith('Error'):
                try:
                    intelligence.learn_from_interaction(user_request, actual_output, user_feedback=None)
                    print("✅ EnhancedIntelligence learned from this interaction")
                except Exception as learn_error:
                    print(f"⚠️ EnhancedIntelligence learning failed (non-critical): {learn_error}")   

            # ================================================================
            # FIX #2: UPDATE CLIENT PROFILE
            # February 4, 2026: Build cumulative client knowledge
            # ================================================================
            if project_id:
                try:
                    # Get project to find client name
                    db_temp = get_db()
                    project = db_temp.execute('SELECT client_name, industry FROM projects WHERE project_id = ?', (project_id,)).fetchone()
                    db_temp.close()
                    
                    if project and project['client_name']:
                        interaction_data = {
                            'approach': orchestrator,
                            'approach_worked': True,  # Assume success unless feedback says otherwise
                            'industry': project['industry'],
                            'preferences': {}
                        }
                        
                        # Detect preferences from request
                        if 'dupont' in user_request.lower():
                            interaction_data['preferences']['schedule_type'] = 'DuPont'
                        elif '12 hour' in user_request.lower() or '12-hour' in user_request.lower():
                            interaction_data['preferences']['shift_length'] = 12
                        
                        update_client_profile(project['client_name'], interaction_data)
                        print(f"👤 Updated profile for {project['client_name']}")
                except Exception as profile_update_error:
                    print(f"⚠️ Client profile update failed (non-critical): {profile_update_error}")

            # ================================================================
            # FIX #4: PROACTIVE SUGGESTIONS
            # February 4, 2026: Generate next-step suggestions
            # ================================================================
            suggestions_generated = []
            try:
                suggester = get_proactive_suggestions()
                
                suggestion_context = {
                    'task_id': task_id,
                    'user_request': user_request,
                    'ai_response': actual_output,
                    'document_created': document_created,
                    'orchestrator': orchestrator
                }
                
                suggestions_generated = suggester.generate_suggestions(suggestion_context)
                
                # Store suggestions
                for suggestion in suggestions_generated:
                    suggester.store_suggestion(conversation_id, task_id, suggestion)
                
                if suggestions_generated:
                    print(f"💡 Generated {len(suggestions_generated)} proactive suggestions")
            except Exception as suggest_error:
                print(f"⚠️ Could not generate suggestions: {suggest_error}")   
                        
            # Auto-learn from this conversation
            try:
                learn_from_conversation(user_request, actual_output if actual_output else '')
            except Exception as learn_error:
                print(f"⚠️ Auto-learning failed (non-critical): {learn_error}")
            
            # ================================================================
            # PHASE 1: PROACTIVE CURIOSITY ENGINE - February 5, 2026
            # Ask curious follow-up questions after completing tasks
            # ================================================================
            curious_question = None
            try:
                curiosity_engine = get_curiosity_engine()
                curiosity_check = curiosity_engine.should_be_curious(
                    conversation_id,
                    {
                        'user_request': user_request,
                        'ai_response': actual_output if actual_output else '',
                        'task_completed': True
                    }
                )
                
                if curiosity_check['should_ask']:
                    curious_question = curiosity_check['question']
                    print(f"🤔 Curious follow-up: {curious_question}")
            except Exception as curiosity_error:
                print(f"⚠️ Curiosity engine failed (non-critical): {curiosity_error}")
            # ================================================================
            
            return jsonify({
                'success': True, 'task_id': task_id, 'conversation_id': conversation_id,
                'result': formatted_output, 'orchestrator': orchestrator,
                'specialists_used': [s.get('specialist') for s in specialist_results] if specialist_results else [],
                'consensus': consensus_result, 'execution_time': total_time,
                'knowledge_applied': knowledge_applied, 'knowledge_used': knowledge_applied,
                'knowledge_sources': knowledge_sources, 'formatting_applied': True,
                'document_created': document_created, 'document_url': document_url,
                'document_id': document_id, 'document_type': document_type, 'suggestions': suggestions,
                'proactive_suggestions': suggestions_generated,
                'curious_question': curious_question
            })
        except Exception as orchestration_error:
            import traceback
            print(f"Orchestration error: {traceback.format_exc()}")
            db.execute('UPDATE tasks SET status = ? WHERE id = ?', ('failed', task_id))
            db.commit()
            db.close()
            add_message(conversation_id, 'assistant', f"Error: {str(orchestration_error)}", task_id, {'error': True})
            return jsonify({'success': False, 'error': f'Orchestration failed: {str(orchestration_error)}',
                           'task_id': task_id, 'conversation_id': conversation_id}), 500
    except Exception as e:
        import traceback
        print(f"CRITICAL ERROR: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'}), 500

def handle_large_excel_initial(file_path, user_request, conversation_id, project_id, mode, file_info, overall_start):
    """
    Handle initial upload of a large Excel file - analyze first 100 rows.
    
    UPDATED February 6, 2026 (v8): MEMORY FIX - Store file in database, not session/memory
    - Saves file to permanent location in database
    - Stores file_id in session instead of file contents
    - Re-extracts chunks on demand instead of keeping in RAM
    - Prevents 512MB+ memory crashes on paid Render plans
    
    UPDATED February 5, 2026 (v7): Reduced max_tokens from 6000 to 3000 to prevent timeout
    """
    try:
        analyzer = get_progressive_analyzer()
        
        chunk_result = analyzer.extract_excel_chunk(file_path, start_row=0, num_rows=100)
        
        if not chunk_result['success']:
            return jsonify({
                'success': False,
                'error': f"Could not analyze Excel file: {chunk_result.get('error')}"
            }), 500
        
        # Create conversation if needed
        if not conversation_id:
            conversation_id = create_conversation(mode=mode, project_id=project_id)
            print(f"Created new conversation: {conversation_id}")
        
        # ================================================================
        # FIX v8: Copy file to PERMANENT location to prevent memory issues
        # Keep file available for follow-up questions without holding in RAM
        # ================================================================
        import shutil
        
        # Create permanent storage directory
        permanent_dir = '/mnt/project/uploaded_files'
        os.makedirs(permanent_dir, exist_ok=True)
        
        # Copy to permanent location with unique name
        original_filename = os.path.basename(file_path)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        permanent_filename = f"{timestamp}_{original_filename}"
        permanent_path = os.path.join(permanent_dir, permanent_filename)
        
        shutil.copy2(file_path, permanent_path)
        print(f"💾 Saved file permanently: {permanent_path}")
        
        
        # Store minimal info in session (just IDs, not data)
        session[f'file_analysis_{conversation_id}'] = {
            'file_path': permanent_path,
            'current_position': chunk_result['end_row'],
            'total_rows': chunk_result['total_rows'],
            'columns': chunk_result['columns'],
            'file_name': original_filename,
            'file_size_mb': file_info['file_size_mb']
        }
        # ================================================================
        
        # Create task
        db = get_db()
        cursor = db.execute('INSERT INTO tasks (user_request, status, conversation_id) VALUES (?, ?, ?)',
                           (user_request, 'processing', conversation_id))
        task_id = cursor.lastrowid
        db.commit()
        
        # Build analysis prompt
        from orchestration.ai_clients import call_gpt4
        
        sheet_names = chunk_result.get('sheet_names', ['Sheet1'])
        num_sheets = chunk_result.get('num_sheets', 1)
        sheets_summary = f"\n📋 **FILE CONTAINS {num_sheets} WORKSHEET(S):** {', '.join(sheet_names)}\n"
        
        analysis_prompt = f"""You are Jim Goodwin, Shiftwork Solutions LLC - 30+ years optimizing 24/7 operations.

**CLIENT FILE:** {file_info['file_size_mb']}MB Excel, {chunk_result['total_rows']:,} total rows
**REQUEST:** {user_request}
{sheets_summary}

**YOUR TASK:** Analyze first 100 rows with SPECIFIC NUMBERS - this is a $16,500/week consulting engagement.

**DELIVER:**
1. **Key Statistics** - Actual totals, ranges, averages (not "varies")
2. **Top Patterns** - What 3 things stand out? Use percentages.
3. **Operational Insights** - Coverage gaps? Shift patterns? Cost drivers?
4. **Red Flags** - What needs immediate attention?
5. **Next Steps** - What would you recommend?

**DATA (First 100 rows):**
{chunk_result['summary']}

{chunk_result['text_preview']}

BE SPECIFIC. Use actual numbers from the data."""

        print(f"📊 Calling GPT-4 with max_tokens=3000...")
        gpt_response = call_gpt4(analysis_prompt, max_tokens=3000)
        
        if not gpt_response.get('error') and gpt_response.get('content'):
            ai_analysis = gpt_response.get('content', '')
            
            # Add continuation prompt
            continuation_prompt = analyzer.generate_continuation_prompt(chunk_result)
            full_response = ai_analysis + continuation_prompt
            
            formatted_output = convert_markdown_to_html(full_response)
                      
            total_time = time.time() - overall_start
            db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                      ('completed', 'gpt4_progressive_excel', total_time, task_id))
            db.commit()
            db.close()
            
            add_message(conversation_id, 'assistant', full_response, task_id,
                       {'orchestrator': 'gpt4_progressive_excel', 'rows_analyzed': 100, 
                        'total_rows': chunk_result['total_rows'], 'execution_time': total_time,
                        'permanent_file': permanent_path})

                        
            return jsonify({
                'success': True,
                'task_id': task_id,
                'conversation_id': conversation_id,
                'result': formatted_output,
                'orchestrator': 'gpt4_progressive_excel',
                'execution_time': total_time,
                'progressive_analysis': True,
                'rows_analyzed': 100,
                'total_rows': chunk_result['total_rows'],
                'rows_remaining': chunk_result['rows_remaining']
            })
        else:
            error_msg = gpt_response.get('content', 'Unknown error')
            print(f"❌ GPT-4 analysis failed: {error_msg}")
            db.close()
            return jsonify({
                'success': False,
                'error': f'Could not analyze Excel file: {error_msg}'
            }), 500
            
    except Exception as e:
        import traceback
        print(f"❌ Large Excel handling error: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500

def handle_progressive_continuation(conversation_id, user_request, continuation_request, file_analysis_state, overall_start):
    """
    Handle user requesting more rows from a large Excel file.
    UPDATED February 6, 2026: Fixed to provide actual analysis instead of hypothetical
    Added: January 31, 2026
    """
    try:
        analyzer = get_progressive_analyzer()
        
        file_path = file_analysis_state['file_path']
        current_position = file_analysis_state['current_position']
        total_rows = file_analysis_state['total_rows']
        file_name = file_analysis_state.get('file_name', 'file')
        
        # Determine how many rows to analyze
        action = continuation_request['action']
        
        if action == 'analyze_next':
            num_rows = continuation_request['num_rows']
            start_row = current_position
        elif action == 'analyze_all':
            num_rows = None  # All remaining
            start_row = current_position
        elif action == 'analyze_range':
            start_row = continuation_request['start_row']
            num_rows = continuation_request['end_row'] - start_row
        else:
            return jsonify({'success': False, 'error': 'Invalid continuation action'}), 400
        
        end_display = start_row + num_rows if num_rows else 'end'
        print(f"📊 Extracting rows {start_row} to {end_display} from {file_name}")
        
        # Extract the requested chunk
        chunk_result = analyzer.extract_excel_chunk(file_path, start_row=start_row, num_rows=num_rows)
        
        if not chunk_result['success']:
            return jsonify({
                'success': False,
                'error': f"Could not extract data: {chunk_result.get('error')}"
            }), 500
        
        # Update session state
        session[f'file_analysis_{conversation_id}']['current_position'] = chunk_result['end_row']
        
        # Create task
        db = get_db()
        cursor = db.execute('INSERT INTO tasks (user_request, status, conversation_id) VALUES (?, ?, ?)',
                           (user_request, 'processing', conversation_id))
        task_id = cursor.lastrowid
        db.commit()
        
        # Build analysis prompt - CRITICAL: Demand actual calculations
        from orchestration.ai_clients import call_gpt4
        
        rows_analyzed = chunk_result['rows_analyzed']
        
        analysis_prompt = f"""You are Jim Goodwin, Shiftwork Solutions LLC - analyzing workforce data.

**USER REQUEST:** {user_request}

**CRITICAL INSTRUCTION:** The user wants ACTUAL CALCULATIONS from the data below, NOT a description of how to do it.
- Calculate actual totals, averages, sums
- Provide real numbers in tables
- DO NOT say "I would calculate" or "hypothetically" or "for example"
- DO NOT provide code examples or theoretical approaches
- Just analyze the data and give the user the actual numbers

**DATA ANALYZED:** Rows {start_row + 1} to {chunk_result['end_row']} ({rows_analyzed:,} rows)

{chunk_result['summary']}

{chunk_result['text_preview']}

**DELIVER ACTUAL ANALYSIS WITH REAL NUMBERS FROM THIS DATA.**"""

        print(f"📊 Calling GPT-4 to analyze {rows_analyzed:,} rows...")
        gpt_response = call_gpt4(analysis_prompt, max_tokens=4000)
        
        if not gpt_response.get('error') and gpt_response.get('content'):
            ai_analysis = gpt_response.get('content', '')
            
            # Add continuation prompt if more rows remain
            if chunk_result['rows_remaining'] > 0:
                continuation_prompt = analyzer.generate_continuation_prompt(chunk_result)
                full_response = ai_analysis + continuation_prompt
            else:
                full_response = ai_analysis + "\n\n✅ **Analysis complete!** All rows have been analyzed."
                # Clear session state
                session.pop(f'file_analysis_{conversation_id}', None)
            
            formatted_output = convert_markdown_to_html(full_response)
            
            total_time = time.time() - overall_start
            db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                      ('completed', 'gpt4_progressive_excel', total_time, task_id))
            db.commit()
            db.close()
            
            add_message(conversation_id, 'assistant', full_response, task_id,
                       {'orchestrator': 'gpt4_progressive_excel', 'rows_analyzed': rows_analyzed,
                        'total_rows': total_rows, 'execution_time': total_time})
            
            return jsonify({
                'success': True,
                'task_id': task_id,
                'conversation_id': conversation_id,
                'result': formatted_output,
                'orchestrator': 'gpt4_progressive_excel',
                'execution_time': total_time,
                'progressive_analysis': True,
                'rows_analyzed': rows_analyzed,
                'total_rows': total_rows,
                'rows_remaining': chunk_result['rows_remaining']
            })
        else:
            db.close()
            return jsonify({
                'success': False,
                'error': 'Could not analyze data'
            }), 500
            
    except Exception as e:
        import traceback
        print(f"Progressive continuation error: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500




# I did no harm and this file is not truncated
