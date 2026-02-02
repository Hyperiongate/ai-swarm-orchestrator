"""
Document Creation Helper - Python-based Professional DOCX Generation
Created: February 2, 2026
Last Updated: February 2, 2026

REPLACES Node.js/docx-js approach with Python python-docx for reliability on Render.
Uses python-docx which is already in requirements.txt - no npm packages needed!

Author: Jim @ Shiftwork Solutions LLC
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


def create_analysis_document(analysis_text, file_names_str, output_path='/tmp/file_analysis.docx'):
    """
    Create a professional file analysis document using python-docx.
    
    Args:
        analysis_text: The AI analysis text to include in the document
        file_names_str: Comma-separated string of analyzed file names
        output_path: Where to save the document (default: /tmp/file_analysis.docx)
    
    Returns:
        dict with 'success', 'filepath', and optional 'error'
    """
    try:
        # Create document
        doc = Document()
        
        # Company Header - SHIFTWORK SOLUTIONS LLC
        header = doc.add_heading('SHIFTWORK SOLUTIONS LLC', level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.runs[0]
        header_run.font.size = Pt(20)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Subtitle
        subtitle = doc.add_heading('File Analysis Report', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(0, 102, 204)  # Medium blue
        
        # Add spacing
        doc.add_paragraph()
        
        # File(s) Analyzed
        p_files = doc.add_paragraph()
        p_files.add_run('File(s) Analyzed: ').bold = True
        p_files.add_run(file_names_str)
        
        # Date
        p_date = doc.add_paragraph()
        p_date.add_run('Date: ').bold = True
        p_date.add_run(datetime.now().strftime('%B %d, %Y'))
        
        # Analyzed by
        p_analyzer = doc.add_paragraph()
        p_analyzer.add_run('Analyzed by: ').bold = True
        p_analyzer.add_run('AI Swarm Orchestrator')
        
        # Add spacing
        doc.add_paragraph()
        
        # Analysis Results heading
        results_heading = doc.add_heading('Analysis Results', level=2)
        results_run = results_heading.runs[0]
        results_run.font.color.rgb = RGBColor(0, 102, 204)
        
        # Add the analysis text
        # Split into paragraphs and preserve formatting
        for line in analysis_text.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                p_run = p.runs[0]
                p_run.font.size = Pt(11)
                p_run.font.name = 'Arial'
        
        # Save document
        doc.save(output_path)
        
        print(f"✅ Document created successfully: {output_path}")
        
        return {
            'success': True,
            'filepath': output_path,
            'message': 'Document created successfully'
        }
        
    except Exception as e:
        print(f"❌ Document creation failed: {e}")
        import traceback
        traceback.print_exc()
        
        return {
            'success': False,
            'error': str(e),
            'message': f'Failed to create document: {e}'
        }


# I did no harm and this file is not truncated
