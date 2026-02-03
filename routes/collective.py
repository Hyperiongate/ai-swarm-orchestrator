"""
COLLECTIVE INTELLIGENCE API ROUTES - Phase 4
Created: February 2, 2026
Last Updated: February 2, 2026

API endpoints for collective intelligence features.
Provides intelligent questionnaires and deliverable generation.

Endpoints:
- POST /api/collective/learn - Learn from existing materials
- GET  /api/collective/status - Get learning status
- POST /api/collective/questionnaire - Get questionnaire for deliverable
- POST /api/collective/generate - Generate deliverable from answers
- GET  /api/collective/patterns - View learned patterns

Author: Jim @ Shiftwork Solutions LLC (managed by Claude Sonnet 4)
"""

from flask import Blueprint, jsonify, request, current_app
from collective_intelligence_engine import get_collective_intelligence
import sqlite3
from datetime import datetime


# Create blueprint
collective_bp = Blueprint('collective', __name__)


@collective_bp.route('/api/collective/learn', methods=['POST'])
def learn_from_materials():
    """
    Trigger learning from existing project materials.
    
    This analyzes all files in the knowledge base and extracts patterns.
    
    NOTE: With auto-learning enabled, this happens automatically.
    You can still trigger it manually anytime.
    """
    try:
        # Get knowledge base from app
        knowledge_base = getattr(current_app, 'knowledge_base', None)
        
        # If knowledge base not initialized, try to create a simple one
        if not knowledge_base:
            # Create minimal knowledge base structure for learning
            class SimpleKnowledgeBase:
                def __init__(self):
                    self.knowledge_index = []
                    self._load_project_files()
                
                def _load_project_files(self):
                    """Load files from project_files/ directory"""
                    import os
                    import json
                    
                    # Try multiple possible locations
                    possible_dirs = [
                        './project_files/',
                        '/app/project_files/',
                        'project_files/',
                        '/mnt/project/'
                    ]
                    
                    project_dir = None
                    for dir_path in possible_dirs:
                        if os.path.exists(dir_path):
                            project_dir = dir_path
                            print(f"📁 Found project files at: {dir_path}")
                            break
                    
                    if not project_dir:
                        print("⚠️  No project_files directory found")
                        return
                    
                    for filename in os.listdir(project_dir):
                        filepath = os.path.join(project_dir, filename)
                        if os.path.isfile(filepath):
                            try:
                                # Try to read file
                                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                                    content = f.read()
                                
                                self.knowledge_index.append({
                                    'title': filename,
                                    'content': content,
                                    'filepath': filepath
                                })
                                print(f"  ✅ Loaded: {filename}")
                            except Exception as e:
                                print(f"  ⚠️  Could not read {filename}: {e}")
            
            knowledge_base = SimpleKnowledgeBase()
            print(f"📚 Created temporary knowledge base with {len(knowledge_base.knowledge_index)} files")
        
        ci = get_collective_intelligence(knowledge_base=knowledge_base)
        results = ci.learn_from_existing_materials()
        
        return jsonify({
            'success': True,
            'results': results,
            'message': f"Learned from existing materials: {results['patterns_discovered']} patterns discovered",
            'note': 'Auto-learning can run this automatically - see /api/collective/auto-learning/status'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@collective_bp.route('/api/collective/auto-learning/status', methods=['GET'])
def get_auto_learning_status():
    """
    Get status of automatic learning system.
    
    Shows when learning last ran, how it's configured, etc.
    """
    try:
        from auto_learning_trigger import get_auto_learning_trigger
        
        trigger = get_auto_learning_trigger()
        status = trigger.get_auto_learning_status()
        
        return jsonify({
            'success': True,
            'auto_learning': status
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@collective_bp.route('/api/collective/auto-learning/configure', methods=['POST'])
def configure_auto_learning():
    """
    Configure automatic learning.
    
    Body:
    {
        "enabled": true,
        "frequency_hours": 168  // How often to re-learn (default: weekly)
    }
    """
    try:
        from auto_learning_trigger import get_auto_learning_trigger
        
        data = request.get_json() or {}
        
        trigger = get_auto_learning_trigger()
        
        # Enable/disable
        if 'enabled' in data:
            if data['enabled']:
                trigger.enable_auto_learning()
            else:
                trigger.disable_auto_learning()
        
        # Set frequency
        if 'frequency_hours' in data:
            trigger.set_learning_frequency(data['frequency_hours'])
        
        # Get updated status
        status = trigger.get_auto_learning_status()
        
        return jsonify({
            'success': True,
            'auto_learning': status,
            'message': 'Auto-learning configuration updated'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@collective_bp.route('/api/collective/auto-learning/trigger', methods=['POST'])
def trigger_auto_learning():
    """
    Manually trigger auto-learning check.
    
    This checks if learning should run and runs it if conditions are met.
    Useful for testing or forcing a learning cycle.
    """
    try:
        from auto_learning_trigger import get_auto_learning_trigger
        
        knowledge_base = getattr(current_app, 'knowledge_base', None)
        
        trigger = get_auto_learning_trigger()
        result = trigger.check_and_trigger(knowledge_base)
        
        if result:
            return jsonify({
                'success': True,
                'learning_triggered': True,
                'result': result
            })
        else:
            return jsonify({
                'success': True,
                'learning_triggered': False,
                'message': 'No trigger conditions met'
            })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@collective_bp.route('/api/collective/status', methods=['GET'])
def get_status():
    """
    Get status of collective intelligence system.
    
    Returns what has been learned and what's available.
    """
    try:
        ci = get_collective_intelligence()
        db = sqlite3.connect(ci.db_path)
        db.row_factory = sqlite3.Row
        cursor = db.cursor()
        
        # Count patterns
        cursor.execute('SELECT COUNT(*) as count FROM collective_patterns')
        patterns = cursor.fetchone()['count']
        
        # Count learned questions
        cursor.execute('SELECT COUNT(*) as count FROM learned_questions')
        questions = cursor.fetchone()['count']
        
        # Count templates
        cursor.execute('SELECT COUNT(*) as count FROM deliverable_templates')
        templates = cursor.fetchone()['count']
        
        # Get pattern categories
        cursor.execute('''
            SELECT pattern_category, COUNT(*) as count
            FROM collective_patterns
            GROUP BY pattern_category
        ''')
        pattern_breakdown = [dict(row) for row in cursor.fetchall()]
        
        db.close()
        
        status = 'learned' if patterns > 0 else 'needs_learning'
        
        return jsonify({
            'success': True,
            'status': status,
            'metrics': {
                'patterns_discovered': patterns,
                'questions_learned': questions,
                'templates_available': templates
            },
            'pattern_breakdown': pattern_breakdown,
            'capabilities': {
                'document_generation': [
                    {
                        'type': 'implementation_manual',
                        'name': 'Implementation Manual',
                        'description': 'Complete implementation manual for schedule changes',
                        'action': 'generate'
                    }
                ],
                'document_analysis': [
                    {
                        'type': 'oaf',
                        'name': 'Operations Assessment Feedback (OAF)',
                        'description': 'Analyze operational assessments with learned intelligence',
                        'action': 'analyze'
                    },
                    {
                        'type': 'eaf',
                        'name': 'Employee Assessment Feedback (EAF)',
                        'description': 'Analyze employee surveys with normative comparisons',
                        'action': 'analyze'
                    }
                ]
            }
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@collective_bp.route('/api/collective/questionnaire', methods=['POST'])
def get_questionnaire():
    """
    Get intelligent questionnaire for a deliverable type.
    
    Body:
    {
        "deliverable_type": "implementation_manual"
    }
    
    Returns learned questionnaire with smart questions.
    """
    try:
        data = request.get_json() or {}
        deliverable_type = data.get('deliverable_type', 'implementation_manual')
        
        # Get knowledge base from app
        knowledge_base = getattr(current_app, 'knowledge_base', None)
        
        ci = get_collective_intelligence(knowledge_base=knowledge_base)
        questions = ci.get_questionnaire(deliverable_type)
        
        if not questions:
            return jsonify({
                'success': False,
                'error': f'No questionnaire available for {deliverable_type}'
            }), 404
        
        return jsonify({
            'success': True,
            'deliverable_type': deliverable_type,
            'questions': questions,
            'total_questions': len(questions),
            'required_questions': len([q for q in questions if q.get('required', False)]),
            'message': f'Questionnaire ready with {len(questions)} questions'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@collective_bp.route('/api/collective/generate', methods=['POST'])
def generate_deliverable():
    """
    Generate a deliverable based on questionnaire answers.
    
    Body:
    {
        "deliverable_type": "implementation_manual",
        "answers": {
            "client_name": "Acme Corp",
            "industry": "Manufacturing",
            ...
        }
    }
    
    Returns generated deliverable content.
    """
    try:
        data = request.get_json() or {}
        deliverable_type = data.get('deliverable_type', 'implementation_manual')
        answers = data.get('answers', {})
        
        if not answers:
            return jsonify({
                'success': False,
                'error': 'No answers provided'
            }), 400
        
        # Get knowledge base from app
        knowledge_base = getattr(current_app, 'knowledge_base', None)
        
        ci = get_collective_intelligence(knowledge_base=knowledge_base)
        content = ci.generate_deliverable(deliverable_type, answers)
        
        if not content:
            return jsonify({
                'success': False,
                'error': f'Could not generate {deliverable_type}'
            }), 500
        
        # Save as document
        from database import save_generated_document
        import os
        from datetime import datetime
        
        # Convert to Word document
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # Add content
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    if line.startswith('#'):
                        # Header
                        level = len(line) - len(line.lstrip('#'))
                        text = line.lstrip('#').strip()
                        doc.add_heading(text, level=min(level, 3))
                    else:
                        # Paragraph
                        p = doc.add_paragraph(line)
                        p.style.font.size = Pt(11)
            
            # Save
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            client = answers.get('client_name', 'Client').replace(' ', '_')
            filename = f'implementation_manual_{client}_{timestamp}.docx'
            filepath = f'/tmp/{filename}'
            
            doc.save(filepath)
            
            file_size = os.path.getsize(filepath)
            doc_id = save_generated_document(
                filename=filename,
                original_name=f"Implementation Manual - {answers.get('client_name', 'Client')}",
                document_type='docx',
                file_path=filepath,
                file_size=file_size,
                task_id=None,
                conversation_id=None,
                project_id=None,
                title=f"Implementation Manual - {answers.get('client_name', 'Client')}",
                description=f"Auto-generated using collective intelligence",
                category='implementation_manual'
            )
            
            return jsonify({
                'success': True,
                'deliverable_type': deliverable_type,
                'content': content,
                'document_id': doc_id,
                'document_url': f'/api/generated-documents/{doc_id}/download',
                'message': f'{deliverable_type} generated successfully'
            })
            
        except Exception as doc_error:
            # Return content even if document creation fails
            return jsonify({
                'success': True,
                'deliverable_type': deliverable_type,
                'content': content,
                'message': f'{deliverable_type} generated (document creation failed: {str(doc_error)})'
            })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@collective_bp.route('/api/collective/patterns', methods=['GET'])
def get_patterns():
    """
    Get learned patterns from collective intelligence.
    
    Query params:
    - category: Filter by category (optional)
    - limit: Max results (default 20)
    """
    try:
        category = request.args.get('category', None)
        limit = request.args.get('limit', 20, type=int)
        
        ci = get_collective_intelligence()
        db = sqlite3.connect(ci.db_path)
        db.row_factory = sqlite3.Row
        cursor = db.cursor()
        
        if category:
            cursor.execute('''
                SELECT * FROM collective_patterns
                WHERE pattern_category = ?
                ORDER BY confidence DESC, created_at DESC
                LIMIT ?
            ''', (category, limit))
        else:
            cursor.execute('''
                SELECT * FROM collective_patterns
                ORDER BY confidence DESC, created_at DESC
                LIMIT ?
            ''', (limit,))
        
        patterns = [dict(row) for row in cursor.fetchall()]
        db.close()
        
        return jsonify({
            'success': True,
            'count': len(patterns),
            'patterns': patterns
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@collective_bp.route('/api/collective/analyze-document', methods=['POST'])
def analyze_document():
    """
    Analyze an uploaded document using collective intelligence.
    
    Body:
    {
        "document_content": "Full text of document",
        "document_type": "eaf" or "oaf",
        "metadata": {
            "client_name": "Acme Corp",
            "industry": "Manufacturing",
            ...
        }
    }
    
    Returns intelligent analysis with insights and recommendations.
    """
    try:
        data = request.get_json() or {}
        document_content = data.get('document_content', '')
        document_type = data.get('document_type', '')
        metadata = data.get('metadata', {})
        
        if not document_content:
            return jsonify({
                'success': False,
                'error': 'No document content provided'
            }), 400
        
        if not document_type:
            return jsonify({
                'success': False,
                'error': 'No document type specified',
                'supported_types': ['eaf', 'oaf']
            }), 400
        
        # Get knowledge base from app
        knowledge_base = getattr(current_app, 'knowledge_base', None)
        
        ci = get_collective_intelligence(knowledge_base=knowledge_base)
        analysis = ci.analyze_document(document_content, document_type, metadata)
        
        if 'error' in analysis:
            return jsonify({
                'success': False,
                'error': analysis['error'],
                'supported_types': analysis.get('supported_types', [])
            }), 400
        
        return jsonify({
            'success': True,
            'document_type': document_type,
            'analysis': analysis,
            'message': f'{document_type.upper()} analysis complete'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@collective_bp.route('/api/collective/conversational-generate', methods=['POST'])
def conversational_generate():
    """
    Start conversational deliverable generation.
    
    Body:
    {
        "request": "I need an implementation manual"
    }
    
    This endpoint intelligently asks questions and builds the deliverable.
    """
    try:
        data = request.get_json() or {}
        user_request = data.get('request', '')
        conversation_state = data.get('conversation_state', {})
        
        # Detect deliverable type from request
        deliverable_type = None
        if 'implementation manual' in user_request.lower():
            deliverable_type = 'implementation_manual'
        elif 'oaf' in user_request.lower() or 'operations assessment' in user_request.lower():
            deliverable_type = 'oaf'
        elif 'eaf' in user_request.lower() or 'employee assessment' in user_request.lower() or 'employee feedback' in user_request.lower():
            deliverable_type = 'eaf'
        
        if not deliverable_type:
            return jsonify({
                'success': False,
                'error': 'Could not determine deliverable type from request',
                'suggestion': 'Try: "I need an implementation manual" or "I need an OAF" or "I need an EAF"'
            }), 400
        
        # Get questionnaire
        knowledge_base = getattr(current_app, 'knowledge_base', None)
        ci = get_collective_intelligence(knowledge_base=knowledge_base)
        questions = ci.get_questionnaire(deliverable_type)
        
        # If no state, start questionnaire
        if not conversation_state:
            return jsonify({
                'success': True,
                'action': 'ask_questions',
                'deliverable_type': deliverable_type,
                'next_question': questions[0] if questions else None,
                'total_questions': len(questions),
                'current_question_index': 0,
                'conversation_state': {
                    'deliverable_type': deliverable_type,
                    'questions': questions,
                    'answers': {},
                    'current_index': 0
                },
                'message': f"I'll help you create an {deliverable_type}. Let me ask you a few questions."
            })
        
        # Continue questionnaire
        current_index = conversation_state.get('current_index', 0)
        questions_list = conversation_state.get('questions', [])
        answers = conversation_state.get('answers', {})
        
        # Check if we have all answers
        if current_index >= len(questions_list):
            # Generate deliverable
            content = ci.generate_deliverable(deliverable_type, answers)
            
            return jsonify({
                'success': True,
                'action': 'generated',
                'deliverable_type': deliverable_type,
                'content': content,
                'answers': answers,
                'message': f'{deliverable_type} generated successfully!'
            })
        
        # Return next question
        next_question = questions_list[current_index]
        
        return jsonify({
            'success': True,
            'action': 'ask_questions',
            'next_question': next_question,
            'current_question_index': current_index,
            'total_questions': len(questions_list),
            'conversation_state': conversation_state
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# I did no harm and this file is not truncated
