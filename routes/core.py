"""
Core Routes
Created: January 21, 2026 
Last Updated: January 26, 2026 - REPLACED OLD SCHEDULE SYSTEM WITH PATTERN-BASED

CHANGES IN THIS VERSION:
- January 26, 2026: COMPLETE SCHEDULE SYSTEM REPLACEMENT
  * Removed old named-schedule system (DuPont, Panama, Pitman buttons)
  * Integrated new pattern-based conversational system
  * Now asks: shift length → pattern selection
  * Only DuPont and Southern Swing keep their names (industry standards)
  * All other patterns described by work/off rhythm (2-2-3, 4-4, etc.)
  * Added session management for multi-turn schedule conversations
  * Generates visual Excel schedules with color-coded shifts

- January 23, 2026: FIXED CLARIFICATION LOOP BUG
  * Added clarification_answers parsing from request (JSON and form data)
  * Modified proactive pre-check to SKIP when clarification_answers provided
  * This fixes the infinite loop where AI kept asking same questions
  * User answers now bypass the SmartQuestioner and go straight to task execution

- January 23, 2026: ADDED MISSING API ENDPOINTS TO FIX FRONTEND ERRORS
  * Added /api/projects - List projects (was causing 404)
  * Added /api/projects/<id> - Get single project
  * Added /api/survey/questions - Get survey questions (was causing 404)
  * Added /api/survey/create - Create new survey
  * Added /api/survey/<id>/responses - Get survey responses
  * Added /api/marketing/status - Get marketing hub status (was causing 404)
  * Added /api/marketing/generate - Generate social content
  * Added /api/marketing/post - Post to social media
  * Added /api/opportunities - List opportunities
  * Added /api/opportunities/status - Get opportunity finder status (was causing 404)
  * Added /api/opportunities/top - Get top opportunities
  * Added /api/opportunities/pitch - Generate opportunity pitch
  * These endpoints integrate with existing utility classes

- January 23, 2026: DOCUMENT MANAGEMENT ENDPOINTS
  * Added GET /api/generated-documents - List all generated documents
  * Added GET /api/generated-documents/<id> - Get single document info
  * Added GET /api/generated-documents/<id>/download - Download document
  * Added GET /api/generated-documents/<id>/print - Get printable version
  * Added DELETE /api/generated-documents/<id> - Delete document
  * Added GET /api/documents/stats - Get document statistics

- January 22, 2026: PERSISTENT CONVERSATION MEMORY
- January 22, 2026: CRITICAL BUG FIX - AI now actually completes tasks
- January 22, 2026: Added proactive intelligence (smart questioning + suggestions)
- January 21, 2026: Added feedback endpoint and learning stats

Author: Jim @ Shiftwork Solutions LLC
"""

from flask import Blueprint, request, jsonify, send_file, current_app, make_response, session
import time
import json
import os
import markdown
import shutil
from datetime import datetime
from schedule_request_handler_combined import get_combined_schedule_handler
from database import (
    get_db, 
    create_conversation, 
    get_conversation, 
    get_conversations,
    update_conversation,
    delete_conversation as db_delete_conversation,
    add_message,
    get_messages,
    get_conversation_context,
    save_generated_document,
    get_generated_documents,
    get_generated_document,
    get_generated_document_by_filename,
    update_document_access,
    delete_generated_document,
    get_document_stats,
    get_schedule_context,        
    save_schedule_context         
)
from orchestration import (
    analyze_task_with_sonnet,
    handle_with_opus,
    execute_specialist_task,
    validate_with_consensus
)
from database_file_management import (
    save_project_file, get_project_files, get_all_projects_with_files,
    get_project_file_by_id, get_project_file_by_name, delete_project_file,
    get_file_stats_by_project
)
from werkzeug.utils import secure_filename
import shutil
from code_assistant_agent import get_code_assistant
from orchestration.proactive_agent import ProactiveAgent

core_bp = Blueprint('core', __name__)


# ============================================================================
# IMPORT OPTIONAL MODULES (Survey, Marketing, Opportunities)
# ============================================================================

# Survey Builder
SURVEY_BUILDER_AVAILABLE = False
survey_builder = None
try:
    from survey_builder import get_survey_builder
    survey_builder = get_survey_builder()
    SURVEY_BUILDER_AVAILABLE = True
    print("✅ Survey Builder loaded for API endpoints")
except ImportError:
    print("ℹ️  Survey Builder not available")
except Exception as e:
    print(f"⚠️  Survey Builder error: {e}")

# Marketing Hub
MARKETING_HUB_AVAILABLE = False
marketing_hub = None
try:
    from marketing_hub import get_marketing_hub
    marketing_hub = get_marketing_hub()
    MARKETING_HUB_AVAILABLE = True
    print("✅ Marketing Hub loaded for API endpoints")
except ImportError:
    print("ℹ️  Marketing Hub not available")
except Exception as e:
    print(f"⚠️  Marketing Hub error: {e}")

# Opportunity Finder
OPPORTUNITY_FINDER_AVAILABLE = False
opportunity_finder = None
try:
    from opportunity_finder import get_opportunity_finder
    opportunity_finder = get_opportunity_finder()
    OPPORTUNITY_FINDER_AVAILABLE = True
    print("✅ Opportunity Finder loaded for API endpoints")
except ImportError:
    print("ℹ️  Opportunity Finder not available")
except Exception as e:
    print(f"⚠️  Opportunity Finder error: {e}")

# Project Manager
PROJECT_MANAGER_AVAILABLE = False
project_manager = None
try:
    from project_manager import ProjectManager
    project_manager = ProjectManager()
    PROJECT_MANAGER_AVAILABLE = True
    print("✅ Project Manager loaded for API endpoints")
except ImportError:
    print("ℹ️  Project Manager not available")
except Exception as e:
    print(f"⚠️  Project Manager error: {e}")


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def convert_markdown_to_html(text):
    """Convert markdown text to styled HTML"""
    if not text:
        return text
    html = markdown.markdown(text, extensions=['extra', 'nl2br'])
    return f'<div style="line-height: 1.8; color: #333;">{html}</div>'


def should_create_document(user_request):
    """Determine if we should create a downloadable document"""
    doc_keywords = ['create', 'generate', 'make', 'write', 'draft', 'prepare',
                    'document', 'report', 'proposal', 'presentation', 'schedule',
                    'contract', 'agreement', 'summary', 'analysis']
    request_lower = user_request.lower()
    for keyword in doc_keywords:
        if keyword in request_lower:
            if 'presentation' in request_lower or 'powerpoint' in request_lower or 'slides' in request_lower:
                return True, 'pptx'
            elif 'spreadsheet' in request_lower or 'excel' in request_lower:
                return True, 'xlsx'
            elif 'pdf' in request_lower:
                return True, 'pdf'
            else:
                return True, 'docx'
    return False, None


def get_knowledge_context_for_prompt(knowledge_base, user_request, max_context=3000):
    """Get knowledge context to include in completion prompts"""
    if not knowledge_base:
        return ""
    try:
        context = knowledge_base.get_context_for_task(user_request, max_context=max_context)
        if context:
            return f"\n\n=== SHIFTWORK SOLUTIONS KNOWLEDGE BASE ===\nUse this information from our 30+ years of expertise:\n\n{context}\n\n=== END KNOWLEDGE BASE ===\n\n"
        return ""
    except Exception as e:
        print(f"Knowledge context retrieval failed: {e}")
        return ""


def generate_document_title(user_request, doc_type):
    """Generate a human-readable title from the user request"""
    title = user_request.strip()
    if title:
        title = title[0].upper() + title[1:]
    if len(title) > 60:
        title = title[:57] + '...'
    return title


def categorize_document(user_request, doc_type):
    """Determine document category based on request content"""
    request_lower = user_request.lower()
    if any(word in request_lower for word in ['schedule', 'dupont', 'panama', 'pitman', '2-2-3']):
        return 'schedule'
    elif any(word in request_lower for word in ['proposal', 'bid', 'quote']):
        return 'proposal'
    elif any(word in request_lower for word in ['report', 'analysis', 'summary']):
        return 'report'
    elif any(word in request_lower for word in ['contract', 'agreement', 'legal']):
        return 'contract'
    elif any(word in request_lower for word in ['checklist', 'list', 'todo']):
        return 'checklist'
    elif any(word in request_lower for word in ['email', 'letter', 'memo']):
        return 'communication'
    elif any(word in request_lower for word in ['presentation', 'slides', 'deck']):
        return 'presentation'
    else:
        return 'general'


# ============================================================================
# PROJECTS API ENDPOINTS
# ============================================================================

@core_bp.route('/api/projects', methods=['GET'])
def list_projects():
    """List all projects."""
    try:
        status_filter = request.args.get('status', 'all')
        limit = request.args.get('limit', 50, type=int)
        
        db = get_db()
        
        if status_filter and status_filter != 'all':
            projects = db.execute('''
                SELECT id, project_id, client_name, industry, facility_type, 
                       project_phase, status, created_at, updated_at
                FROM projects
                WHERE status = ?
                ORDER BY created_at DESC
                LIMIT ?
            ''', (status_filter, limit)).fetchall()
        else:
            projects = db.execute('''
                SELECT id, project_id, client_name, industry, facility_type,
                       project_phase, status, created_at, updated_at
                FROM projects
                ORDER BY created_at DESC
                LIMIT ?
            ''', (limit,)).fetchall()
        
        db.close()
        
        project_list = []
        for proj in projects:
            project_list.append({
                'id': proj['id'],
                'project_id': proj['project_id'],
                'client_name': proj['client_name'],
                'industry': proj['industry'],
                'facility_type': proj['facility_type'],
                'project_phase': proj['project_phase'],
                'status': proj['status'],
                'created_at': proj['created_at'],
                'updated_at': proj['updated_at']
            })
        
        return jsonify({
            'success': True,
            'projects': project_list,
            'count': len(project_list)
        })
    except Exception as e:
        print(f"Error listing projects: {e}")
        return jsonify({
            'success': True,
            'projects': [],
            'count': 0,
            'message': 'No projects found or projects table not initialized'
        })


@core_bp.route('/api/projects', methods=['POST'])
def create_project():
    """Create a new project"""
    try:
        data = request.json or {}
        
        import uuid
        project_id = str(uuid.uuid4())[:8]
        
        client_name = data.get('client_name', 'New Client')
        industry = data.get('industry', 'Manufacturing')
        facility_type = data.get('facility_type', 'Production')
        
        db = get_db()
        db.execute('''
            INSERT INTO projects (project_id, client_name, industry, facility_type, status, project_phase)
            VALUES (?, ?, ?, ?, 'active', 'initial')
        ''', (project_id, client_name, industry, facility_type))
        db.commit()
        db.close()
        
        return jsonify({
            'success': True,
            'project_id': project_id,
            'message': 'Project created successfully'
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# LEGACY PROJECT ENDPOINTS - ADDED January 28, 2026
# These support the old swarm-app.js interface
# ============================================================================

@core_bp.route('/api/project/start', methods=['POST'])
def start_project_legacy():
    """
    Start a new project (singular endpoint - used by swarm-app.js)
    FIXED: January 28, 2026 - Added missing endpoint
    """
    try:
        data = request.json or {}
        
        import uuid
        project_id = str(uuid.uuid4())[:8]
        
        client_name = data.get('client_name', 'New Client')
        industry = data.get('industry', 'Manufacturing')
        facility_type = data.get('facility_type', 'Production')
        
        db = get_db()
        db.execute('''
            INSERT INTO projects (project_id, client_name, industry, facility_type, status, project_phase)
            VALUES (?, ?, ?, ?, 'active', 'initial')
        ''', (project_id, client_name, industry, facility_type))
        db.commit()
        db.close()
        
        return jsonify({
            'success': True,
            'project_id': project_id,
            'client_name': client_name,
            'message': f'Project workflow started for {client_name}',
            'suggested_first_step': 'Upload data collection documents or employee surveys'
        })
    except Exception as e:
        print(f"Error starting project: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/project/<project_id>/context', methods=['GET'])
def get_project_context_legacy(project_id):
    """
    Get project context (singular endpoint - used by swarm-app.js)
    FIXED: January 28, 2026 - Added missing endpoint
    """
    try:
        db = get_db()
        
        project = db.execute('''
            SELECT * FROM projects WHERE project_id = ? OR id = ?
        ''', (project_id, project_id)).fetchone()
        
        db.close()
        
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        
        return jsonify({
            'success': True,
            'project_id': project['project_id'],
            'client_name': project['client_name'],
            'industry': project['industry'],
            'facility_type': project['facility_type'],
            'phase': project['project_phase'],
            'project_phase': project['project_phase'],
            'status': project['status'],
            'created_at': project['created_at']
        })
    except Exception as e:
        print(f"Error getting project context: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@core_bp.route('/api/projects/<project_id>', methods=['GET'])
def get_project(project_id):
    """Get a specific project by ID"""
    try:
        db = get_db()
        
        project = db.execute('''
            SELECT * FROM projects WHERE project_id = ? OR id = ?
        ''', (project_id, project_id)).fetchone()
        
        db.close()
        
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        
        return jsonify({
            'success': True,
            'project': dict(project)
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/projects/<project_id>', methods=['PUT'])
def update_project(project_id):
    """Update a project"""
    try:
        data = request.json or {}
        
        db = get_db()
        
        updates = []
        values = []
        
        if 'client_name' in data:
            updates.append('client_name = ?')
            values.append(data['client_name'])
        if 'industry' in data:
            updates.append('industry = ?')
            values.append(data['industry'])
        if 'status' in data:
            updates.append('status = ?')
            values.append(data['status'])
        if 'project_phase' in data:
            updates.append('project_phase = ?')
            values.append(data['project_phase'])
        
        updates.append('updated_at = CURRENT_TIMESTAMP')
        
        if updates:
            values.append(project_id)
            db.execute(f'''
                UPDATE projects SET {', '.join(updates)}
                WHERE project_id = ? OR id = ?
            ''', values + [project_id])
            db.commit()
        
        db.close()
        
        return jsonify({'success': True, 'message': 'Project updated'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# SURVEY API ENDPOINTS
# ============================================================================

@core_bp.route('/api/survey/questions', methods=['GET'])
def get_survey_questions():
    """Get available survey questions from the question bank."""
    try:
        if not SURVEY_BUILDER_AVAILABLE or not survey_builder:
            return jsonify({
                'success': True,
                'questions': [],
                'categories': [],
                'total_count': 0,
                'message': 'Survey Builder not available'
            })
        
        questions = []
        categories = set()
        
        for q_id, q_data in survey_builder.question_bank.items():
            questions.append({
                'id': q_id,
                'text': q_data.get('text', ''),
                'type': q_data.get('type', 'text'),
                'category': q_data.get('category', 'general'),
                'required': q_data.get('required', False),
                'options': q_data.get('options', [])
            })
            categories.add(q_data.get('category', 'general'))
        
        return jsonify({
            'success': True,
            'questions': questions,
            'categories': list(categories),
            'total_count': len(questions)
        })
    except Exception as e:
        print(f"Error getting survey questions: {e}")
        return jsonify({
            'success': True,
            'questions': [],
            'categories': [],
            'total_count': 0,
            'message': 'Could not load survey questions'
        })


@core_bp.route('/api/survey/create', methods=['POST'])
def create_survey():
    """Create a new survey"""
    try:
        if not SURVEY_BUILDER_AVAILABLE or not survey_builder:
            return jsonify({'success': False, 'error': 'Survey Builder not available'}), 503
        
        data = request.json or {}
        project_name = data.get('project_name', 'New Survey')
        selected_questions = data.get('questions', [])
        custom_questions = data.get('custom_questions', [])
        
        survey = survey_builder.create_survey(
            project_name=project_name,
            selected_questions=selected_questions,
            custom_questions=custom_questions
        )
        
        return jsonify({
            'success': True,
            'survey': survey
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/survey/<survey_id>', methods=['GET'])
def get_survey(survey_id):
    """Get a specific survey"""
    try:
        if not SURVEY_BUILDER_AVAILABLE or not survey_builder:
            return jsonify({'success': False, 'error': 'Survey Builder not available'}), 503
        
        survey = survey_builder.get_survey(survey_id)
        
        if not survey:
            return jsonify({'success': False, 'error': 'Survey not found'}), 404
        
        return jsonify({
            'success': True,
            'survey': survey
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/survey/<survey_id>/submit', methods=['POST'])
def submit_survey_response(survey_id):
    """Submit a response to a survey"""
    try:
        if not SURVEY_BUILDER_AVAILABLE or not survey_builder:
            return jsonify({'success': False, 'error': 'Survey Builder not available'}), 503
        
        data = request.json or {}
        answers = data.get('answers', {})
        respondent_id = data.get('respondent_id')
        
        result = survey_builder.submit_response(survey_id, answers, respondent_id)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/survey/<survey_id>/analyze', methods=['GET'])
def analyze_survey(survey_id):
    """Analyze survey responses"""
    try:
        if not SURVEY_BUILDER_AVAILABLE or not survey_builder:
            return jsonify({'success': False, 'error': 'Survey Builder not available'}), 503
        
        analysis = survey_builder.analyze_responses(survey_id)
        
        return jsonify(analysis)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/survey/status', methods=['GET'])
def get_survey_status():
    """Get survey system status"""
    return jsonify({
        'success': True,
        'available': SURVEY_BUILDER_AVAILABLE,
        'question_count': len(survey_builder.question_bank) if survey_builder else 0,
        'active_surveys': len(survey_builder.surveys) if survey_builder else 0
    })


# ============================================================================
# MARKETING API ENDPOINTS
# ============================================================================

@core_bp.route('/api/marketing/status', methods=['GET'])
def get_marketing_status():
    """Get marketing hub status."""
    try:
        if not MARKETING_HUB_AVAILABLE or not marketing_hub:
            return jsonify({
                'success': True,
                'available': False,
                'platforms': {
                    'linkedin': False,
                    'twitter': False,
                    'facebook': False
                },
                'message': 'Marketing Hub not available'
            })
        
        status = marketing_hub.get_status()
        
        return jsonify({
            'success': True,
            'available': True,
            'platforms': status.get('platforms', {}),
            'total_configured': status.get('total_configured', 0),
            'ready': status.get('ready', False)
        })
    except Exception as e:
        print(f"Error getting marketing status: {e}")
        return jsonify({
            'success': True,
            'available': False,
            'platforms': {},
            'message': 'Could not get marketing status'
        })


@core_bp.route('/api/marketing/generate', methods=['POST'])
def generate_marketing_content():
    """Generate social media content"""
    try:
        if not MARKETING_HUB_AVAILABLE or not marketing_hub:
            return jsonify({'success': False, 'error': 'Marketing Hub not available'}), 503
        
        data = request.json or {}
        topic = data.get('topic', '')
        platform = data.get('platform', 'linkedin')
        
        if not topic:
            return jsonify({'success': False, 'error': 'Topic is required'}), 400
        
        from orchestration.ai_clients import call_claude_sonnet
        
        result = marketing_hub.generate_social_content(topic, platform, call_claude_sonnet)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/marketing/post', methods=['POST'])
def post_to_social():
    """Post content to social media"""
    try:
        if not MARKETING_HUB_AVAILABLE or not marketing_hub:
            return jsonify({'success': False, 'error': 'Marketing Hub not available'}), 503
        
        data = request.json or {}
        content = data.get('content', '')
        platform = data.get('platform', 'linkedin')
        
        if not content:
            return jsonify({'success': False, 'error': 'Content is required'}), 400
        
        if platform == 'linkedin':
            result = marketing_hub.post_to_linkedin(content)
        else:
            result = {'success': False, 'error': f'Platform {platform} not supported yet'}
        
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/marketing/research', methods=['POST'])
def conduct_market_research():
    """Conduct AI-powered market research"""
    try:
        if not MARKETING_HUB_AVAILABLE or not marketing_hub:
            return jsonify({'success': False, 'error': 'Marketing Hub not available'}), 503
        
        data = request.json or {}
        topic = data.get('topic', '')
        
        if not topic:
            return jsonify({'success': False, 'error': 'Topic is required'}), 400
        
        from orchestration.ai_clients import call_claude_sonnet
        
        result = marketing_hub.conduct_market_research(topic, call_claude_sonnet)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# OPPORTUNITIES API ENDPOINTS
# ============================================================================

@core_bp.route('/api/opportunities', methods=['GET'])
def list_opportunities():
    """List all business opportunities"""
    try:
        if not OPPORTUNITY_FINDER_AVAILABLE or not opportunity_finder:
            return jsonify({
                'success': True,
                'opportunities': [],
                'message': 'Opportunity Finder not available'
            })
        
        opportunities = []
        for opp in opportunity_finder.opportunity_templates:
            opportunities.append({
                'name': opp['name'],
                'tagline': opp['tagline'],
                'target_market': opp['target_market'],
                'revenue_model': opp['revenue_model'],
                'estimated_revenue': opp['estimated_revenue'],
                'effort_to_launch': opp['effort_to_launch']
            })
        
        return jsonify({
            'success': True,
            'opportunities': opportunities,
            'count': len(opportunities)
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/opportunities/status', methods=['GET'])
def get_opportunities_status():
    """Get opportunity finder status."""
    try:
        if not OPPORTUNITY_FINDER_AVAILABLE or not opportunity_finder:
            return jsonify({
                'success': True,
                'available': False,
                'total_opportunities': 0,
                'capabilities_analyzed': 0,
                'schedules_count': 0,
                'metrics_count': 0,
                'message': 'Opportunity Finder not available'
            })
        
        stats = opportunity_finder.get_stats()
        
        return jsonify({
            'success': True,
            'available': True,
            'total_opportunities': stats.get('total_opportunities', 0),
            'capabilities_analyzed': stats.get('capabilities_analyzed', 0),
            'market_trends_tracked': stats.get('market_trends_tracked', 0),
            'schedules_count': stats.get('schedules_count', 0),
            'metrics_count': stats.get('metrics_count', 0),
            'top_opportunities': stats.get('top_opportunities', [])
        })
    except Exception as e:
        print(f"Error getting opportunities status: {e}")
        return jsonify({
            'success': True,
            'available': False,
            'total_opportunities': 0,
            'schedules_count': 0,
            'metrics_count': 0,
            'message': 'Could not get opportunities status'
        })


@core_bp.route('/api/opportunities/top', methods=['GET'])
def get_top_opportunities():
    """Get top-ranked business opportunities"""
    try:
        if not OPPORTUNITY_FINDER_AVAILABLE or not opportunity_finder:
            return jsonify({'success': False, 'error': 'Opportunity Finder not available'}), 503
        
        limit = request.args.get('limit', 5, type=int)
        
        top_opps = opportunity_finder.get_top_opportunities(limit=limit)
        
        return jsonify({
            'success': True,
            'opportunities': top_opps
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/opportunities/pitch', methods=['POST'])
def generate_opportunity_pitch():
    """Generate a pitch for a specific opportunity"""
    try:
        if not OPPORTUNITY_FINDER_AVAILABLE or not opportunity_finder:
            return jsonify({'success': False, 'error': 'Opportunity Finder not available'}), 503
        
        data = request.json or {}
        opportunity_name = data.get('name')
        
        pitch = opportunity_finder.generate_pitch(opportunity_name)
        
        return jsonify({
            'success': True,
            'pitch': pitch
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/opportunities/analyze', methods=['POST'])
def analyze_opportunity():
    """Analyze market fit for an opportunity"""
    try:
        if not OPPORTUNITY_FINDER_AVAILABLE or not opportunity_finder:
            return jsonify({'success': False, 'error': 'Opportunity Finder not available'}), 503
        
        data = request.json or {}
        opportunity_name = data.get('name')
        
        if not opportunity_name:
            return jsonify({'success': False, 'error': 'Opportunity name required'}), 400
        
        opportunity = next(
            (o for o in opportunity_finder.opportunity_templates 
             if o['name'].lower() == opportunity_name.lower()),
            None
        )
        
        if not opportunity:
            return jsonify({'success': False, 'error': 'Opportunity not found'}), 404
        
        analysis = opportunity_finder.analyze_market_fit(opportunity)
        
        return jsonify({
            'success': True,
            'analysis': analysis
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# GENERATED DOCUMENTS MANAGEMENT ENDPOINTS
# ============================================================================

@core_bp.route('/api/generated-documents', methods=['GET'])
def list_generated_documents():
    """Get list of all generated documents for display in sidebar."""
    try:
        limit = request.args.get('limit', 50, type=int)
        document_type = request.args.get('type')
        project_id = request.args.get('project_id')
        conversation_id = request.args.get('conversation_id')
        
        documents = get_generated_documents(
            limit=limit,
            document_type=document_type,
            project_id=project_id,
            conversation_id=conversation_id
        )
        
        formatted_docs = []
        for doc in documents:
            formatted_docs.append({
                'id': doc['id'],
                'filename': doc['filename'],
                'title': doc['title'] or doc['original_name'],
                'type': doc['document_type'],
                'size': doc['file_size'],
                'category': doc['category'],
                'created_at': doc['created_at'],
                'download_count': doc['download_count'],
                'download_url': f"/api/generated-documents/{doc['id']}/download",
                'print_url': f"/api/generated-documents/{doc['id']}/print"
            })
        
        return jsonify({
            'success': True,
            'documents': formatted_docs,
            'count': len(formatted_docs)
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/generated-documents/<int:document_id>', methods=['GET'])
def get_document_info(document_id):
    """Get information about a specific document"""
    try:
        doc = get_generated_document(document_id)
        if not doc:
            return jsonify({'success': False, 'error': 'Document not found'}), 404
        
        return jsonify({
            'success': True,
            'document': {
                'id': doc['id'],
                'filename': doc['filename'],
                'title': doc['title'] or doc['original_name'],
                'type': doc['document_type'],
                'size': doc['file_size'],
                'category': doc['category'],
                'description': doc['description'],
                'created_at': doc['created_at'],
                'last_accessed': doc['last_accessed'],
                'download_count': doc['download_count'],
                'task_id': doc['task_id'],
                'conversation_id': doc['conversation_id'],
                'project_id': doc['project_id'],
                'download_url': f"/api/generated-documents/{doc['id']}/download",
                'print_url': f"/api/generated-documents/{doc['id']}/print"
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/generated-documents/<int:document_id>/download', methods=['GET'])
def download_generated_document(document_id):
    """Download a generated document by ID"""
    try:
        doc = get_generated_document(document_id)
        if not doc:
            return jsonify({'error': 'Document not found'}), 404
        
        file_path = doc['file_path']
        if not os.path.exists(file_path):
            alternative_paths = [
                f"/mnt/user-data/outputs/{doc['filename']}",
                f"/tmp/{doc['filename']}",
                f"./{doc['filename']}"
            ]
            file_path = None
            for alt_path in alternative_paths:
                if os.path.exists(alt_path):
                    file_path = alt_path
                    break
            if not file_path:
                return jsonify({'error': 'File not found on server'}), 404
        
        update_document_access(document_id)
        
        mime_types = {
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'pdf': 'application/pdf',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
            'txt': 'text/plain',
            'csv': 'text/csv'
        }
        mime_type = mime_types.get(doc['document_type'], 'application/octet-stream')
        
        return send_file(file_path, as_attachment=True, download_name=doc['filename'], mimetype=mime_type)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@core_bp.route('/api/generated-documents/<int:document_id>/print', methods=['GET'])
def print_generated_document(document_id):
    """Get a printable version of the document."""
    try:
        doc = get_generated_document(document_id)
        if not doc:
            return jsonify({'error': 'Document not found'}), 404
        
        file_path = doc['file_path']
        if not os.path.exists(file_path):
            alternative_paths = [
                f"/mnt/user-data/outputs/{doc['filename']}",
                f"/tmp/{doc['filename']}",
                f"./{doc['filename']}"
            ]
            file_path = None
            for alt_path in alternative_paths:
                if os.path.exists(alt_path):
                    file_path = alt_path
                    break
            if not file_path:
                return jsonify({'error': 'File not found on server'}), 404
        
        update_document_access(document_id)
        
        if doc['document_type'] == 'docx':
            try:
                from docx import Document
                docx_doc = Document(file_path)
                content_html = ""
                for para in docx_doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        level = para.style.name[-1] if para.style.name[-1].isdigit() else '2'
                        content_html += f"<h{level}>{para.text}</h{level}>\n"
                    else:
                        if para.text.strip():
                            content_html += f"<p>{para.text}</p>\n"
                
                print_html = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{doc['title'] or doc['original_name']} - Print</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; max-width: 800px; margin: 0 auto; padding: 40px; line-height: 1.6; }}
        h1 {{ color: #333; border-bottom: 2px solid #667eea; padding-bottom: 10px; }}
        h2 {{ color: #667eea; margin-top: 30px; }}
        h3 {{ color: #764ba2; }}
        p {{ margin: 10px 0; text-align: justify; }}
        .header {{ text-align: center; margin-bottom: 40px; padding-bottom: 20px; border-bottom: 1px solid #ccc; }}
        .header h1 {{ border-bottom: none; }}
        .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #ccc; text-align: center; font-size: 12px; color: #666; }}
        @media print {{ body {{ padding: 20px; }} .no-print {{ display: none; }} }}
    </style>
</head>
<body>
    <div class="no-print" style="background: #f0f0f0; padding: 15px; margin-bottom: 20px; border-radius: 8px;">
        <button onclick="window.print()" style="padding: 10px 20px; background: #667eea; color: white; border: none; border-radius: 6px; cursor: pointer; font-weight: bold;">Print Document</button>
        <button onclick="window.close()" style="padding: 10px 20px; background: #e0e0e0; color: #333; border: none; border-radius: 6px; cursor: pointer; margin-left: 10px;">Close</button>
    </div>
    <div class="header">
        <h1>{doc['title'] or doc['original_name']}</h1>
        <p>Shiftwork Solutions LLC</p>
        <p style="font-size: 12px; color: #666;">Generated: {doc['created_at']}</p>
    </div>
    {content_html}
    <div class="footer">
        <p>Shiftwork Solutions LLC - Hundreds of Facilities Served</p>
        <p>www.shiftworksolutions.com</p>
    </div>
</body>
</html>'''
                response = make_response(print_html)
                response.headers['Content-Type'] = 'text/html'
                return response
            except Exception as docx_error:
                print(f"DOCX parsing failed: {docx_error}")
                return jsonify({
                    'success': False,
                    'error': 'Could not create printable version. Please download the document instead.',
                    'download_url': f"/api/generated-documents/{document_id}/download"
                }), 400
        
        return jsonify({
            'success': False,
            'error': f'{doc["document_type"].upper()} files must be downloaded and printed from the native application.',
            'download_url': f"/api/generated-documents/{document_id}/download"
        }), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@core_bp.route('/api/generated-documents/<int:document_id>', methods=['DELETE'])
def delete_document(document_id):
    """Delete a generated document"""
    try:
        doc = get_generated_document(document_id)
        if not doc:
            return jsonify({'success': False, 'error': 'Document not found'}), 404
        
        hard_delete = request.args.get('hard', 'false').lower() == 'true'
        success = delete_generated_document(document_id, hard_delete=hard_delete)
        
        if success:
            return jsonify({'success': True, 'message': 'Document deleted successfully'})
        else:
            return jsonify({'success': False, 'error': 'Failed to delete document'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/documents/stats', methods=['GET'])
def get_docs_stats():
    """Get document statistics"""
    try:
        stats = get_document_stats()
        return jsonify({'success': True, 'stats': stats})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# CONVERSATION MEMORY ENDPOINTS
# ============================================================================

@core_bp.route('/api/conversations', methods=['GET'])
def list_conversations():
    """Get list of recent conversations"""
    try:
        limit = request.args.get('limit', 20, type=int)
        project_id = request.args.get('project_id')
        include_archived = request.args.get('include_archived', 'false').lower() == 'true'
        conversations = get_conversations(limit=limit, project_id=project_id, include_archived=include_archived)
        return jsonify({'success': True, 'conversations': conversations, 'count': len(conversations)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/conversations', methods=['POST'])
def create_new_conversation():
    """Create a new conversation"""
    try:
        data = request.json or {}
        mode = data.get('mode', 'quick')
        project_id = data.get('project_id')
        title = data.get('title')
        conversation_id = create_conversation(mode=mode, project_id=project_id, title=title)
        return jsonify({'success': True, 'conversation_id': conversation_id, 'message': 'Conversation created successfully'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/conversations/<conversation_id>', methods=['GET'])
def get_single_conversation(conversation_id):
    """Get a specific conversation with its messages"""
    try:
        conversation = get_conversation(conversation_id)
        if not conversation:
            return jsonify({'success': False, 'error': 'Conversation not found'}), 404
        messages = get_messages(conversation_id, limit=100)
        return jsonify({'success': True, 'conversation': conversation, 'messages': messages})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/conversations/<conversation_id>', methods=['PUT'])
def update_single_conversation(conversation_id):
    """Update conversation metadata"""
    try:
        conversation = get_conversation(conversation_id)
        if not conversation:
            return jsonify({'success': False, 'error': 'Conversation not found'}), 404
        data = request.json or {}
        update_conversation(conversation_id, title=data.get('title'), mode=data.get('mode'),
                          project_id=data.get('project_id'), is_archived=data.get('is_archived'))
        return jsonify({'success': True, 'message': 'Conversation updated successfully'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/conversations/<conversation_id>', methods=['DELETE'])
def delete_single_conversation(conversation_id):
    """Delete a conversation and all its messages"""
    try:
        conversation = get_conversation(conversation_id)
        if not conversation:
            return jsonify({'success': False, 'error': 'Conversation not found'}), 404
        db_delete_conversation(conversation_id)
        return jsonify({'success': True, 'message': 'Conversation deleted successfully'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/conversations/<conversation_id>/messages', methods=['POST'])
def add_conversation_message(conversation_id):
    """Add a message to a conversation"""
    try:
        conversation = get_conversation(conversation_id)
        if not conversation:
            return jsonify({'success': False, 'error': 'Conversation not found'}), 404
        data = request.json or {}
        role = data.get('role', 'user')
        content = data.get('content')
        task_id = data.get('task_id')
        metadata = data.get('metadata')
        if not content:
            return jsonify({'success': False, 'error': 'Content is required'}), 400
        add_message(conversation_id, role, content, task_id, metadata)
        return jsonify({'success': True, 'message': 'Message added successfully'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# MAIN ORCHESTRATION ENDPOINT - UPDATED January 26, 2026
# ============================================================================

@core_bp.route('/api/orchestrate', methods=['POST'])
def orchestrate():
    """
    Main orchestration endpoint with proactive intelligence and conversation memory.
    
    UPDATED January 26, 2026: Replaced old schedule system with pattern-based conversational system
    - Asks for shift length (12-hour or 8-hour)
    - Then shows available patterns for that shift length
    - Generates visual Excel schedules with color coding
    - Only DuPont and Southern Swing keep their names (industry standards)
    """
    try:
        if request.is_json:
            data = request.json
            user_request = data.get('request')
            enable_consensus = data.get('enable_consensus', True)
            project_id = data.get('project_id')
            conversation_id = data.get('conversation_id')
            mode = data.get('mode', 'quick')
        else:
            user_request = request.form.get('request')
            enable_consensus = request.form.get('enable_consensus', 'true').lower() == 'true'
            project_id = request.form.get('project_id')
            conversation_id = request.form.get('conversation_id')
            mode = request.form.get('mode', 'quick')
        
        if not user_request:
            return jsonify({'success': False, 'error': 'Request text required'}), 400
        
        overall_start = time.time()
        
        # Check for clarification answers
        clarification_answers = None
        if request.is_json:
            clarification_answers_raw = request.json.get('clarification_answers')
            if clarification_answers_raw:
                if isinstance(clarification_answers_raw, str):
                    try:
                        clarification_answers = json.loads(clarification_answers_raw)
                    except:
                        clarification_answers = None
                else:
                    clarification_answers = clarification_answers_raw
        else:
            clarification_answers_str = request.form.get('clarification_answers')
            if clarification_answers_str:
                try:
                    clarification_answers = json.loads(clarification_answers_str)
                except:
                    clarification_answers = None
        
        if clarification_answers:
            print(f"✅ Clarification answers received: {clarification_answers}")
            context_additions = []
            for field, value in clarification_answers.items():
                context_additions.append(f"{field}: {value}")
            if context_additions:
                user_request = f"{user_request}\n\nAdditional context from user:\n" + "\n".join(context_additions)
                print(f"Enhanced request with clarification: {user_request[:200]}...")
        
        # Create conversation if not provided
        if not conversation_id:
            conversation_id = create_conversation(mode=mode, project_id=project_id)
            print(f"Created new conversation: {conversation_id}")
        
        add_message(conversation_id, 'user', user_request)
        
        # Initialize proactive agent
        proactive = None
        try:
            proactive = ProactiveAgent()
        except Exception as proactive_init_error:
            print(f"Proactive agent init failed: {proactive_init_error}")
        
        # Proactive pre-check - SKIP if clarification answers provided
        if proactive and not clarification_answers:
            try:
                pre_check = proactive.pre_process_request(user_request)
                if pre_check['action'] == 'ask_questions':
                    db = get_db()
                    cursor = db.execute('INSERT INTO tasks (user_request, status, conversation_id) VALUES (?, ?, ?)',
                                       (user_request, 'needs_clarification', conversation_id))
                    task_id = cursor.lastrowid
                    db.commit()
                    db.close()
                    return jsonify({'success': True, 'needs_clarification': True, 'clarification_data': pre_check['data'],
                                   'task_id': task_id, 'conversation_id': conversation_id})
                if pre_check['action'] == 'detect_project':
                    return jsonify({'success': True, 'project_detected': True, 'project_data': pre_check['data'],
                                   'conversation_id': conversation_id})
            except Exception as proactive_error:
                print(f"Proactive check failed: {proactive_error}")
        elif clarification_answers:
            print("Skipping proactive check - user provided clarification answers")
        
        import sys
        app_module = sys.modules.get('app')
        knowledge_base = getattr(app_module, 'knowledge_base', None) if app_module else None
        conversation_context = get_conversation_context(conversation_id, max_messages=10)
        
        db = get_db()
        cursor = db.execute('INSERT INTO tasks (user_request, status, conversation_id) VALUES (?, ?, ?)',
                           (user_request, 'processing', conversation_id))
        task_id = cursor.lastrowid
        db.commit()
        # =============================================================================
        # CODE ASSISTANT - CHECK IF USER IS GIVING CODE FEEDBACK
        # =============================================================================

        # Get code assistant
        try:
            code_assistant_agent = get_code_assistant(knowledge_base=knowledge_base)
            
            # Check if this is code feedback
            feedback_check = code_assistant_agent.detect_code_feedback(user_request)
            
            if feedback_check['is_code_feedback']:
                # This is code feedback - route to code assistant
                print(f"🔧 Code feedback detected for: {feedback_check['target_file']}")
                
                from orchestration.ai_clients import call_claude_sonnet
                
                code_result = code_assistant_agent.process_code_feedback(
                    user_request,
                    call_claude_sonnet
                )
                
                if code_result['success']:
                    # Code was fixed!
                    deployment_pkg = code_result['deployment_package']
                 
                    # FIXED January 28, 2026: Initialize these variables BEFORE the try block
                    doc_id = None
                    document_url = None
                    
                    # Save fixed file
                    if code_result.get('output_path'):
                        try:
                            file_size = os.path.getsize(code_result['output_path'])
                            doc_id = save_generated_document(
                                filename=os.path.basename(code_result['output_path']),
                                original_name=f"Fixed: {code_result['target_file']}",
                                document_type='py',
                                file_path=code_result['output_path'],
                                file_size=file_size,
                                task_id=task_id,
                                conversation_id=conversation_id,
                                project_id=project_id,
                                title=f"Code Fix: {code_result['target_file']}",
                                description=deployment_pkg['change_summary'],
                                category='code'
                            )
                            
                            document_url = f'/api/generated-documents/{doc_id}/download'
                        except Exception as doc_error:
                            print(f"Could not save code file: {doc_error}")
                            document_url = None
                    else:
                        document_url = None
                    
                    # Convert response to HTML
                    response_html = convert_markdown_to_html(code_result['message'])
                    
                    # Update task as completed
                    db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                              ('completed', 'code_assistant', time.time() - overall_start, task_id))
                    db.commit()
                    
                    # Add message to conversation
                    add_message(conversation_id, 'assistant', code_result['message'], task_id,
                               {'document_created': True, 'document_type': 'py', 'document_id': doc_id,
                                'orchestrator': 'code_assistant', 'target_file': code_result['target_file']})
                    
                    # Get suggestions
                    suggestions = []
                    if proactive:
                        try:
                            suggestions = proactive.post_process_result(task_id, user_request, code_result['message'])
                        except:
                            pass
                    
                    db.close()
                    
                    # Return code fix response
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'conversation_id': conversation_id,
                        'result': response_html,
                        'document_url': document_url,
                        'document_id': doc_id,
                        'document_created': True,
                        'document_type': 'py',
                        'execution_time': time.time() - overall_start,
                        'orchestrator': 'code_assistant',
                        'target_file': code_result['target_file'],
                        'deployment_instructions': deployment_pkg['deployment_instructions'],
                        'suggestions': suggestions
                    })
        
        except Exception as code_assistant_error:
            print(f"Code Assistant failed: {code_assistant_error}")
            # Continue to normal orchestration
        # =============================================================================
        # NEW PATTERN-BASED SCHEDULE SYSTEM - START
        # =============================================================================
        
        # Pattern-based schedule generator intercept
        schedule_handler = get_combined_schedule_handler()
        
        # Get conversation context from session if it exists
        # FIXED: Get conversation context from DATABASE instead of session
        schedule_context = get_schedule_context(conversation_id)
        
        print(f"🔍 DEBUG: Loaded schedule context from DB: {schedule_context}")
        
        # Process the schedule request
        schedule_result = schedule_handler.process_request(user_request, schedule_context)
        
        print(f"🔍 DEBUG: Schedule result action: {schedule_result['action']}")
        
        if schedule_result['action'] != 'not_schedule_request':
            # This IS a schedule-related request
            
            # FIXED: Save returned context to DATABASE
            if 'context' in schedule_result:
                save_schedule_context(conversation_id, schedule_result['context'])
                print(f"💾 DEBUG: Saved context: {schedule_result['context']}")
            
            if schedule_result['action'] == 'generate_schedule':
                # Schedule was generated successfully
                source_filepath = schedule_result['filepath']
                filename = os.path.basename(source_filepath)
                
                # FIXED: Use the file directly from where it was saved (/tmp/)
                # Don't try to copy to /mnt/user-data/outputs (doesn't exist on Render)
                file_path = source_filepath
                
                # Save to database as generated document
                file_size = os.path.getsize(file_path)
                shift_length = schedule_result['shift_length']
                pattern_key = schedule_result['pattern_key']
                doc_title = f"{shift_length}-Hour {pattern_key.upper().replace('_', ' ')} Schedule Pattern"
                
                doc_id = save_generated_document(
                    filename=filename,
                    original_name=doc_title,
                    document_type='xlsx',
                    file_path=file_path,
                    file_size=file_size,
                    task_id=task_id,
                    conversation_id=conversation_id,
                    project_id=project_id,
                    title=doc_title,
                    description=f"Visual {shift_length}-hour {pattern_key} schedule pattern with color-coded shifts",
                    category='schedule'
                )
                
                # Convert markdown response to HTML
                response_html = convert_markdown_to_html(schedule_result['message'])
                
                # Update task as completed
                db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                          ('completed', 'pattern_schedule_generator', time.time() - overall_start, task_id))
                db.commit()
                
                # Add message to conversation
                add_message(conversation_id, 'assistant', schedule_result['message'], task_id,
                           {'document_created': True, 'document_type': 'xlsx', 'document_id': doc_id, 
                            'orchestrator': 'pattern_schedule_generator', 'shift_length': shift_length,
                            'pattern': pattern_key})
                
                # Clear schedule context from session
                session.pop('schedule_context', None)
                
                # Get proactive suggestions if available
                suggestions = []
                if proactive:
                    try:
                        suggestions = proactive.post_process_result(task_id, user_request, schedule_result['message'])
                    except:
                        pass
                
                db.close()
                
                # Return response with document
                return jsonify({
                    'success': True,
                    'task_id': task_id,
                    'conversation_id': conversation_id,
                    'result': response_html,
                    'document_url': f'/api/generated-documents/{doc_id}/download',
                    'document_id': doc_id,
                    'document_created': True,
                    'document_type': 'xlsx',
                    'execution_time': time.time() - overall_start,
                    'orchestrator': 'pattern_schedule_generator',
                    'knowledge_applied': False,
                    'formatting_applied': True,
                    'specialists_used': [],
                    'consensus': None,
                    'suggestions': suggestions,
                    'shift_length': shift_length,
                    'pattern': pattern_key
                })
            
            else:
                # Need more information - asking user a question
                # Context already saved to database above
                
                print(f"❓ DEBUG: Asking follow-up. Context in DB for conv {conversation_id}")
                
                # Convert response to HTML
                response_html = convert_markdown_to_html(schedule_result['message'])
                
                # Update task as in-progress
                db.execute('UPDATE tasks SET status = ? WHERE id = ?', ('in_progress', task_id))
                db.commit()
                
                # Add message to conversation
                add_message(conversation_id, 'assistant', schedule_result['message'], task_id,
                           {'waiting_for_input': True, 'orchestrator': 'pattern_schedule_generator',
                            'waiting_for': schedule_result.get('waiting_for')})
                
                db.close()
                
                # Return question to user
                return jsonify({
                    'success': True,
                    'task_id': task_id,
                    'conversation_id': conversation_id,
                    'result': response_html,
                    'needs_input': True,
                    'waiting_for': schedule_result.get('waiting_for'),
                    'orchestrator': 'pattern_schedule_generator',
                    'execution_time': time.time() - overall_start
                })
        
        # =============================================================================
        # NEW PATTERN-BASED SCHEDULE SYSTEM - END
        # =============================================================================
# =============================================================================
        # INTROSPECTION DETECTION - FIXED January 29, 2026
        # Must come BEFORE regular AI orchestration to intercept introspection requests
        # =============================================================================
        
        try:
            from introspection import is_introspection_request, get_introspection_engine
            
            introspection_check = is_introspection_request(user_request)
            
            if introspection_check['is_introspection']:
                print(f"🔍 Introspection request detected: {introspection_check['action']}")
                
                intro_engine = get_introspection_engine()
                action = introspection_check['action']
                
                if action == 'run':
                    print("Running introspection analysis...")
                    
                    days = 7
                    if 'monthly' in user_request.lower() or '30' in user_request:
                        days = 30
                    
                    report = intro_engine.run_introspection(days=days, is_monthly=(days >= 28))
                    response_html = convert_markdown_to_html(report.get('reflection', ''))
                    
                    db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                              ('completed', 'introspection_engine', time.time() - overall_start, task_id))
                    db.commit()
                    
                    add_message(conversation_id, 'assistant', report.get('reflection', ''), task_id,
                               {'orchestrator': 'introspection_engine', 'introspection_id': report.get('insight_id')})
                    
                    db.close()
                    
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'conversation_id': conversation_id,
                        'result': response_html,
                        'orchestrator': 'introspection_engine',
                        'execution_time': time.time() - overall_start,
                        'introspection_report': {
                            'id': report.get('insight_id'),
                            'health_score': report.get('summary', {}).get('health_score', 0),
                            'period_days': report.get('period_days', days),
                            'type': report.get('introspection_type', 'weekly')
                        }
                    })
                
                elif action == 'show_latest':
                    latest = intro_engine.get_latest_introspection()
                    
                    if not latest or not latest.get('full_report'):
                        response_text = "No introspection has been run yet. Would you like me to run one now? Just say 'run introspection'."
                    else:
                        response_text = latest.get('full_report', {}).get('reflection', 'No reflection available.')
                    
                    response_html = convert_markdown_to_html(response_text)
                    
                    db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ? WHERE id = ?',
                              ('completed', 'introspection_engine', task_id))
                    db.commit()
                    
                    add_message(conversation_id, 'assistant', response_text, task_id,
                               {'orchestrator': 'introspection_engine'})
                    
                    db.close()
                    
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'conversation_id': conversation_id,
                        'result': response_html,
                        'orchestrator': 'introspection_engine',
                        'execution_time': time.time() - overall_start
                    })
        
        except ImportError:
            print("ℹ️  Introspection not available - continuing to regular orchestration")
        except Exception as intro_error:
            print(f"⚠️  Introspection detection failed: {intro_error}")
        
        # =============================================================================
        # END INTROSPECTION DETECTION
        # =============================================================================
     # =============================================================================
        # INTROSPECTION DETECTION - FIXED January 29, 2026
        # Must come BEFORE regular AI orchestration to intercept introspection requests
        # =============================================================================
        
        try:
            from introspection import is_introspection_request, get_introspection_engine
            
            introspection_check = is_introspection_request(user_request)
            
            if introspection_check['is_introspection']:
                print(f"🔍 Introspection request detected: {introspection_check['action']}")
                
                intro_engine = get_introspection_engine()
                action = introspection_check['action']
                
                if action == 'run':
                    print("Running introspection analysis...")
                    
                    days = 7
                    if 'monthly' in user_request.lower() or '30' in user_request:
                        days = 30
                    
                    report = intro_engine.run_introspection(days=days, is_monthly=(days >= 28))
                    response_html = convert_markdown_to_html(report.get('reflection', ''))
                    
                    db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                              ('completed', 'introspection_engine', time.time() - overall_start, task_id))
                    db.commit()
                    
                    add_message(conversation_id, 'assistant', report.get('reflection', ''), task_id,
                               {'orchestrator': 'introspection_engine', 'introspection_id': report.get('insight_id')})
                    
                    db.close()
                    
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'conversation_id': conversation_id,
                        'result': response_html,
                        'orchestrator': 'introspection_engine',
                        'execution_time': time.time() - overall_start,
                        'introspection_report': {
                            'id': report.get('insight_id'),
                            'health_score': report.get('summary', {}).get('health_score', 0),
                            'period_days': report.get('period_days', days),
                            'type': report.get('introspection_type', 'weekly')
                        }
                    })
                
                elif action == 'show_latest':
                    latest = intro_engine.get_latest_introspection()
                    
                    if not latest or not latest.get('full_report'):
                        response_text = "No introspection has been run yet. Would you like me to run one now? Just say 'run introspection'."
                    else:
                        response_text = latest.get('full_report', {}).get('reflection', 'No reflection available.')
                    
                    response_html = convert_markdown_to_html(response_text)
                    
                    db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ? WHERE id = ?',
                              ('completed', 'introspection_engine', task_id))
                    db.commit()
                    
                    add_message(conversation_id, 'assistant', response_text, task_id,
                               {'orchestrator': 'introspection_engine'})
                    
                    db.close()
                    
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'conversation_id': conversation_id,
                        'result': response_html,
                        'orchestrator': 'introspection_engine',
                        'execution_time': time.time() - overall_start
                    })
        
        except ImportError:
            print("ℹ️  Introspection not available - continuing to regular orchestration")
        except Exception as intro_error:
            print(f"⚠️  Introspection detection failed: {intro_error}")
        
        # =============================================================================
        # END INTROSPECTION DETECTION
        # =============================================================================
        # Regular AI orchestration
        try:
            print(f"Analyzing task: {user_request[:100]}...")
            analysis = analyze_task_with_sonnet(user_request, knowledge_base=knowledge_base)
            task_type = analysis.get('task_type', 'general')
            confidence = analysis.get('confidence', 0.5)
            escalate = analysis.get('escalate_to_opus', False)
            specialists_needed = analysis.get('specialists_needed', [])
            knowledge_applied = analysis.get('knowledge_applied', False)
            knowledge_sources = analysis.get('knowledge_sources', [])
            
            if specialists_needed:
                specialists_needed = [s for s in specialists_needed if s and s.lower() != 'none']
            
            orchestrator = 'sonnet'
            opus_guidance = None
            
            if escalate:
                print("Escalating to Opus for strategic guidance...")
                orchestrator = 'opus'
                try:
                    opus_result = handle_with_opus(user_request, analysis, knowledge_base=knowledge_base)
                    opus_guidance = opus_result.get('strategic_analysis', '')
                    if opus_result.get('specialist_assignments'):
                        for assignment in opus_result.get('specialist_assignments', []):
                            specialist = assignment.get('ai') or assignment.get('specialist')
                            if specialist and specialist.lower() != 'none':
                                specialists_needed.append(specialist)
                except Exception as opus_error:
                    print(f"Opus guidance failed: {opus_error}")
            
            specialist_results = []
            specialist_output = None
            if specialists_needed:
                for specialist_info in specialists_needed:
                    if isinstance(specialist_info, dict):
                        specialist = specialist_info.get('specialist') or specialist_info.get('ai')
                        specialist_task = specialist_info.get('task', user_request)
                    else:
                        specialist = specialist_info
                        specialist_task = user_request
                    if specialist and specialist.lower() != 'none':
                        result = execute_specialist_task(specialist, specialist_task)
                        specialist_results.append(result)
                        if result.get('success') and result.get('output'):
                            specialist_output = result.get('output')
      
            from orchestration.ai_clients import call_claude_opus, call_claude_sonnet
            knowledge_context = get_knowledge_context_for_prompt(knowledge_base, user_request)
            
            conversation_history = ""
            if conversation_context and len(conversation_context) > 1:
                conversation_history = "\n\n=== CONVERSATION HISTORY ===\n"
                for msg in conversation_context[:-1]:
                    role_label = "User" if msg['role'] == 'user' else "Assistant"
                    content_preview = msg['content'][:500] + '...' if len(msg['content']) > 500 else msg['content']
                    conversation_history += f"{role_label}: {content_preview}\n"
                conversation_history += "=== END CONVERSATION HISTORY ===\n\n"
            
            if specialist_output:
                actual_output = specialist_output
            else:
                completion_prompt = f"""{knowledge_context}{conversation_history}

USER REQUEST: {user_request}

Please complete this request fully. Provide the actual deliverable the user is asking for.
Do not describe what you would do - actually do it.
Do not provide meta-commentary about the task - provide the actual content requested.

Be comprehensive and professional in your response."""
                if opus_guidance:
                    completion_prompt += f"\n\nSTRATEGIC GUIDANCE (from senior AI):\n{opus_guidance}\n\nUse this guidance to inform your response."
                
                if orchestrator == 'opus':
                    response = call_claude_opus(completion_prompt, conversation_history=conversation_context)
                else:
                    response = call_claude_sonnet(completion_prompt, conversation_history=conversation_context)
                
                if isinstance(response, dict):
                    if response.get('error'):
                        actual_output = f"Error completing task: {response.get('content', 'Unknown error')}"
                    else:
                        actual_output = response.get('content', '')
                else:
                    actual_output = str(response)
            
            print(f"Task completed. Output length: {len(actual_output) if actual_output else 0} chars")
            
            # Document creation
            document_created = False
            document_url = None
            document_type = None
            document_id = None
            
            should_create, doc_type = should_create_document(user_request)
            if should_create and actual_output and not actual_output.startswith('Error'):
                try:
                    from docx import Document
                    from docx.shared import Pt
                    doc = Document()
                    title = doc.add_heading('Shiftwork Solutions LLC', 0)
                    title.alignment = 1
                    lines = actual_output.split('\n')
                    for line in lines:
                        if line.strip():
                            if line.startswith('#'):
                                level = len(line) - len(line.lstrip('#'))
                                text = line.lstrip('#').strip()
                                doc.add_heading(text, level=min(level, 3))
                            else:
                                p = doc.add_paragraph(line)
                                p.style.font.size = Pt(11)
                    
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = f'shiftwork_document_{timestamp}.docx'
                    save_paths = [f'/mnt/user-data/outputs/{filename}', f'/tmp/{filename}', f'./{filename}']
                    
                    saved = False
                    saved_path = None
                    for filepath in save_paths:
                        try:
                            os.makedirs(os.path.dirname(filepath), exist_ok=True)
                            doc.save(filepath)
                            saved = True
                            saved_path = filepath
                            print(f"Document created: {filepath}")
                            break
                        except Exception as save_error:
                            print(f"Could not save to {filepath}: {save_error}")
                    
                    if saved and saved_path:
                        file_size = os.path.getsize(saved_path)
                        doc_title = generate_document_title(user_request, 'docx')
                        doc_category = categorize_document(user_request, 'docx')
                        document_id = save_generated_document(
                            filename=filename, original_name=doc_title, document_type='docx',
                            file_path=saved_path, file_size=file_size, task_id=task_id,
                            conversation_id=conversation_id, project_id=project_id,
                            title=doc_title, description=f"Generated from request: {user_request[:100]}",
                            category=doc_category
                        )
                        document_created = True
                        document_url = f'/api/generated-documents/{document_id}/download'
                        document_type = 'docx'
                except Exception as doc_error:
                    print(f"Document creation failed: {doc_error}")
            
            formatted_output = convert_markdown_to_html(actual_output)
            
            consensus_result = None
            if enable_consensus and actual_output and not actual_output.startswith('Error'):
                try:
                    consensus_result = validate_with_consensus(actual_output)
                except Exception as consensus_error:
                    print(f"Consensus validation failed: {consensus_error}")
            
            total_time = time.time() - overall_start
            db.execute('UPDATE tasks SET status = ?, assigned_orchestrator = ?, execution_time_seconds = ? WHERE id = ?',
                      ('completed', orchestrator, total_time, task_id))
            db.commit()
            db.close()
            
            add_message(conversation_id, 'assistant', actual_output, task_id,
                       {'orchestrator': orchestrator, 'knowledge_applied': knowledge_applied,
                        'document_created': document_created, 'document_type': document_type,
                        'document_id': document_id, 'execution_time': total_time})
            
            suggestions = []
            if proactive:
                try:
                    suggestions = proactive.post_process_result(task_id, user_request, actual_output if actual_output else '')
                except Exception as suggest_error:
                    print(f"Suggestion generation failed: {suggest_error}")
            
            return jsonify({
                'success': True, 'task_id': task_id, 'conversation_id': conversation_id,
                'result': formatted_output, 'orchestrator': orchestrator,
                'specialists_used': [s.get('specialist') for s in specialist_results] if specialist_results else [],
                'consensus': consensus_result, 'execution_time': total_time,
                'knowledge_applied': knowledge_applied, 'knowledge_used': knowledge_applied,
                'knowledge_sources': knowledge_sources, 'formatting_applied': True,
                'document_created': document_created, 'document_url': document_url,
                'document_id': document_id, 'document_type': document_type, 'suggestions': suggestions
            })
        except Exception as orchestration_error:
            import traceback
            print(f"Orchestration error: {traceback.format_exc()}")
            db.execute('UPDATE tasks SET status = ? WHERE id = ?', ('failed', task_id))
            db.commit()
            db.close()
            add_message(conversation_id, 'assistant', f"Error: {str(orchestration_error)}", task_id, {'error': True})
            return jsonify({'success': False, 'error': f'Orchestration failed: {str(orchestration_error)}',
                           'task_id': task_id, 'conversation_id': conversation_id}), 500
    except Exception as e:
        import traceback
        print(f"CRITICAL ERROR: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'}), 500


# ============================================================================
# OTHER EXISTING ENDPOINTS
# ============================================================================

@core_bp.route('/api/tasks', methods=['GET'])
def get_tasks():
    """Get recent tasks"""
    try:
        limit = request.args.get('limit', 20, type=int)
        db = get_db()
        tasks = db.execute('SELECT * FROM tasks ORDER BY created_at DESC LIMIT ?', (limit,)).fetchall()
        db.close()
        return jsonify({'tasks': [dict(task) for task in tasks]})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@core_bp.route('/api/task/<int:task_id>', methods=['GET'])
def get_task(task_id):
    """Get specific task details"""
    try:
        db = get_db()
        task = db.execute('SELECT * FROM tasks WHERE id = ?', (task_id,)).fetchone()
        db.close()
        if not task:
            return jsonify({'error': 'Task not found'}), 404
        return jsonify(dict(task))
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@core_bp.route('/api/stats', methods=['GET'])
def get_stats():
    """Get system statistics"""
    try:
        db = get_db()
        total_tasks = db.execute('SELECT COUNT(*) FROM tasks').fetchone()[0]
        completed_tasks = db.execute('SELECT COUNT(*) FROM tasks WHERE status = ?', ('completed',)).fetchone()[0]
        total_conversations = db.execute('SELECT COUNT(*) FROM conversations').fetchone()[0]
        total_messages = db.execute('SELECT COUNT(*) FROM conversation_messages').fetchone()[0]
        total_documents = db.execute('SELECT COUNT(*) FROM generated_documents WHERE is_deleted = 0').fetchone()[0]
        db.close()
        return jsonify({
            'total_tasks': total_tasks,
            'completed_tasks': completed_tasks,
            'success_rate': (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0,
            'total_conversations': total_conversations,
            'total_messages': total_messages,
            'total_documents': total_documents
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@core_bp.route('/api/documents', methods=['GET'])
def get_documents():
    """Get documents list - returns BOTH knowledge base docs AND generated documents."""
    try:
        import sys
        app_module = sys.modules.get('app')
        knowledge_base = getattr(app_module, 'knowledge_base', None) if app_module else None
        
        kb_documents = []
        if knowledge_base:
            for filename, doc_data in knowledge_base.knowledge_index.items():
                kb_documents.append({
                    'id': f'kb_{filename}',
                    'filename': filename,
                    'title': doc_data.get('title', filename),
                    'type': 'knowledge',
                    'word_count': doc_data.get('word_count', 0),
                    'category': doc_data.get('category', 'reference'),
                    'source': 'knowledge_base'
                })
        
        generated_docs = get_generated_documents(limit=50)
        gen_documents = []
        for doc in generated_docs:
            gen_documents.append({
                'id': doc['id'],
                'filename': doc['filename'],
                'title': doc['title'] or doc['original_name'],
                'type': doc['document_type'],
                'size': doc['file_size'],
                'category': doc['category'],
                'created_at': doc['created_at'],
                'download_url': f"/api/generated-documents/{doc['id']}/download",
                'print_url': f"/api/generated-documents/{doc['id']}/print",
                'source': 'generated'
            })
        
        return jsonify({
            'success': True,
            'knowledge_base_documents': kb_documents,
            'generated_documents': gen_documents,
            'knowledge_base_count': len(kb_documents),
            'generated_count': len(gen_documents),
            'status': 'available'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@core_bp.route('/api/feedback', methods=['POST'])
def submit_feedback():
    """Submit feedback on task results"""
    try:
        data = request.json
        task_id = data.get('task_id')
        overall_rating = data.get('overall_rating')
        quality_rating = data.get('quality_rating')
        accuracy_rating = data.get('accuracy_rating')
        usefulness_rating = data.get('usefulness_rating')
        improvement_categories = data.get('improvement_categories', [])
        user_comment = data.get('user_comment', '')
        output_used = data.get('output_used', False)
        
        if not task_id or not overall_rating:
            return jsonify({'error': 'task_id and overall_rating required'}), 400
        
        db = get_db()
        try:
            task = db.execute('SELECT * FROM tasks WHERE id = ?', (task_id,)).fetchone()
            if not task:
                return jsonify({'error': 'Task not found'}), 404
            
            consensus = db.execute('SELECT agreement_score FROM consensus_validations WHERE task_id = ?', (task_id,)).fetchone()
            consensus_score = consensus['agreement_score'] if consensus else None
            
            consensus_was_accurate = None
            if consensus_score is not None:
                avg_rating = (quality_rating + accuracy_rating + usefulness_rating) / 3
                if consensus_score >= 0.7 and avg_rating >= 3.5:
                    consensus_was_accurate = True
                elif consensus_score < 0.7 and avg_rating < 3.5:
                    consensus_was_accurate = True
                else:
                    consensus_was_accurate = False
            
            db.execute('''INSERT INTO user_feedback 
                (task_id, overall_rating, quality_rating, accuracy_rating, usefulness_rating,
                 consensus_was_accurate, improvement_categories, user_comment, output_used)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                (task_id, overall_rating, quality_rating, accuracy_rating, usefulness_rating,
                 consensus_was_accurate, json.dumps(improvement_categories), user_comment, output_used))
            
            orchestrator = task['assigned_orchestrator']
            task_type = 'general'
            avg_rating = (quality_rating + accuracy_rating + usefulness_rating) / 3
            pattern_key = f"{orchestrator}_{task_type}"
            
            existing_pattern = db.execute('SELECT * FROM learning_records WHERE pattern_type = ?', (pattern_key,)).fetchone()
            if existing_pattern:
                new_avg = (existing_pattern['success_rate'] * existing_pattern['times_applied'] + avg_rating / 5.0) / (existing_pattern['times_applied'] + 1)
                db.execute('UPDATE learning_records SET success_rate = ?, times_applied = times_applied + 1 WHERE pattern_type = ?',
                          (new_avg, pattern_key))
            else:
                db.execute('INSERT INTO learning_records (pattern_type, success_rate, times_applied, pattern_data) VALUES (?, ?, ?, ?)',
                          (pattern_key, avg_rating / 5.0, 1, json.dumps({'orchestrator': orchestrator, 'task_type': task_type,
                                                                        'last_rating': avg_rating, 'last_consensus': consensus_score})))
            
            db.commit()
            return jsonify({'success': True, 'message': 'Feedback recorded - system is learning!',
                           'consensus_was_accurate': consensus_was_accurate, 'learning_updated': True})
        except Exception as e:
            db.rollback()
            import traceback
            print(f"Feedback error: {traceback.format_exc()}")
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            db.close()
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/learning/stats', methods=['GET'])
def get_learning_stats():
    """Get learning system statistics"""
    try:
        db = get_db()
        total_feedback = db.execute('SELECT COUNT(*) as count FROM user_feedback').fetchone()
        avg_overall = db.execute('SELECT AVG(overall_rating) as avg FROM user_feedback').fetchone()
        consensus_accuracy = db.execute('''SELECT COUNT(*) as total,
            SUM(CASE WHEN consensus_was_accurate = 1 THEN 1 ELSE 0 END) as accurate
            FROM user_feedback WHERE consensus_was_accurate IS NOT NULL''').fetchone()
        
        consensus_accuracy_rate = None
        if consensus_accuracy and consensus_accuracy['total'] > 0:
            consensus_accuracy_rate = int((consensus_accuracy['accurate'] / consensus_accuracy['total']) * 100)
        
        db.close()
        return jsonify({
            'success': True,
            'total_feedback_count': total_feedback['count'] if total_feedback else 0,
            'average_overall_rating': round(avg_overall['avg'], 2) if avg_overall and avg_overall['avg'] else 0,
            'consensus_accuracy_rate': consensus_accuracy_rate
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@core_bp.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download generated files (legacy endpoint)"""
    try:
        possible_paths = [f'/mnt/user-data/outputs/{filename}', f'/tmp/{filename}', f'./{filename}']
        for file_path in possible_paths:
            if os.path.exists(file_path):
                return send_file(file_path, as_attachment=True, download_name=filename)
        return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ============================================================================
# PROJECT FILE MANAGEMENT ENDPOINTS
# ============================================================================

@core_bp.route('/api/projects/<project_id>/files/upload', methods=['POST'])
def upload_project_files(project_id):
    """
    Upload files to a specific project.
    Creates project folder if it doesn't exist.
    """
    try:
        # Import here to avoid circular dependency
        from database_file_management import save_project_file
        from werkzeug.utils import secure_filename
        
        # Check if project exists
        db = get_db()
        project = db.execute(
            'SELECT * FROM projects WHERE project_id = ?',
            (project_id,)
        ).fetchone()
        db.close()
        
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        
        # Check if files were uploaded
        if 'files' not in request.files:
            return jsonify({'success': False, 'error': 'No files provided'}), 400
        
        files = request.files.getlist('files')
        
        if not files or files[0].filename == '':
            return jsonify({'success': False, 'error': 'No files selected'}), 400
        
        # Create project directory
        project_dir = f'/tmp/projects/{project_id}'
        os.makedirs(project_dir, exist_ok=True)
        
        uploaded_files = []
        
        for file in files:
            if file and file.filename:
                # Secure the filename
                original_filename = file.filename
                filename = secure_filename(original_filename)
                
                # Add timestamp to avoid collisions
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                name, ext = os.path.splitext(filename)
                filename = f"{name}_{timestamp}{ext}"
                
                # Save file to disk
                file_path = os.path.join(project_dir, filename)
                file.save(file_path)
                
                # Get file size
                file_size = os.path.getsize(file_path)
                
                # Determine file type
                file_type = ext[1:].lower() if ext else 'unknown'
                
                # Get optional metadata from request
                conversation_id = request.form.get('conversation_id')
                task_id = request.form.get('task_id')
                description = request.form.get('description')
                
                # Save to database
                file_id = save_project_file(
                    project_id=project_id,
                    filename=filename,
                    original_filename=original_filename,
                    file_type=file_type,
                    file_path=file_path,
                    file_size=file_size,
                    uploaded_by='user',
                    is_generated=False,
                    task_id=task_id,
                    conversation_id=conversation_id,
                    description=description
                )
                
                uploaded_files.append({
                    'id': file_id,
                    'filename': filename,
                    'original_filename': original_filename,
                    'file_type': file_type,
                    'file_size': file_size,
                    'download_url': f'/api/projects/{project_id}/files/{file_id}/download'
                })
        
        return jsonify({
            'success': True,
            'files': uploaded_files,
            'count': len(uploaded_files),
            'message': f'Uploaded {len(uploaded_files)} file(s) to project'
        })
        
    except Exception as e:
        import traceback
        print(f"File upload error: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/projects/<project_id>/files', methods=['GET'])
def list_project_files(project_id):
    """
    List all files in a project.
    This powers the folder view in the UI.
    """
    try:
        from database_file_management import get_project_files, get_file_stats_by_project
        
        # Check if project exists
        db = get_db()
        project = db.execute(
            'SELECT * FROM projects WHERE project_id = ?',
            (project_id,)
        ).fetchone()
        db.close()
        
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        
        # Get all files for this project
        files = get_project_files(project_id)
        
        # Get file statistics
        stats = get_file_stats_by_project(project_id)
        
        # Format files for UI
        formatted_files = []
        for file in files:
            formatted_files.append({
                'id': file['id'],
                'filename': file['filename'],
                'original_filename': file['original_filename'],
                'file_type': file['file_type'],
                'file_size': file['file_size'],
                'is_generated': file['is_generated'],
                'uploaded_at': file['uploaded_at'],
                'category': file['category'],
                'description': file['description'],
                'download_url': f'/api/projects/{project_id}/files/{file["id"]}/download'
            })
        
        return jsonify({
            'success': True,
            'project_id': project_id,
            'client_name': project['client_name'],
            'files': formatted_files,
            'count': len(formatted_files),
            'stats': stats
        })
        
    except Exception as e:
        import traceback
        print(f"List files error: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/projects/folders', methods=['GET'])
def list_project_folders():
    """
    List all projects with their file counts.
    This creates the folder view in the documents tab.
    """
    try:
        from database_file_management import get_all_projects_with_files
        
        projects = get_all_projects_with_files()
        
        formatted_projects = []
        for project in projects:
            formatted_projects.append({
                'project_id': project['project_id'],
                'client_name': project['client_name'] or 'Unnamed Project',
                'industry': project['industry'],
                'project_phase': project['project_phase'],
                'file_count': project['file_count'],
                'folder_icon': '📁',
                'view_url': f'/api/projects/{project["project_id"]}/files'
            })
        
        return jsonify({
            'success': True,
            'projects': formatted_projects,
            'count': len(formatted_projects)
        })
        
    except Exception as e:
        import traceback
        print(f"List folders error: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/projects/<project_id>/files/<int:file_id>/download', methods=['GET'])
def download_project_file(project_id, file_id):
    """Download a file from a project"""
    try:
        from database_file_management import get_project_file_by_id
        
        # Get file from database
        file = get_project_file_by_id(file_id)
        
        if not file:
            return jsonify({'error': 'File not found'}), 404
        
        # Verify file belongs to this project
        if file['project_id'] != project_id:
            return jsonify({'error': 'File does not belong to this project'}), 403
        
        # Check if file exists on disk
        file_path = file['file_path']
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found on server'}), 404
        
        # Send file
        return send_file(
            file_path,
            as_attachment=True,
            download_name=file['original_filename']
        )
        
    except Exception as e:
        import traceback
        print(f"Download error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@core_bp.route('/api/projects/<project_id>/files/<int:file_id>', methods=['DELETE'])
def delete_project_file_endpoint(project_id, file_id):
    """Delete a file from a project"""
    try:
        from database_file_management import get_project_file_by_id, delete_project_file
        
        # Get file from database
        file = get_project_file_by_id(file_id)
        
        if not file:
            return jsonify({'success': False, 'error': 'File not found'}), 404
        
        # Verify file belongs to this project
        if file['project_id'] != project_id:
            return jsonify({'success': False, 'error': 'File does not belong to this project'}), 403
        
        # Delete file (soft delete by default)
        hard_delete = request.args.get('hard', 'false').lower() == 'true'
        delete_project_file(file_id, hard_delete=hard_delete)
        
        return jsonify({
            'success': True,
            'message': 'File deleted successfully'
        })
        
    except Exception as e:
        import traceback
        print(f"Delete error: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500


@core_bp.route('/api/projects/<project_id>/files/find', methods=['POST'])
def find_project_file(project_id):
    """
    Find a file by name in a project.
    This is used by the AI to locate files when user says "look at the budget file"
    """
    try:
        from database_file_management import get_project_file_by_name
        
        data = request.json or {}
        filename = data.get('filename', '').strip()
        
        if not filename:
            return jsonify({'success': False, 'error': 'Filename required'}), 400
        
        # Find file
        file = get_project_file_by_name(project_id, filename)
        
        if not file:
            return jsonify({
                'success': False,
                'found': False,
                'message': f'No file matching "{filename}" found in this project'
            })
        
        # Return file info and content path
        return jsonify({
            'success': True,
            'found': True,
            'file': {
                'id': file['id'],
                'filename': file['filename'],
                'original_filename': file['original_filename'],
                'file_type': file['file_type'],
                'file_size': file['file_size'],
                'file_path': file['file_path'],
                'uploaded_at': file['uploaded_at'],
                'description': file['description']
            }
        })
        
    except Exception as e:
        import traceback
        print(f"Find file error: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# HELPER: Save generated files to project
# ============================================================================

def save_generated_file_to_project(project_id, filename, file_path, file_type, 
                                   task_id=None, conversation_id=None, 
                                   description=None):
    """
    Helper function to save AI-generated files to a project.
    Call this when the AI generates a document for a project.
    """
    try:
        from database_file_management import save_project_file
        from werkzeug.utils import secure_filename
        
        if not os.path.exists(file_path):
            print(f"⚠️ File does not exist: {file_path}")
            return None
        
        # Create project directory
        project_dir = f'/tmp/projects/{project_id}'
        os.makedirs(project_dir, exist_ok=True)
        
        # Copy file to project directory
        secure_name = secure_filename(filename)
        new_path = os.path.join(project_dir, secure_name)
        
        import shutil
        shutil.copy2(file_path, new_path)
        
        # Get file size
        file_size = os.path.getsize(new_path)
        
        # Save to database
        file_id = save_project_file(
            project_id=project_id,
            filename=secure_name,
            original_filename=filename,
            file_type=file_type,
            file_path=new_path,
            file_size=file_size,
            uploaded_by='ai',
            is_generated=True,
            task_id=task_id,
            conversation_id=conversation_id,
            description=description
        )
        
        return file_id
        
    except Exception as e:
        print(f"Error saving generated file to project: {e}")
        return None


# I did no harm and this file is not truncated

# I did no harm and this file is not truncated
