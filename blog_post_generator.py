"""
BLOG POST GENERATOR - AI-Powered SEO & AI-Search Optimized Blog Posts
Shiftwork Solutions LLC  

PURPOSE:
    Generates conversational, SEO-optimized blog posts for Shiftwork Solutions
    website. Posts are written in an accessible, human style (someone reads
    these over coffee, not in a boardroom), while rigorously following the same
    SEO and AI-search optimization rules as the Case Study Generator.
    
    NOW INCLUDES: URL slug generation, meta descriptions, FAQ sections,
    numbered lists, and perfect 100/100 SEO optimization for AI search engines.

TARGET AUDIENCE:
    General Managers, Directors, HR Managers — people who own the workforce
    problem, control the budget, and feel pressure from above and below.

CHANGE LOG:
    February 23, 2026 - ENHANCED FOR PERFECT SEO (100/100 optimization)
                        * Added URL slug generation
                        * Added AI-generated meta descriptions (under 160 chars)
                        * Updated prompt to require FAQ sections
                        * Updated prompt to require numbered lists
                        * Added SEO metadata to database
                        * Complete website-ready SEO package

    February 23, 2026 - Initial creation. Blog post generation engine with
                        DOCX export, database persistence, and library view.

AUTHOR: Jim @ Shiftwork Solutions LLC
LAST UPDATED: February 23, 2026
"""

import os
import json
import re
from datetime import datetime


# ============================================================================
# TOPIC DEFINITIONS
# ============================================================================
BLOG_TOPICS = {
    'workforce_resistance': {
        'display': 'Workforce Resistance to Schedule Change',
        'primary_keyword': 'workforce resistance to schedule change',
        'secondary_keywords': [
            'employee pushback shift schedule', 'change management shift work',
            'getting employee buy-in scheduling', 'schedule change resistance'
        ],
        'prompt_hint': (
            'Why employees resist schedule changes, what makes it worse, '
            'and how to prevent it through early engagement and real involvement '
            'in the process rather than top-down announcement.'
        )
    },
    'employee_surveys': {
        'display': 'Employee Survey Best Practices',
        'primary_keyword': 'employee survey shift schedule',
        'secondary_keywords': [
            'workforce preference survey', 'shift schedule survey design',
            'employee input scheduling', 'workforce engagement survey'
        ],
        'prompt_hint': (
            'How to design, administer, and use employee surveys when changing '
            'shift schedules. Why most surveys fail (the question is wrong, or '
            'the results are filed away), and what good survey practice looks like.'
        )
    },
    'overtime_cost_reduction': {
        'display': 'Overtime Cost Reduction',
        'primary_keyword': 'overtime cost reduction manufacturing',
        'secondary_keywords': [
            'overtime management shift work', 'reducing overtime 24/7 operations',
            'overtime scheduling strategy', 'workforce overtime optimization'
        ],
        'prompt_hint': (
            'The real reasons overtime gets out of control in 24/7 operations, '
            'how schedule design contributes to chronic overtime, the 20/60/20 rule '
            'for overtime preference distribution, and practical strategies to '
            'reduce overtime without burning out the workforce.'
        )
    },
    'shift_schedule_design': {
        'display': 'Shift Schedule Design',
        'primary_keyword': 'shift schedule design 24/7 operations',
        'secondary_keywords': [
            '12-hour shift pattern', 'rotating shift schedule', 'fixed shift schedule',
            'shift rotation design', 'continuous operations scheduling'
        ],
        'prompt_hint': (
            'Principles of effective shift schedule design for continuous '
            'operations. Tradeoffs between fixed vs rotating schedules, 8-hour '
            'vs 12-hour shifts, and why the "best" schedule depends entirely '
            'on your workforce, not just your coverage requirements.'
        )
    },
    'twelve_vs_eight_hour': {
        'display': '12-Hour vs 8-Hour Shifts',
        'primary_keyword': '12-hour vs 8-hour shifts manufacturing',
        'secondary_keywords': [
            '12-hour shift advantages', '8-hour shift schedule', 'compressed shift schedule',
            'shift length comparison', '12-hour shift fatigue'
        ],
        'prompt_hint': (
            'A practical, balanced look at 12-hour vs 8-hour shift schedules. '
            'Why more than 80 percent of continuous operations workers prefer '
            '12-hour shifts when asked, the fatigue considerations that matter, '
            'and why the answer is different for every workforce and operation.'
        )
    },
    'work_life_balance': {
        'display': 'Work-Life Balance in Shift Work',
        'primary_keyword': 'work-life balance shift work',
        'secondary_keywords': [
            'shift worker work-life balance', 'employee schedule preferences',
            'shift work quality of life', 'workforce schedule satisfaction'
        ],
        'prompt_hint': (
            'What work-life balance actually means to shift workers — and why '
            'it means something different to every person on your floor. The '
            '400 employees, 400 definitions problem, how surveys reveal what '
            'workers actually value, and why assumptions about employee '
            'preferences are usually wrong.'
        )
    },
    'change_management': {
        'display': 'Change Management for Operations',
        'primary_keyword': 'change management shift schedule operations',
        'secondary_keywords': [
            'schedule transition management', 'workforce change management',
            'shift change communication', 'employee change resistance operations'
        ],
        'prompt_hint': (
            'Why schedule changes so often fail — even when the new schedule '
            'is objectively better. The importance of process, communication, '
            'and genuine employee involvement in creating lasting change. '
            'How top-down announcements trigger resistance even for good changes.'
        )
    },
    'turnover_reduction': {
        'display': 'Turnover Reduction in 24/7 Operations',
        'primary_keyword': 'turnover reduction 24/7 operations',
        'secondary_keywords': [
            'shift worker retention', 'reduce turnover manufacturing',
            'schedule-driven turnover', 'workforce retention shift work'
        ],
        'prompt_hint': (
            'How poorly designed shift schedules drive turnover — especially '
            'on night shifts. The hidden costs of turnover in continuous '
            'operations, which schedule-related factors employees cite when '
            'quitting, and how schedule redesign can meaningfully reduce '
            'avoidable turnover.'
        )
    },
    'fatigue_shift_work': {
        'display': 'Fatigue and Shift Work',
        'primary_keyword': 'fatigue shift work management',
        'secondary_keywords': [
            'shift worker fatigue', 'night shift fatigue management',
            'fatigue risk management operations', 'alertness shift work'
        ],
        'prompt_hint': (
            'The science and practical reality of fatigue in shift work. '
            'Why circadian rhythm matters, how shift start times affect sleep '
            'duration, why not starting shifts before 6am matters, and what '
            'managers can actually do to reduce fatigue-related risk without '
            'overhauling the entire schedule.'
        )
    },
    'choosing_consultant': {
        'display': 'How to Choose a Shift Scheduling Consultant',
        'primary_keyword': 'shift scheduling consultant',
        'secondary_keywords': [
            'shift work consultant', 'workforce scheduling consulting firm',
            'choose shiftwork consultant', 'shift schedule consulting services'
        ],
        'prompt_hint': (
            'What to look for when hiring a shift scheduling consultant. '
            'The difference between firms that impose solutions vs those that '
            'involve employees in the process. Why experience and methodology '
            'matter more than software, and the questions every buyer should '
            'ask before signing a contract.'
        )
    },
    'cost_of_wrong_schedule': {
        'display': 'The Cost of Getting Scheduling Wrong',
        'primary_keyword': 'cost of poor shift scheduling',
        'secondary_keywords': [
            'shift schedule ROI', 'scheduling impact operations',
            'cost bad shift schedule', 'workforce scheduling financial impact'
        ],
        'prompt_hint': (
            'The full financial and human cost of a poorly designed or poorly '
            'implemented shift schedule. Overtime explosion, turnover, absenteeism, '
            'quality problems, and workforce morale — and why these costs are '
            'usually invisible on the P&L until they become a crisis.'
        )
    },
    'other': {
        'display': 'Other',
        'primary_keyword': 'shift scheduling 24/7 operations',
        'secondary_keywords': [
            'workforce scheduling', 'continuous operations management',
            'shift work consulting', 'employee scheduling optimization'
        ],
        'prompt_hint': (
            'A general shiftwork operations topic related to scheduling, workforce '
            'management, change management, or consulting best practices.'
        )
    }
}


def generate_url_slug(title: str) -> str:
    """
    Generate a clean URL slug from the blog post title.
    Example: "12-Hour vs 8-Hour Shifts" -> "12-hour-vs-8-hour-shifts"
    """
    # Convert to lowercase
    slug = title.lower()
    
    # Remove special characters, keep alphanumeric, spaces, and hyphens
    slug = re.sub(r'[^\w\s-]', '', slug)
    
    # Replace spaces and multiple hyphens with single hyphen
    slug = re.sub(r'[\s_]+', '-', slug)
    slug = re.sub(r'-+', '-', slug)
    
    # Remove leading/trailing hyphens
    slug = slug.strip('-')
    
    # Limit length to 60 characters
    if len(slug) > 60:
        slug = slug[:60].rstrip('-')
    
    return slug


def get_blog_post_prompt(topic: str, custom_topic: str, angle: str) -> str:
    """
    Build the AI prompt for blog post generation.
    
    NOW INCLUDES: Requirements for FAQ sections, numbered lists,
    and perfect structure for 100/100 SEO optimization.
    """

    # Resolve topic data
    topic_data = BLOG_TOPICS.get(topic, BLOG_TOPICS['other'])
    topic_display = topic_data['display']
    primary_keyword = topic_data['primary_keyword']
    secondary_keywords = ', '.join(topic_data['secondary_keywords'][:3])
    prompt_hint = topic_data['prompt_hint']

    # If topic is 'other' and custom topic was provided, use it
    if topic == 'other' and custom_topic:
        topic_display = custom_topic

    # The specific angle or additional context Jim provided
    angle_section = f"\nSPECIFIC ANGLE OR FOCUS FOR THIS POST:\n{angle}\n" if angle else ""

    prompt = f"""You are a senior workforce management consultant and content writer for Shiftwork Solutions LLC. You're writing a blog post for the company website — conversational, expert, and deeply practical.

Your readers are General Managers, Plant Directors, and HR Managers at companies with 24/7 shift operations. They are dealing with real pressure from executives to control costs and from employees who are worried about schedule changes affecting their lives.

TOPIC: {topic_display}
PRIMARY KEYWORD: {primary_keyword}
SECONDARY KEYWORDS: {secondary_keywords}
TOPIC KNOWLEDGE: {prompt_hint}
{angle_section}

===== ABOUT SHIFTWORK SOLUTIONS LLC =====
- Specialized consulting firm with hundreds of completed 24/7 shift operations engagements
- Partners: Jim Dillingham (San Rafael, CA), Dan Capshaw (Ivins, UT), Ethan Franklin
- Core philosophy: the best shift schedule is the one employees actually choose
- Core process: employee surveys, data analysis, collaborative schedule design, change management
- Typical engagement: 6 weeks, about $16,500/week
- Unique value: we walk the shop floor, talk to workers directly, and earn their trust
- Industries: manufacturing, pharmaceutical, food processing, mining, logistics, and many more

===== THE EMPLOYEE ENGAGEMENT RULE =====
EVERY BLOG POST must connect back to the employee engagement dimension — because this is
the #1 reason companies hire us. Clients are not afraid of the schedule math. They are
afraid of their employees pushing back.

===== 100/100 SEO OPTIMIZATION REQUIREMENTS =====
This post MUST achieve perfect SEO and AI search optimization. Every element matters.

1. TITLE: Lead with the topic or a common question. Use "{primary_keyword}" or close variant.
   Keep under 70 characters.

2. OPENING PARAGRAPH: Use "{primary_keyword}" in the first 2 sentences. Hook immediately.

3. SECONDARY KEYWORDS: Weave these into headings AND body: {secondary_keywords}
   Use at least TWO secondary keywords in section headings.

4. NUMBERS: Use specific, conservative, believable statistics throughout.
   Examples: "roughly 80 percent", "15-20 percent chronic overtime", "turnover rates of 25-35 percent"

5. QUOTABLE INSIGHTS: Include at least 5 specific, memorable insights that AI search
   engines would pull out to answer questions. Make them concrete and data-backed.

6. FAQ SECTION: REQUIRED. Must include 3 common questions using primary/secondary keywords.

7. NUMBERED LISTS: REQUIRED. Must include at least one 5-item numbered list with
   bold headings for each item.

===== REQUIRED STRUCTURE =====
Use this EXACT structure (this is MANDATORY for 100/100 SEO score):

# [Title: keyword-rich, under 70 characters]

[Opening paragraph - hook + primary keyword in first 2 sentences]

## [First subheading - use a secondary keyword or natural question]

[2-3 paragraphs on the first main point]

## [Second subheading - NUMBERED LIST with secondary keyword if possible]
[Example: "5 Warning Signs [Secondary Keyword] Is Failing"]

[Brief intro paragraph explaining the list]

1. **[Bold item heading]** - [2-3 sentences with specific details and data]
2. **[Bold item heading]** - [2-3 sentences with specific details and data]
3. **[Bold item heading]** - [2-3 sentences with specific details and data]
4. **[Bold item heading]** - [2-3 sentences with specific details and data]
5. **[Bold item heading]** - [2-3 sentences with specific details and data]

## [Third subheading - include secondary keyword]

[2-3 paragraphs - this is where employee engagement dimension should appear]

## Frequently Asked Questions

**Q: [Question using primary keyword - what people actually search for]**

A: [2-3 sentence answer with specific data, conservative numbers]

**Q: [Question using secondary keyword]**

A: [2-3 sentence answer with specific, quotable insight]

**Q: [Third question - common concern in this topic area]**

A: [2-3 sentence answer with actionable information]

## What This Means for You

[1-2 paragraphs - practical takeaway. What should they do differently?]

## About Shiftwork Solutions LLC

[DYNAMIC CLOSING - unique to this post, not boilerplate.
Write 2-3 sentences that:
  1. Connect this specific topic to what Shiftwork Solutions has seen across hundreds of engagements
  2. Invite reader to reach out - conversational, not salesy
  3. Reference employee-centered philosophy

Every post must have a DIFFERENT closing.]

===== TONE AND STYLE =====
- Conversational but authoritative
- Use "you" to address reader directly
- Short paragraphs (2-4 sentences)
- Occasional first-person ("In our experience...")
- No jargon unless audience uses it
- Direct and honest
- Target: 800-1000 words

===== OUTPUT =====
Return ONLY the markdown content. Start with # title.
No preamble, no explanation, no JSON wrapper.
Just clean markdown."""

    return prompt


def generate_meta_description(title: str, content: str, primary_keyword: str) -> str:
    """
    Generate an SEO-optimized meta description using AI.
    Must be under 160 characters and include the primary keyword.
    """
    try:
        from config import ANTHROPIC_API_KEY, CLAUDE_SONNET_MODEL
        import anthropic

        if not ANTHROPIC_API_KEY:
            # Fallback: create a simple meta description
            return f"{title[:100]}. Expert insights on {primary_keyword}."

        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        prompt = f"""Generate a meta description for this blog post. Requirements:
- MAXIMUM 155 characters (hard limit)
- Must include the keyword: "{primary_keyword}"
- Compelling and action-oriented
- Mention a specific benefit or outcome

Blog Post Title: {title}

First paragraph: {content[:500]}

Return ONLY the meta description text, nothing else. No quotes, no explanation."""

        message = client.messages.create(
            model=CLAUDE_SONNET_MODEL,
            max_tokens=100,
            messages=[{"role": "user", "content": prompt}]
        )

        meta_desc = message.content[0].text.strip()
        
        # Remove quotes if present
        meta_desc = meta_desc.strip('"').strip("'")
        
        # Ensure it's under 160 characters
        if len(meta_desc) > 155:
            meta_desc = meta_desc[:152] + "..."
        
        return meta_desc

    except Exception as e:
        print(f"[BlogPost] Meta description generation failed: {e}")
        # Fallback
        return f"{title[:100]}. Expert insights on {primary_keyword}."


def generate_blog_post(topic: str, custom_topic: str, angle: str) -> dict:
    """
    Generate a blog post with complete SEO metadata.

    Returns:
        {
            'success': True/False,
            'id': int,
            'title': str,
            'url_slug': str,               # NEW
            'meta_description': str,       # NEW
            'topic': str,
            'topic_display': str,
            'content': str,
            'word_count': int
        }
    """
    try:
        from config import ANTHROPIC_API_KEY, CLAUDE_SONNET_MODEL
        import anthropic

        if not ANTHROPIC_API_KEY:
            return {'success': False, 'error': 'Anthropic API key not configured'}

        # Validate topic
        if topic not in BLOG_TOPICS:
            topic = 'other'

        # For 'other' topic, require custom_topic
        if topic == 'other' and not custom_topic:
            return {'success': False, 'error': 'Please describe the topic for "Other"'}

        prompt = get_blog_post_prompt(topic, custom_topic, angle)

        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        message = client.messages.create(
            model=CLAUDE_SONNET_MODEL,
            max_tokens=3000,  # Increased for longer, more detailed posts
            messages=[{"role": "user", "content": prompt}]
        )

        content = message.content[0].text.strip()

        # Extract title from the first # heading
        title = 'Blog Post'
        for line in content.split('\n'):
            line = line.strip()
            if line.startswith('# ') and not line.startswith('## '):
                title = line[2:].strip()
                break

        # Generate URL slug
        url_slug = generate_url_slug(title)

        # Get primary keyword for meta description
        topic_data = BLOG_TOPICS.get(topic, BLOG_TOPICS['other'])
        primary_keyword = topic_data['primary_keyword']

        # Generate meta description
        meta_description = generate_meta_description(title, content, primary_keyword)

        # Word count
        word_count = len(content.split())

        # Determine display name
        if topic == 'other' and custom_topic:
            topic_display = custom_topic
        else:
            topic_display = topic_data['display']

        # Save to database (now includes SEO metadata)
        post_id = save_blog_post_to_db(
            topic=topic,
            topic_display=topic_display,
            title=title,
            url_slug=url_slug,
            meta_description=meta_description,
            content=content,
            angle=angle or ''
        )

        print(f"[BlogPost] Generated: id={post_id}, slug={url_slug}, "
              f"title='{title[:40]}...', words={word_count}")

        return {
            'success': True,
            'id': post_id,
            'title': title,
            'url_slug': url_slug,
            'meta_description': meta_description,
            'topic': topic,
            'topic_display': topic_display,
            'content': content,
            'word_count': word_count
        }

    except Exception as e:
        import traceback
        return {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }


def generate_blog_post_docx(post_content: str, title: str, topic_display: str) -> bytes:
    """
    Convert markdown blog post content to a professional Word document.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Document styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(20)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1A, 0x53, 0x7A)

    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = 'Shiftwork Solutions LLC  |  Blog Post'
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.runs[0].font.size = Pt(9)
    header_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f'© Shiftwork Solutions LLC  |  {topic_display}  |  shift-work.com'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Helper: inline bold formatting
    def add_paragraph_with_formatting(doc, text, style_name='Normal'):
        para = doc.add_paragraph(style=style_name)
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for i, part in enumerate(parts):
            run = para.add_run(part)
            if i % 2 == 1:
                run.bold = True
        return para

    # Parse and render markdown
    lines = post_content.strip().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith('# ') and not line.startswith('## '):
            heading_text = line[2:].strip()
            para = doc.add_heading(heading_text, level=1)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph('')
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=3)
        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            add_paragraph_with_formatting(doc, bullet_text, 'List Bullet')
        elif re.match(r'^\d+\.\s', line):
            item_text = re.sub(r'^\d+\.\s', '', line).strip()
            add_paragraph_with_formatting(doc, item_text, 'List Number')
        elif line.startswith('---') or line.startswith('==='):
            pass
        elif line == '':
            pass
        else:
            if line.strip():
                add_paragraph_with_formatting(doc, line.strip())
        i += 1

    # Final metadata
    doc.add_paragraph('')
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_para.add_run(
        f'Generated by Shiftwork Solutions LLC AI System  |  '
        f'{datetime.now().strftime("%B %d, %Y")}'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.italic = True

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def save_blog_post_to_db(topic: str, topic_display: str, title: str,
                          url_slug: str, meta_description: str,
                          content: str, angle: str) -> int:
    """Save a generated blog post with SEO metadata to the database."""
    from database import get_db
    db = get_db()
    cursor = db.execute('''
        INSERT INTO blog_posts (
            topic, topic_display, title, url_slug, meta_description,
            content, angle, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        topic, topic_display, title, url_slug, meta_description,
        content, angle,
        datetime.now().isoformat(), datetime.now().isoformat()
    ))
    db.commit()
    record_id = cursor.lastrowid
    db.close()
    return record_id


def get_all_blog_posts() -> list:
    """Retrieve all saved blog posts, newest first."""
    from database import get_db
    db = get_db()
    rows = db.execute('''
        SELECT id, topic, topic_display, title, url_slug, 
               meta_description, angle, created_at
        FROM blog_posts
        ORDER BY created_at DESC
    ''').fetchall()
    db.close()
    return [dict(row) for row in rows]


def get_blog_post_by_id(post_id: int) -> dict:
    """Retrieve a single blog post by ID."""
    from database import get_db
    db = get_db()
    row = db.execute(
        'SELECT * FROM blog_posts WHERE id = ?', (post_id,)
    ).fetchone()
    db.close()
    return dict(row) if row else None


def delete_blog_post(post_id: int) -> bool:
    """Delete a blog post by ID."""
    from database import get_db
    db = get_db()
    db.execute('DELETE FROM blog_posts WHERE id = ?', (post_id,))
    db.commit()
    db.close()
    return True


def init_blog_posts_table():
    """Create the blog_posts table if it does not exist."""
    from database import get_db
    db = get_db()
    db.execute('''
        CREATE TABLE IF NOT EXISTS blog_posts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            topic TEXT NOT NULL,
            topic_display TEXT NOT NULL,
            title TEXT NOT NULL,
            url_slug TEXT NOT NULL,
            meta_description TEXT NOT NULL,
            content TEXT NOT NULL,
            angle TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    db.execute('''
        CREATE INDEX IF NOT EXISTS idx_blog_posts_topic
        ON blog_posts(topic)
    ''')
    db.execute('''
        CREATE INDEX IF NOT EXISTS idx_blog_posts_created
        ON blog_posts(created_at DESC)
    ''')
    db.execute('''
        CREATE INDEX IF NOT EXISTS idx_blog_posts_slug
        ON blog_posts(url_slug)
    ''')
    db.commit()
    db.close()
    print("  [BlogPost] blog_posts table ready with SEO fields")

# I did no harm and this file is not truncated
