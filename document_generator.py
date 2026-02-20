"""
Document Generator - AI Response to Downloadable Word Document
Created: February 20, 2026

PURPOSE:
    When the AI generates a checklist, report, proposal, guide, or other
    document-type response in Handler 10, this module converts that text
    into an actual downloadable .docx file using python-docx.

    This fixes the bug where the AI would describe or list document content
    in the chat window but tell users "I cannot provide downloadable files"
    because no file was actually being created.

CHANGELOG:
- February 20, 2026: Initial creation
  * Converts markdown AI responses to professional Word documents
  * Detects document-type requests (checklist, report, proposal, etc.)
  * Parses headings (H1/H2/H3), bullet lists, numbered lists, checkboxes,
    bold/italic text, and plain paragraphs
  * Saves to /tmp/outputs/ so existing download_handler.py can serve them
  * Returns document_url, document_id for Handler 10 JSON response
  * Integrates with existing save_generated_document() database function

Author: Jim @ Shiftwork Solutions LLC
"""

import os
import re
import uuid
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# DOCUMENT TYPE DETECTION
# ---------------------------------------------------------------------------

# Keywords that indicate the user wants a deliverable document file
DOCUMENT_REQUEST_KEYWORDS = [
    'checklist', 'check list',
    'report', 'reports',
    'proposal', 'proposals',
    'document', 'documents',
    'template', 'templates',
    'manual', 'manuals',
    'guide', 'guides',
    'summary', 'summaries',
    'plan', 'plans',
    'agenda', 'agendas',
    'outline', 'outlines',
    'memo', 'memos',
    'letter', 'letters',
    'word doc', 'word document',
    '.docx',
    'downloadable',
    'download',
    'save as',
    'export',
    'formatted document',
]

# Phrases that explicitly ask for a file to be generated
EXPLICIT_FILE_PHRASES = [
    'create a', 'create an',
    'generate a', 'generate an',
    'make a', 'make an',
    'write a', 'write an',
    'build a', 'build an',
    'produce a', 'produce an',
    'give me a', 'give me an',
    'provide a', 'provide an',
    'prepare a', 'prepare an',
    'draft a', 'draft an',
    'put together a', 'put together an',
    'put together the',
    'can you create', 'can you make', 'can you generate',
    'can you write', 'can you build',
    'i need a', 'i need an',
    'i want a', 'i want an',
    'i would like a', 'i would like an',
]


def is_document_request(user_request):
    """
    Detect whether the user's request is for a downloadable document.

    Returns True if the request contains both:
    1. An action phrase (create, generate, make, etc.)
    2. A document keyword (checklist, report, proposal, etc.)

    OR if it contains an explicit download/file keyword.

    Args:
        user_request (str): The user's original request text

    Returns:
        bool: True if this looks like a document creation request
    """
    if not user_request:
        return False

    req_lower = user_request.lower()

    # Explicit download requests always qualify
    for kw in ['downloadable', 'download', '.docx', 'word doc', 'word document',
               'save as', 'export', 'formatted document']:
        if kw in req_lower:
            return True

    # Require both an action phrase AND a document keyword
    has_action = any(phrase in req_lower for phrase in EXPLICIT_FILE_PHRASES)
    has_doc_keyword = any(kw in req_lower for kw in DOCUMENT_REQUEST_KEYWORDS)

    return has_action and has_doc_keyword


# ---------------------------------------------------------------------------
# DOCUMENT TITLE EXTRACTION
# ---------------------------------------------------------------------------

def extract_document_title(user_request, ai_response):
    """
    Extract a meaningful document title from the request or AI response.

    Looks for:
    1. H1 heading in AI response (# Title)
    2. Key noun phrase in user request
    3. Falls back to generic "Document" title

    Args:
        user_request (str): Original user request
        ai_response (str): AI-generated response text

    Returns:
        str: Document title
    """
    # Try to find H1 heading in AI response
    h1_match = re.search(r'^#\s+(.+)$', ai_response, re.MULTILINE)
    if h1_match:
        return h1_match.group(1).strip()

    # Try to extract from user request
    req_lower = user_request.lower()

    # Common patterns: "create a [X] for [client]"
    patterns = [
        r'(?:create|make|generate|write|build|produce|prepare|draft)\s+(?:a|an|the)\s+(.+?)(?:\s+for\s+|\s+about\s+|\s+on\s+|$)',
        r'(?:give me|provide|i need|i want)\s+(?:a|an)\s+(.+?)(?:\s+for\s+|\s+about\s+|\s+on\s+|$)',
    ]

    for pattern in patterns:
        match = re.search(pattern, req_lower)
        if match:
            raw_title = match.group(1).strip()
            # Capitalize first letter of each word
            return raw_title.title()

    # Timestamp-based fallback
    timestamp = datetime.now().strftime('%B %d, %Y')
    return f"Document - {timestamp}"


# ---------------------------------------------------------------------------
# MARKDOWN TO DOCX CONVERTER
# ---------------------------------------------------------------------------

def markdown_to_docx(title, markdown_text):
    """
    Convert markdown-formatted AI response text into a Word document.

    Handles:
    - H1 (#), H2 (##), H3 (###) headings
    - Bullet lists (- item, * item)
    - Numbered lists (1. item)
    - Checkbox lists (- [ ] item, - [x] item)
    - Bold text (**text**)
    - Italic text (*text*)
    - Horizontal rules (---, ***)
    - Plain paragraphs

    Args:
        title (str): Document title for the cover/heading
        markdown_text (str): Markdown-formatted text from AI response

    Returns:
        Document: python-docx Document object ready to save
    """
    doc = Document()

    # Set page margins (1.25" sides, 1" top/bottom - professional standard)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # --- Document title ---
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add company subtitle
    subtitle = doc.add_paragraph('Shiftwork Solutions LLC')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

    # Add date
    date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Spacer
    doc.add_paragraph('')

    # --- Parse and add content line by line ---
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.rstrip()

        # Skip blank lines (add spacing naturally)
        if not line.strip():
            continue

        # Skip lines that are just the same as the title (avoid duplication)
        if line.strip() == f'# {title}' or line.strip() == title:
            continue

        # H1
        if re.match(r'^#\s+', line):
            text = re.sub(r'^#\s+', '', line).strip()
            doc.add_heading(_strip_markdown_inline(text), level=1)

        # H2
        elif re.match(r'^##\s+', line):
            text = re.sub(r'^##\s+', '', line).strip()
            doc.add_heading(_strip_markdown_inline(text), level=2)

        # H3
        elif re.match(r'^###\s+', line):
            text = re.sub(r'^###\s+', '', line).strip()
            doc.add_heading(_strip_markdown_inline(text), level=3)

        # Horizontal rule
        elif re.match(r'^[-*_]{3,}$', line.strip()):
            # Add a thin paragraph with a bottom border as a visual separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        # Checkbox (- [ ] or - [x])
        elif re.match(r'^-\s+\[[ xX]\]\s+', line):
            checked = bool(re.match(r'^-\s+\[[xX]\]', line))
            text = re.sub(r'^-\s+\[[ xX]\]\s+', '', line).strip()
            text = _strip_markdown_inline(text)
            prefix = '[x]' if checked else '[ ]'
            p = doc.add_paragraph(style='List Bullet')
            _add_run_with_formatting(p, f'{prefix}  {text}')

        # Bullet list (- or *)
        elif re.match(r'^[-*]\s+', line):
            text = re.sub(r'^[-*]\s+', '', line).strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_runs(p, text)

        # Numbered list (1. 2. etc.)
        elif re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            _add_inline_runs(p, text)

        # Indented bullet (  - item)
        elif re.match(r'^\s{2,}[-*]\s+', line):
            text = re.sub(r'^\s+[-*]\s+', '', line).strip()
            p = doc.add_paragraph(style='List Bullet 2')
            _add_inline_runs(p, text)

        # Plain paragraph
        else:
            text = line.strip()
            if text:
                p = doc.add_paragraph()
                _add_inline_runs(p, text)

    return doc


def _strip_markdown_inline(text):
    """
    Remove markdown inline formatting (bold, italic, code) from text.
    Used where python-docx style handles formatting (headings).

    Args:
        text (str): Text with possible markdown formatting

    Returns:
        str: Plain text
    """
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    return text


def _add_run_with_formatting(paragraph, text):
    """
    Add a simple text run to a paragraph with no inline formatting.

    Args:
        paragraph: python-docx Paragraph
        text (str): Plain text
    """
    paragraph.add_run(text)


def _add_inline_runs(paragraph, text):
    """
    Parse inline markdown formatting and add appropriately styled runs.

    Handles:
    - **bold** -> bold run
    - *italic* -> italic run
    - `code` -> monospace run
    - plain text -> normal run

    Args:
        paragraph: python-docx Paragraph
        text (str): Text with possible inline markdown
    """
    # Pattern matches bold, italic, code, or plain text segments
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+)'
    segments = re.findall(pattern, text)

    if not segments:
        paragraph.add_run(text)
        return

    for segment in segments:
        if segment.startswith('**') and segment.endswith('**'):
            # Bold
            run = paragraph.add_run(segment[2:-2])
            run.bold = True
        elif segment.startswith('*') and segment.endswith('*') and len(segment) > 2:
            # Italic
            run = paragraph.add_run(segment[1:-1])
            run.italic = True
        elif segment.startswith('`') and segment.endswith('`'):
            # Code - monospace
            run = paragraph.add_run(segment[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(segment)


# ---------------------------------------------------------------------------
# MAIN GENERATION FUNCTION
# ---------------------------------------------------------------------------

def generate_document(user_request, ai_response_text, task_id=None,
                      conversation_id=None, project_id=None):
    """
    Generate a downloadable Word document from an AI response.

    This is the main entry point called from orchestration_handler.py
    Handler 10 when a document-type request is detected.

    Steps:
    1. Extract a meaningful title from the request/response
    2. Convert markdown AI response to python-docx Document
    3. Save to /tmp/outputs/ so download_handler.py can serve it
    4. Register in database via save_generated_document()
    5. Return document metadata for JSON response

    Args:
        user_request (str): Original user request text
        ai_response_text (str): AI-generated markdown response
        task_id (int, optional): Task ID for database record
        conversation_id (str, optional): Conversation ID for database record
        project_id (str, optional): Project ID for database record

    Returns:
        dict: {
            'success': bool,
            'document_url': str,      # e.g. '/api/download/doc_20260220_123456.docx'
            'document_id': int,       # database ID
            'document_type': 'docx',
            'filename': str,
            'error': str              # only on failure
        }
    """
    try:
        # 1. Extract title
        title = extract_document_title(user_request, ai_response_text)
        print(f"Document generator: title='{title}'")

        # 2. Convert markdown to Word document
        doc = markdown_to_docx(title, ai_response_text)

        # 3. Save to /tmp/outputs/
        os.makedirs('/tmp/outputs', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:40]
        filename = f"doc_{timestamp}_{safe_title}.docx"
        file_path = os.path.join('/tmp/outputs', filename)

        doc.save(file_path)

        file_size = os.path.getsize(file_path)
        print(f"Document generator: saved {filename} ({file_size} bytes)")

        # 4. Register in database
        doc_id = None
        try:
            from database import save_generated_document
            doc_id = save_generated_document(
                filename=filename,
                original_name=title,
                document_type='docx',
                file_path=file_path,
                file_size=file_size,
                task_id=task_id,
                conversation_id=conversation_id,
                project_id=project_id,
                title=title,
                description=f"Generated from request: {user_request[:200]}",
                category='document'
            )
            print(f"Document generator: registered as doc_id={doc_id}")
        except Exception as db_error:
            # Non-fatal - file exists even if DB save fails
            print(f"Document generator: DB save failed (non-critical): {db_error}")

        # 5. Build download URL
        # Use /api/download/ which is served by existing download_handler.py
        document_url = f'/api/download/{filename}'

        return {
            'success': True,
            'document_url': document_url,
            'document_id': doc_id,
            'document_type': 'docx',
            'filename': filename,
            'file_size': file_size,
            'title': title
        }

    except Exception as e:
        import traceback
        print(f"Document generator ERROR: {traceback.format_exc()}")
        return {
            'success': False,
            'error': str(e)
        }


# I did no harm and this file is not truncated
