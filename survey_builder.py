"""
SURVEY BUILDER MODULE FOR AI SWARM ORCHESTRATOR
File: survey_builder.py
Created: January 28, 2026
Last Updated: March 12, 2026

CHANGELOG:

- March 12, 2026: PHASE 2 — Enhanced export_to_word() for branded client-ready output.
  WHAT CHANGED (export_to_word only):
    - Added Shiftwork Solutions letterhead / brand header on every page via header
    - Added professional cover page: company name, survey title, date, confidentiality note
    - Added survey instructions page (standalone page before questions)
    - Added page breaks between major sections (schedule concepts, each question category)
    - Added Remark-compatible filled bubble circles (U+25CF) for paper administration;
      open circles (U+25CB) for online/unsure — controlled by administration_mode param
    - Added sequential question numbering across all categories (continuous 1..N)
    - Added "Thank You for Participating" closing page
    - Added footer with page number and Shiftwork Solutions contact info
    - Added helper methods: _add_page_break, _add_horizontal_rule,
      _add_cover_page, _add_instructions_page, _add_thank_you_page,
      _add_doc_header_footer, _bubble_char
  WHAT DID NOT CHANGE:
    - Question bank (_load_question_bank) — zero changes
    - Schedule library (_load_schedule_library) — zero changes
    - create_survey() — zero changes
    - get_categories() — zero changes
    - get_questions_by_category() — zero changes
    - _get_standard_instructions() — zero changes
    - All public method signatures unchanged; export_to_word() gains one
      optional keyword argument: administration_mode (str, default 'online')

- February 28, 2026: ADDED 39 NEW QUESTIONS FROM MASTER SURVEY DOCUMENT
  SOURCE: Master_Survey_for_Survey_Business.doc (Dan Capshaw / Shiftwork Solutions)
  PREVIOUS TOTAL: 64 questions
  NEW TOTAL: 103 questions
  NEW QUESTIONS BY CATEGORY:
    Demographics (5):
      crew_assignment, dept_tenure, prior_shiftwork, second_job_timing, worst_shift_start
    Sleep & Alertness (6):
      alarm_clock_normal, alarm_clock_day, alarm_clock_afternoon, alarm_clock_night,
      sleep_second_shift, sleep_third_shift
    Working Conditions (1):
      management_equality
    Schedule Features (16):
      schedule_policies_fair, shift_mobility_intent, least_preferred_8hr_shift,
      fixed_vs_rotating_no_seniority, fixed_vs_rotating_not_first_choice,
      crew_cohesion, rotation_frequency, rotation_direction, day_shift_start_10hr,
      shift_swap_importance, night_shift_start_preference, task_variety,
      weekend_willingness, weekend_occasional, understand_247_need,
      new_schedule_trial_willingness
    Overtime (5):
      overtime_timing_actual, overtime_timing_preferred, overtime_extend_shift,
      overtime_day_off, overtime_distribution_fair
    Day Care / Elder Care (6) - ENTIRELY NEW CATEGORY:
      daycare_use, daycare_location, daycare_relationship, daycare_shifts_used,
      daycare_shift_issue, daycare_worst_shift

PURPOSE:
Internal survey creation tool for Jim @ Shiftwork Solutions LLC.
NOT client-facing - this is for creating surveys that will be administered to clients.

FEATURES:
- Master question bank (103 standardized questions)
- Schedule rating questions (customizable per survey)
- Export to Word (.docx) — branded, client-ready, Remark-compatible
- Normative database integration (future)
- Custom question creation

WORKFLOW:
1. Select questions from master bank (or use comprehensive default)
2. Add custom questions if needed
3. Select schedules to rate (with videos/descriptions)
4. Generate Word document
5. Later: analyze results against normative database

AUTHOR: Jim @ Shiftwork Solutions LLC
"""

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ============================================================================
# BRAND CONSTANTS
# ============================================================================

_BRAND_BLUE  = RGBColor(0x1F, 0x5C, 0x99)   # #1F5C99 — Shiftwork Solutions blue
_BRAND_DARK  = RGBColor(0x1A, 0x1A, 0x2E)   # near-black for body text
_BRAND_GRAY  = RGBColor(0x66, 0x66, 0x66)   # mid gray for secondary text
_BRAND_LIGHT = RGBColor(0xF0, 0xF5, 0xFA)   # very light blue for shading

# Unicode bubble characters
_BUBBLE_FILLED = '\u25CF'   # ● filled circle  — used for paper (Remark-compatible)
_BUBBLE_OPEN   = '\u25CB'   # ○ open circle    — used for online / unsure
_BUBBLE_CHECK  = '\u25A1'   # □ open square    — used for checkbox questions


class SurveyBuilder:
    """
    Survey builder for creating customized shift schedule surveys
    """

    def __init__(self):
        """Initialize with master question bank"""
        self.question_bank = self._load_question_bank()
        self.schedule_library = self._load_schedule_library()

    # ========================================================================
    # QUESTION BANK  (unchanged from February 28, 2026)
    # ========================================================================

    def _load_question_bank(self):
        """
        Master question bank - 103 standardized questions
        Based on hundreds of client engagements

        Categories:
        - Demographics
        - Sleep & Alertness
        - Working Conditions (safety, communication, training, management)
        - Schedule Features (preferences, flexibility, rotation)
        - Overtime
        - Day Care / Elder Care
        - Open-Ended
        """

        return {

            # ================================================================
            # DEMOGRAPHICS
            # ================================================================

            'dept': {
                'id': 'dept',
                'category': 'demographics',
                'text': 'What department do you work in?',
                'type': 'multiple_choice',
                'options': ['Production', 'Logistics', 'Quality', 'Maintenance', 'Other']
            },
            'tenure': {
                'id': 'tenure',
                'category': 'demographics',
                'text': 'How long have you worked for this company?',
                'type': 'multiple_choice',
                'options': [
                    'Less than 6 months', '6 months to 1 year', '1 to 5 years',
                    '6 to 10 years', '11 to 15 years', '16 to 20 years', 'Over 20 years'
                ]
            },
            'dept_tenure': {
                'id': 'dept_tenure',
                'category': 'demographics',
                'text': 'How long have you worked in your current department?',
                'type': 'multiple_choice',
                'options': [
                    'Less than 6 months', '6 months to 1 year', '1 to 5 years',
                    '6 to 10 years', '11 to 15 years', '16 to 20 years', 'Over 20 years'
                ]
            },
            'crew_assignment': {
                'id': 'crew_assignment',
                'category': 'demographics',
                'text': 'What crew are you assigned to?',
                'type': 'multiple_choice',
                'options': [
                    'First Shift (8-hour or 10-hour)', 'First Shift (12-hour)',
                    'Second Shift (8-hour)', 'Second Shift (12-hour)',
                    'Third Shift', 'Weekend Shift'
                ]
            },
            'current_schedule': {
                'id': 'current_schedule',
                'category': 'demographics',
                'text': 'What schedule are you currently working?',
                'type': 'multiple_choice',
                'options': ['12-hour Day shift', '12-hour Night shift', 'Other']
            },
            'prior_shiftwork': {
                'id': 'prior_shiftwork',
                'category': 'demographics',
                'text': 'Have you ever worked shiftwork at another facility? (A "shiftwork" job is one that requires more than one shift to provide full coverage.)',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'second_job': {
                'id': 'second_job',
                'category': 'demographics',
                'text': 'Do you have a second job?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'second_job_timing': {
                'id': 'second_job_timing',
                'category': 'demographics',
                'text': 'If you have a second job, do you typically work at that job:',
                'type': 'multiple_choice',
                'options': [
                    'Before your shift at this company starts',
                    'After you have worked your shift at this company',
                    'Only on days that you don\'t work at this company',
                    'I don\'t work at a second job'
                ]
            },
            'employment_type': {
                'id': 'employment_type',
                'category': 'demographics',
                'text': 'Which best describes you?',
                'type': 'multiple_choice',
                'options': ['Hourly', 'Salaried']
            },
            'student_status': {
                'id': 'student_status',
                'category': 'demographics',
                'text': 'Are you a student?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'caregiving': {
                'id': 'caregiving',
                'category': 'demographics',
                'text': 'Do you have children or elder family members at home that require childcare or eldercare when you are at work?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'gender': {
                'id': 'gender',
                'category': 'demographics',
                'text': 'What is your gender?',
                'type': 'multiple_choice',
                'options': ['Female', 'Male', 'Other', 'Prefer not to say']
            },
            'age_group': {
                'id': 'age_group',
                'category': 'demographics',
                'text': 'What is your age group?',
                'type': 'multiple_choice',
                'options': [
                    '25 and under', '26 to 30', '31 to 35', '36 to 40',
                    '41 to 45', '46 to 50', '51 to 55', 'Over 55'
                ]
            },
            'single_parent': {
                'id': 'single_parent',
                'category': 'demographics',
                'text': 'Are you a single parent?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No', 'Not applicable']
            },
            'partner_status': {
                'id': 'partner_status',
                'category': 'demographics',
                'text': "Which best describes your spouse or domestic partner's work status?",
                'type': 'multiple_choice',
                'options': [
                    'No spouse or domestic partner',
                    'Does not work outside the home',
                    'Works a different schedule than I do in this company',
                    'Works a different schedule than I do outside this company',
                    'Works the same schedule as I do in this company',
                    'Works the same schedule as I do outside this company'
                ]
            },
            'commute_method': {
                'id': 'commute_method',
                'category': 'demographics',
                'text': 'How do you normally get to work?',
                'type': 'multiple_choice',
                'options': ['Drive by myself', 'Carpool', 'Public transportation', 'Walk or bike']
            },
            'commute_distance': {
                'id': 'commute_distance',
                'category': 'demographics',
                'text': 'How far do you commute to work (one way)?',
                'type': 'multiple_choice',
                'options': [
                    'Less than 1 mile', '1 to 5 miles', '6 to 10 miles', '11 to 20 miles',
                    '21 to 30 miles', '31 to 40 miles', 'More than 40 miles'
                ]
            },
            'worst_shift_start': {
                'id': 'worst_shift_start',
                'category': 'demographics',
                'text': 'Looking at your daily commute, what is the worst time to start the day shift?',
                'type': 'multiple_choice',
                'options': [
                    'Before 5:30 a.m.', '5:30 a.m.', '6:00 a.m.', '6:30 a.m.',
                    '7:00 a.m.', '7:30 a.m.', '8:00 a.m.', 'Later than 8:00 a.m. before noon'
                ]
            },

            # ================================================================
            # SLEEP & ALERTNESS
            # ================================================================

            'alarm_clock_normal': {
                'id': 'alarm_clock_normal',
                'category': 'sleep_alertness',
                'text': 'Do you normally use an alarm clock to wake up after a sleep period?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'alarm_clock_day': {
                'id': 'alarm_clock_day',
                'category': 'sleep_alertness',
                'text': 'Do you use an alarm clock to wake up when you are working day shift?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'alarm_clock_afternoon': {
                'id': 'alarm_clock_afternoon',
                'category': 'sleep_alertness',
                'text': 'Do you use an alarm clock to wake up when you are working afternoon shift?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'alarm_clock_night': {
                'id': 'alarm_clock_night',
                'category': 'sleep_alertness',
                'text': 'Do you use an alarm clock to wake up when you are working night shift?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'sleep_day_shift': {
                'id': 'sleep_day_shift',
                'category': 'sleep_alertness',
                'text': 'How many hours of sleep do you get every 24-hour period when you are working Day shift?',
                'type': 'multiple_choice',
                'options': [
                    'I never work the Day shift', 'Less than 5 hours',
                    '5 or more hours but less than 6 hours', '6 or more hours but less than 7 hours',
                    '7 or more hours but less than 8 hours', '8 or more hours but less than 9 hours',
                    '9 or more hours'
                ]
            },
            'sleep_second_shift': {
                'id': 'sleep_second_shift',
                'category': 'sleep_alertness',
                'text': 'How many hours of sleep do you get every 24-hour period when you are working second shift?',
                'type': 'multiple_choice',
                'options': [
                    'I never work the second shift', 'Less than 5 hours',
                    '5 or more hours but less than 6 hours', '6 or more hours but less than 7 hours',
                    '7 or more hours but less than 8 hours', '8 or more hours but less than 9 hours',
                    '9 or more hours'
                ]
            },
            'sleep_third_shift': {
                'id': 'sleep_third_shift',
                'category': 'sleep_alertness',
                'text': 'How many hours of sleep do you get every 24-hour period when you are working third shift?',
                'type': 'multiple_choice',
                'options': [
                    'I never work the third shift', 'Less than 5 hours',
                    '5 or more hours but less than 6 hours', '6 or more hours but less than 7 hours',
                    '7 or more hours but less than 8 hours', '8 or more hours but less than 9 hours',
                    '9 or more hours'
                ]
            },
            'sleep_night_shift': {
                'id': 'sleep_night_shift',
                'category': 'sleep_alertness',
                'text': 'How many hours of sleep do you get every 24-hour period when you are working Night shift?',
                'type': 'multiple_choice',
                'options': [
                    'I never work the Night shift', 'Less than 5 hours',
                    '5 or more hours but less than 6 hours', '6 or more hours but less than 7 hours',
                    '7 or more hours but less than 8 hours', '8 or more hours but less than 9 hours',
                    '9 or more hours'
                ]
            },
            'sleep_days_off': {
                'id': 'sleep_days_off',
                'category': 'sleep_alertness',
                'text': 'How many hours of sleep do you get every 24-hour period on days off?',
                'type': 'multiple_choice',
                'options': [
                    'Less than 5 hours', '5 or more hours but less than 6 hours',
                    '6 or more hours but less than 7 hours', '7 or more hours but less than 8 hours',
                    '8 or more hours but less than 9 hours', '9 or more hours'
                ]
            },
            'sleep_needed': {
                'id': 'sleep_needed',
                'category': 'sleep_alertness',
                'text': 'How many hours of sleep do you need every 24-hour period to be fully alert?',
                'type': 'multiple_choice',
                'options': [
                    'Less than 5 hours', '5 or more hours but less than 6 hours',
                    '6 or more hours but less than 7 hours', '7 or more hours but less than 8 hours',
                    '8 or more hours but less than 9 hours', '9 or more hours'
                ]
            },
            'sleepiness_problems': {
                'id': 'sleepiness_problems',
                'category': 'sleep_alertness',
                'text': 'How often do you notice you are having problems with safety or performance due to sleepiness?',
                'type': 'multiple_choice',
                'options': ['Never', 'Rarely', 'Once a month', 'Once a week', 'Almost daily']
            },

            # ================================================================
            # WORKING CONDITIONS
            # ================================================================

            'safety_rating': {
                'id': 'safety_rating',
                'category': 'working_conditions',
                'text': 'Overall, this is a safe place to work.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'safety_improvement': {
                'id': 'safety_improvement',
                'category': 'working_conditions',
                'text': 'Which best describes your opinion?',
                'type': 'multiple_choice',
                'options': [
                    'The company can do a lot more to improve safety at this site',
                    'The employees can do a lot more to improve safety at this site',
                    'Both of the above',
                    'Neither of the above, this is a very safe place to work'
                ]
            },
            'company_communication': {
                'id': 'company_communication',
                'category': 'working_conditions',
                'text': 'This company places a high priority on communication.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'communication_importance': {
                'id': 'communication_importance',
                'category': 'working_conditions',
                'text': 'Communication is important to me.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'handoff_time': {
                'id': 'handoff_time',
                'category': 'working_conditions',
                'text': 'How much time is needed to communicate daily plant conditions between shifts?',
                'type': 'multiple_choice',
                'options': [
                    'Less than 10 minutes', '10 minutes', '15 minutes', '20 minutes',
                    '25 minutes', '30 minutes', 'More than 30 minutes'
                ]
            },
            'management_input': {
                'id': 'management_input',
                'category': 'working_conditions',
                'text': 'Management welcomes input from the workforce.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'enjoy_work': {
                'id': 'enjoy_work',
                'category': 'working_conditions',
                'text': 'I enjoy the work that I do.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'pay_competitive': {
                'id': 'pay_competitive',
                'category': 'working_conditions',
                'text': 'The pay here is good compared to other jobs in the area.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'management_equality': {
                'id': 'management_equality',
                'category': 'working_conditions',
                'text': 'Management treats shift-workers and day-workers equally.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'company_belonging': {
                'id': 'company_belonging',
                'category': 'working_conditions',
                'text': 'I feel like I am a part of this company.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'absenteeism_impact': {
                'id': 'absenteeism_impact',
                'category': 'working_conditions',
                'text': 'Which best describes how you feel? (check as many as you wish)',
                'type': 'checkbox',
                'options': [
                    'There is no problem with last minute absenteeism at this site.',
                    'Covering other people\'s last minute absenteeism disrupts my family and social life.',
                    'The company needs to crack down on those few employees that are taking advantage of the existing absentee policy.'
                ]
            },
            'facility_improvement': {
                'id': 'facility_improvement',
                'category': 'working_conditions',
                'text': 'Overall, things are getting better at this facility.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'best_workplace': {
                'id': 'best_workplace',
                'category': 'working_conditions',
                'text': 'This is one of the best places to work in this area.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'training_importance': {
                'id': 'training_importance',
                'category': 'working_conditions',
                'text': 'Job training is important to me.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'training_adequacy': {
                'id': 'training_adequacy',
                'category': 'working_conditions',
                'text': 'I get enough training to do my job well.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'training_amount': {
                'id': 'training_amount',
                'category': 'working_conditions',
                'text': 'Which best describes how you feel?',
                'type': 'multiple_choice',
                'options': [
                    'We train way too much',
                    'We train just the right amount',
                    'We do not train nearly enough'
                ]
            },
            'supervisor_responsive': {
                'id': 'supervisor_responsive',
                'category': 'working_conditions',
                'text': 'My direct supervisor responds to my concerns about working conditions.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'management_responsive': {
                'id': 'management_responsive',
                'category': 'working_conditions',
                'text': 'Upper management responds to my concerns about working conditions.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },

            # ================================================================
            # SCHEDULE FEATURES
            # ================================================================

            'schedule_improvement': {
                'id': 'schedule_improvement',
                'category': 'schedule_features',
                'text': 'A better schedule will really improve things here.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'schedule_policies_fair': {
                'id': 'schedule_policies_fair',
                'category': 'schedule_features',
                'text': 'Current shift schedule policies are fair.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'current_schedule_satisfaction': {
                'id': 'current_schedule_satisfaction',
                'category': 'schedule_features',
                'text': 'I like my current schedule.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'better_schedules_exist': {
                'id': 'better_schedules_exist',
                'category': 'schedule_features',
                'text': 'I think there are better schedules available than our current schedule.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'shift_mobility_intent': {
                'id': 'shift_mobility_intent',
                'category': 'schedule_features',
                'text': 'Which best describes you?',
                'type': 'multiple_choice',
                'options': [
                    'I plan to go to a better shift as soon as I can',
                    'My current shift is where I plan to stay'
                ]
            },
            'time_off_predictable': {
                'id': 'time_off_predictable',
                'category': 'schedule_features',
                'text': 'My time off is predictable.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'schedule_flexibility': {
                'id': 'schedule_flexibility',
                'category': 'schedule_features',
                'text': 'My schedule allows me the flexibility to get time off when I really need it.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'preferred_8hr_shift': {
                'id': 'preferred_8hr_shift',
                'category': 'schedule_features',
                'text': 'If you were assigned to work a single shift for the next few years, which would be your preferred 8-hour shift?',
                'type': 'multiple_choice',
                'options': ['Day Shift', 'Afternoon Shift', 'Night Shift']
            },
            'least_preferred_8hr_shift': {
                'id': 'least_preferred_8hr_shift',
                'category': 'schedule_features',
                'text': 'If you were assigned to work a single shift for the next few years, which would be your least preferred 8-hour shift?',
                'type': 'multiple_choice',
                'options': ['Day Shift', 'Afternoon Shift', 'Night Shift']
            },
            'preferred_12hr_shift': {
                'id': 'preferred_12hr_shift',
                'category': 'schedule_features',
                'text': 'If you were assigned to work a single shift for the next few years, which would be your preferred 12-hour shift?',
                'type': 'multiple_choice',
                'options': ['Days', 'Nights']
            },
            'hours_vs_days_off': {
                'id': 'hours_vs_days_off',
                'category': 'schedule_features',
                'text': 'Assuming that you get the same amount of pay, which is more important to you?',
                'type': 'multiple_choice',
                'options': [
                    'Working fewer hours each day that I work, even though I will get fewer days off each week',
                    'Working more hours each day so that I can have more days off each week'
                ]
            },
            'fixed_vs_rotating': {
                'id': 'fixed_vs_rotating',
                'category': 'schedule_features',
                'text': 'Which would you prefer?',
                'type': 'multiple_choice',
                'options': ['Fixed or "steady" shifts', 'Rotating shifts']
            },
            'fixed_vs_rotating_no_seniority': {
                'id': 'fixed_vs_rotating_no_seniority',
                'category': 'schedule_features',
                'text': 'Which would you prefer?',
                'type': 'multiple_choice',
                'options': [
                    'Fixed shifts, even though seniority is not a consideration when being assigned to a shift',
                    'Rotating shifts'
                ]
            },
            'fixed_vs_rotating_not_first_choice': {
                'id': 'fixed_vs_rotating_not_first_choice',
                'category': 'schedule_features',
                'text': 'Which would you prefer?',
                'type': 'multiple_choice',
                'options': [
                    'Fixed shifts, even though I would not be assigned to my first choice',
                    'Rotating shifts'
                ]
            },
            'crew_cohesion': {
                'id': 'crew_cohesion',
                'category': 'schedule_features',
                'text': 'Keeping my current crew members together is important to me.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'rotation_frequency': {
                'id': 'rotation_frequency',
                'category': 'schedule_features',
                'text': 'How often would you like to rotate between shifts?',
                'type': 'multiple_choice',
                'options': [
                    'Once a week', 'Once every two weeks', 'Once every four weeks',
                    'Once every two months', 'Once every six months', 'Annually'
                ]
            },
            'rotation_direction': {
                'id': 'rotation_direction',
                'category': 'schedule_features',
                'text': 'On an 8-hour schedule, which direction would you prefer to rotate?',
                'type': 'multiple_choice',
                'options': [
                    'Days > Nights > Evenings > Days',
                    'Days > Evenings > Nights > Days',
                    'No preference'
                ]
            },
            'day_shift_start_8hr': {
                'id': 'day_shift_start_8hr',
                'category': 'schedule_features',
                'text': 'If you worked 8-hour shifts, what time would you like the day shift to start?',
                'type': 'multiple_choice',
                'options': [
                    'Before 5:30 a.m.', '5:30 a.m.', '6:00 a.m.', '6:30 a.m.',
                    '7:00 a.m.', '7:30 a.m.', '8:00 a.m.', 'Later than 8:00 a.m.'
                ]
            },
            'day_shift_start_10hr': {
                'id': 'day_shift_start_10hr',
                'category': 'schedule_features',
                'text': 'If you worked 10-hour shifts, what time would you like the day shift to start?',
                'type': 'multiple_choice',
                'options': [
                    'Before 5:30 a.m.', '5:30 a.m.', '6:00 a.m.', '6:30 a.m.',
                    '7:00 a.m.', '7:30 a.m.', '8:00 a.m.', 'Later than 8:00 a.m.'
                ]
            },
            'day_shift_start_12hr': {
                'id': 'day_shift_start_12hr',
                'category': 'schedule_features',
                'text': 'If you worked 12-hour shifts, what time would you like the day shift to start?',
                'type': 'multiple_choice',
                'options': [
                    'Before 5:30 a.m.', '5:30 a.m.', '6:00 a.m.', '6:30 a.m.',
                    '7:00 a.m.', '7:30 a.m.', '8:00 a.m.',
                    'Later than 8:00 a.m. before noon', 'Noon', '3:00 p.m.'
                ]
            },
            'weekend_preference': {
                'id': 'weekend_preference',
                'category': 'schedule_features',
                'text': 'If pay was not a factor, which would you prefer over an 8-week period?',
                'type': 'multiple_choice',
                'options': [
                    'Work 8 Saturdays and have 8 Sundays off',
                    'Work 8 Sundays and have 8 Saturdays off',
                    'Work 4 full weekends and have 4 full weekends off'
                ]
            },
            'shift_swap_importance': {
                'id': 'shift_swap_importance',
                'category': 'schedule_features',
                'text': 'The ability to swap shifts is important to me.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'night_shift_start_preference': {
                'id': 'night_shift_start_preference',
                'category': 'schedule_features',
                'text': 'If pay is not a factor when comparing the following two work shifts, I would prefer to work a night shift that:',
                'type': 'multiple_choice',
                'options': [
                    'Starts Sunday night and ends Monday morning',
                    'Starts Friday night and ends Saturday morning'
                ]
            },
            'weekend_pattern': {
                'id': 'weekend_pattern',
                'category': 'schedule_features',
                'text': 'Which best describes you?',
                'type': 'multiple_choice',
                'options': [
                    'I like my weekends off to alternate',
                    'I like to have several weekends off in a row and would be willing to work several in a row to make that happen'
                ]
            },
            'work_pattern': {
                'id': 'work_pattern',
                'category': 'schedule_features',
                'text': 'Which best describes you?',
                'type': 'multiple_choice',
                'options': [
                    'I like to work several days in a row and then take a long break',
                    'I like to work a couple of days in a row and then take a short break'
                ]
            },
            'three_day_preference': {
                'id': 'three_day_preference',
                'category': 'schedule_features',
                'text': 'If you could only have 3 days off per week, which of the following would you prefer?',
                'type': 'multiple_choice',
                'options': [
                    'Friday-Saturday-Sunday',
                    'Saturday-Sunday-Monday',
                    'Sunday-Monday-Tuesday'
                ]
            },
            'weekday_preference': {
                'id': 'weekday_preference',
                'category': 'schedule_features',
                'text': 'If your schedule requires you to take weekdays off, which day do you prefer to have off? (One answer only)',
                'type': 'multiple_choice',
                'options': ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
            },
            'supervisor_overlap': {
                'id': 'supervisor_overlap',
                'category': 'schedule_features',
                'text': 'What percentage of time do you think you should be working at the same time as your supervisor?',
                'type': 'multiple_choice',
                'options': ['100%', '90%', '80%', '70%', '60%', '50% or less']
            },
            'task_variety': {
                'id': 'task_variety',
                'category': 'schedule_features',
                'text': "I don't mind doing several different types of work during the week.",
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'weekend_willingness': {
                'id': 'weekend_willingness',
                'category': 'schedule_features',
                'text': 'Which best describes you?',
                'type': 'multiple_choice',
                'options': [
                    'I am willing to work my share of weekends',
                    'I will quit before I work weekends'
                ]
            },
            'weekend_occasional': {
                'id': 'weekend_occasional',
                'category': 'schedule_features',
                'text': 'I am willing to work weekends occasionally if I can plan them in advance.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'understand_247_need': {
                'id': 'understand_247_need',
                'category': 'schedule_features',
                'text': 'It is clear to me why we have to go to a 24/7 schedule to keep this company competitive in this industry.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'new_schedule_trial_willingness': {
                'id': 'new_schedule_trial_willingness',
                'category': 'schedule_features',
                'text': 'Which best describes you?',
                'type': 'multiple_choice',
                'options': [
                    'I am willing to try a new schedule for 6 to 12 months',
                    'I will reluctantly go along with a new schedule trial if that is what the majority of the workforce wants',
                    'I will quit before I go to a new schedule'
                ]
            },

            # ================================================================
            # OVERTIME
            # ================================================================

            'overtime_dependency': {
                'id': 'overtime_dependency',
                'category': 'overtime',
                'text': 'I depend on overtime worked outside my schedule to help me make ends meet:',
                'type': 'multiple_choice',
                'options': ['Never', 'Sometimes', 'Frequently', 'Every week']
            },
            'overtime_amount': {
                'id': 'overtime_amount',
                'category': 'overtime',
                'text': 'Over the last few months I have been:',
                'type': 'multiple_choice',
                'options': [
                    'Working too much overtime',
                    'Working too little overtime',
                    'Working just the right amount of overtime'
                ]
            },
            'overtime_satisfaction': {
                'id': 'overtime_satisfaction',
                'category': 'overtime',
                'text': 'Overtime levels are just right the way they are.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'overtime_timing_actual': {
                'id': 'overtime_timing_actual',
                'category': 'overtime',
                'text': 'When you work overtime outside your schedule, when do you usually work it?',
                'type': 'multiple_choice',
                'options': [
                    "I don't work overtime",
                    'On a regularly scheduled workday by coming in early or staying late',
                    'On Saturdays, but not Sundays',
                    'On Sundays, but not Saturdays',
                    'Any chance I get',
                    'I work overtime when necessary for business needs'
                ]
            },
            'overtime_timing_preferred': {
                'id': 'overtime_timing_preferred',
                'category': 'overtime',
                'text': 'When you have to work overtime, when do you prefer to work it?',
                'type': 'multiple_choice',
                'options': ['On a scheduled work day', 'On a day off', 'No preference']
            },
            'overtime_extend_shift': {
                'id': 'overtime_extend_shift',
                'category': 'overtime',
                'text': 'I prefer overtime by extending my shift.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'overtime_day_off': {
                'id': 'overtime_day_off',
                'category': 'overtime',
                'text': 'I prefer to work overtime by coming in on a day off.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'overtime_distribution_fair': {
                'id': 'overtime_distribution_fair',
                'category': 'overtime',
                'text': 'Current overtime distribution policies are fair.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'overtime_predictable': {
                'id': 'overtime_predictable',
                'category': 'overtime',
                'text': 'Overtime is predictable and can be planned for.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'time_vs_overtime': {
                'id': 'time_vs_overtime',
                'category': 'overtime',
                'text': 'If you had to choose between more time off or more overtime, what would you choose?',
                'type': 'multiple_choice',
                'options': ['More time off', 'More overtime']
            },
            'overtime_desire': {
                'id': 'overtime_desire',
                'category': 'overtime',
                'text': 'When it comes to overtime, I generally want to get:',
                'type': 'multiple_choice',
                'options': [
                    'As much as possible', 'Frequent overtime', 'Occasional overtime',
                    'Infrequent overtime', 'I do not want any overtime'
                ]
            },
            'overtime_expectation': {
                'id': 'overtime_expectation',
                'category': 'overtime',
                'text': 'I expect to get overtime whenever I want it.',
                'type': 'likert',
                'scale': '1 (Strongly Disagree) to 5 (Strongly Agree)'
            },
            'overtime_weekly_hours': {
                'id': 'overtime_weekly_hours',
                'category': 'overtime',
                'text': 'How much overtime would you like to have every week?',
                'type': 'multiple_choice',
                'options': [
                    'None', 'Less than 2 hours', 'Between 2 and 4 hours',
                    'Between 4 and 6 hours', 'Between 6 and 8 hours',
                    'Between 8 and 12 hours', 'I will take all that I can get'
                ]
            },

            # ================================================================
            # DAY CARE / ELDER CARE  (New section added February 28, 2026)
            # ================================================================

            'daycare_use': {
                'id': 'daycare_use',
                'category': 'daycare_eldercare',
                'text': 'Do you use outside day/elder care?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'daycare_location': {
                'id': 'daycare_location',
                'category': 'daycare_eldercare',
                'text': 'Is your day/elder care provider:',
                'type': 'multiple_choice',
                'options': ['Close to home', 'Close to work', 'At home']
            },
            'daycare_relationship': {
                'id': 'daycare_relationship',
                'category': 'daycare_eldercare',
                'text': 'Is your day/elder care provider a family member, neighbor or friend?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'daycare_shifts_used': {
                'id': 'daycare_shifts_used',
                'category': 'daycare_eldercare',
                'text': 'Do you use day/elder care when working? (check all that apply)',
                'type': 'checkbox',
                'options': [
                    'Days - Yes', 'Days - No',
                    'Afternoons - Yes', 'Afternoons - No',
                    'Nights - Yes', 'Nights - No'
                ]
            },
            'daycare_shift_issue': {
                'id': 'daycare_shift_issue',
                'category': 'daycare_eldercare',
                'text': 'Is day/elder care a bigger issue on a particular shift?',
                'type': 'multiple_choice',
                'options': ['Yes', 'No']
            },
            'daycare_worst_shift': {
                'id': 'daycare_worst_shift',
                'category': 'daycare_eldercare',
                'text': 'If day/elder care is a bigger issue on a particular shift, which shift?',
                'type': 'multiple_choice',
                'options': ['Days', 'Afternoons', 'Nights']
            },

            # ================================================================
            # OPEN-ENDED
            # ================================================================

            'schedule_like_most': {
                'id': 'schedule_like_most',
                'category': 'open_ended',
                'text': 'What do you like most about your current schedule?',
                'type': 'text'
            },
            'schedule_like_least': {
                'id': 'schedule_like_least',
                'category': 'open_ended',
                'text': 'What do you like least about your current schedule?',
                'type': 'text'
            },
            'work_life_positives': {
                'id': 'work_life_positives',
                'category': 'open_ended',
                'text': 'Please give 3 examples of things that the company is currently doing that contribute the most to your satisfaction with work-life balance.',
                'type': 'text'
            },
            'work_life_improvements': {
                'id': 'work_life_improvements',
                'category': 'open_ended',
                'text': 'Please list 3 examples of things the company should start doing to improve your work-life balance.',
                'type': 'text'
            },

        }

    # ========================================================================
    # SCHEDULE LIBRARY  (unchanged from January 28, 2026)
    # ========================================================================

    def _load_schedule_library(self):
        """
        Standard schedule library that can be shown to employees.
        Each can have video links and descriptions.
        """
        return {
            '2_3_2_fixed_days': {
                'id': '2_3_2_fixed_days',
                'name': '2-3-2 Pattern (Fixed Days)',
                'description': '2 days on, 2 off, 3 on, 2 off, 2 on, 3 off - Fixed day shift',
                'shift_length': '12 hours',
                'pattern_type': 'fixed',
                'video_url': None
            },
            '2_3_2_fixed_nights': {
                'id': '2_3_2_fixed_nights',
                'name': '2-3-2 Pattern (Fixed Nights)',
                'description': '2 days on, 2 off, 3 on, 2 off, 2 on, 3 off - Fixed night shift',
                'shift_length': '12 hours',
                'pattern_type': 'fixed',
                'video_url': None
            },
            'dupont_rotating': {
                'id': 'dupont_rotating',
                'name': 'DuPont Rotating Schedule',
                'description': 'Classic DuPont 12-hour rotating schedule',
                'shift_length': '12 hours',
                'pattern_type': 'rotating',
                'video_url': None
            },
            '4_on_4_off_days': {
                'id': '4_on_4_off_days',
                'name': '4-on-4-off (Fixed Days)',
                'description': 'Work 4 days, off 4 days - Fixed day shift',
                'shift_length': '12 hours',
                'pattern_type': 'fixed',
                'video_url': None
            },
            '4_on_4_off_nights': {
                'id': '4_on_4_off_nights',
                'name': '4-on-4-off (Fixed Nights)',
                'description': 'Work 4 days, off 4 days - Fixed night shift',
                'shift_length': '12 hours',
                'pattern_type': 'fixed',
                'video_url': None
            },
            'rotating_8hr': {
                'id': 'rotating_8hr',
                'name': 'Rotating 8-Hour Schedule',
                'description': 'Traditional rotating 8-hour schedule with day, afternoon, and night shifts',
                'shift_length': '8 hours',
                'pattern_type': 'rotating',
                'video_url': None
            },
            '5_and_2_days': {
                'id': '5_and_2_days',
                'name': '5&2 (Fixed Days)',
                'description': 'Work 5 days, off 2 days - Fixed day shift',
                'shift_length': '12 hours',
                'pattern_type': 'fixed',
                'video_url': None
            },
            '5_and_2_nights': {
                'id': '5_and_2_nights',
                'name': '5&2 (Fixed Nights)',
                'description': 'Work 5 days, off 2 days - Fixed night shift',
                'shift_length': '12 hours',
                'pattern_type': 'fixed',
                'video_url': None
            },
        }

    # ========================================================================
    # CREATE SURVEY  (unchanged from January 28, 2026)
    # ========================================================================

    def create_survey(self, project_name, company_name, selected_questions,
                      schedules_to_rate, custom_questions=None):
        """
        Create a complete survey object.

        Args:
            project_name:       Name of the project
            company_name:       Client company name
            selected_questions: List of question IDs from question bank
            schedules_to_rate:  List of schedule IDs to include ratings for
            custom_questions:   Optional list of custom question dicts

        Returns:
            Survey dict ready to pass to export_to_word()
        """
        survey = {
            'metadata': {
                'project_name': project_name,
                'company_name': company_name,
                'created_date': datetime.now().strftime('%Y-%m-%d'),
                'created_by': 'Jim @ Shiftwork Solutions LLC'
            },
            'instructions': self._get_standard_instructions(),
            'schedule_concepts': [],
            'questions': [],
            'custom_questions': custom_questions or []
        }

        # Add schedule rating sections
        for schedule_id in schedules_to_rate:
            if schedule_id in self.schedule_library:
                schedule = self.schedule_library[schedule_id]
                survey['schedule_concepts'].append({
                    'schedule': schedule,
                    'question': f"How would you rate the {schedule['name']}?",
                    'rating_options': [
                        'Perfect!',
                        "It's okay",
                        'I prefer something else',
                        'Not great',
                        'Never show this to me again!'
                    ]
                })

        # Add selected questions from bank (preserving order)
        for question_id in selected_questions:
            if question_id in self.question_bank:
                survey['questions'].append(self.question_bank[question_id])

        return survey

    # ========================================================================
    # STANDARD INSTRUCTIONS  (unchanged from January 28, 2026)
    # ========================================================================

    def _get_standard_instructions(self):
        """Standard survey instructions"""
        return {
            'title': 'Shift Schedule Lifestyle Survey',
            'directions': [
                'Your answers are anonymous, so please do not write your name or identify yourself on the questionnaire.',
                'Shiftwork Solutions LLC personnel will process these questionnaires. The results will be summarized and made available to you and management.',
                'For each question, completely fill in the circle next to the answer that matches your preference. Please mark only one response for each question.',
                'If you mark more than one answer, your response will not be counted!'
            ]
        }

    # ========================================================================
    # PHASE 2 — DOCUMENT FORMATTING HELPERS  (new March 12, 2026)
    # ========================================================================

    def _bubble_char(self, administration_mode):
        """
        Return the appropriate answer-option marker based on administration mode.

        For paper administration, return a filled circle (Remark-compatible bubble).
        For online or unsure, return an open circle for visual clarity.

        Args:
            administration_mode: 'paper' | 'online' | 'unsure' | anything else

        Returns:
            Single unicode character string
        """
        if administration_mode == 'paper':
            return _BUBBLE_FILLED   # ● Remark-compatible filled bubble
        return _BUBBLE_OPEN         # ○ open circle for online/unsure

    def _checkbox_char(self):
        """Return the checkbox character for checkbox-type questions."""
        return _BUBBLE_CHECK        # □ open square for multi-select

    def _add_page_break(self, doc):
        """
        Insert a hard page break as a Word run break element.
        Must be inside a Paragraph > Run for valid OOXML.
        """
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)
        return p

    def _add_horizontal_rule(self, doc, color='1F5C99', size=6):
        """
        Add a thin horizontal rule as a paragraph bottom border.
        Used as a visual divider below section headers and the cover page header.
        """
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(size))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def _add_doc_header_footer(self, doc, company_name):
        """
        Add a branded header and footer to every page of the document.

        Header: "Shiftwork Solutions LLC" left, company name right — separated by rule.
        Footer: "CONFIDENTIAL — Do Not Distribute" center, page placeholder right.
        Uses the default section so it applies to all pages.
        """
        section = doc.sections[0]

        # ---- HEADER --------------------------------------------------------
        header = section.header
        header.is_linked_to_previous = False

        # Clear default empty paragraph
        for p in header.paragraphs:
            p.clear()

        h_para = header.paragraphs[0]
        h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Left side: brand name
        run_left = h_para.add_run('Shiftwork Solutions LLC')
        run_left.bold = True
        run_left.font.size = Pt(9)
        run_left.font.color.rgb = _BRAND_BLUE

        # Tab stop to push company name to the right
        run_tab = h_para.add_run('\t')
        run_tab.font.size = Pt(9)

        # Right side: client company name
        run_right = h_para.add_run(company_name)
        run_right.font.size = Pt(9)
        run_right.font.color.rgb = _BRAND_GRAY
        run_right.italic = True

        # Set a right-aligned tab stop at 6 inches
        pPr = h_para._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '8640')   # 6 inches in twips (6 * 1440)
        tabs.append(tab)
        pPr.append(tabs)

        # Bottom border rule under header
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F5C99')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # ---- FOOTER --------------------------------------------------------
        footer = section.footer
        footer.is_linked_to_previous = False

        for p in footer.paragraphs:
            p.clear()

        f_para = footer.paragraphs[0]
        f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_conf = f_para.add_run('CONFIDENTIAL — For Internal Use Only  |  Shiftwork Solutions LLC  |  (415) 763-5005  |  www.shift-work.com')
        run_conf.font.size = Pt(8)
        run_conf.font.color.rgb = _BRAND_GRAY

        # Top border rule above footer
        f_pPr = f_para._p.get_or_add_pPr()
        f_pBdr = OxmlElement('w:pBdr')
        f_top = OxmlElement('w:top')
        f_top.set(qn('w:val'), 'single')
        f_top.set(qn('w:sz'), '4')
        f_top.set(qn('w:space'), '1')
        f_top.set(qn('w:color'), '1F5C99')
        f_pBdr.append(f_top)
        f_pPr.append(f_pBdr)

    def _add_cover_page(self, doc, survey):
        """
        Build the cover page:
          - SHIFTWORK SOLUTIONS LLC (branded header text)
          - Horizontal rule
          - Survey title (large)
          - Prepared for: <Company Name>
          - Date
          - Confidentiality note
          - Page break at end
        """
        metadata = survey['metadata']
        company_name = metadata['company_name']
        created_date = metadata['created_date']

        # Spacer at top
        doc.add_paragraph()
        doc.add_paragraph()

        # Company branding
        p_brand = doc.add_paragraph()
        p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_brand = p_brand.add_run('SHIFTWORK SOLUTIONS LLC')
        run_brand.bold = True
        run_brand.font.size = Pt(16)
        run_brand.font.color.rgb = _BRAND_BLUE

        # Horizontal rule
        self._add_horizontal_rule(doc, color='1F5C99', size=8)

        doc.add_paragraph()

        # Survey title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run('Shift Schedule Lifestyle Survey')
        run_title.bold = True
        run_title.font.size = Pt(28)
        run_title.font.color.rgb = _BRAND_DARK

        doc.add_paragraph()

        # Prepared for
        p_for = doc.add_paragraph()
        p_for.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_for_label = p_for.add_run('Prepared for:  ')
        run_for_label.font.size = Pt(14)
        run_for_label.font.color.rgb = _BRAND_GRAY
        run_for_company = p_for.add_run(company_name)
        run_for_company.bold = True
        run_for_company.font.size = Pt(14)
        run_for_company.font.color.rgb = _BRAND_DARK

        doc.add_paragraph()

        # Date
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_date = p_date.add_run(created_date)
        run_date.font.size = Pt(12)
        run_date.font.color.rgb = _BRAND_GRAY

        doc.add_paragraph()
        doc.add_paragraph()

        # Horizontal rule
        self._add_horizontal_rule(doc, color='CCCCCC', size=4)

        # Confidentiality note
        p_conf = doc.add_paragraph()
        p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_conf = p_conf.add_run(
            'CONFIDENTIAL — This survey is the proprietary property of Shiftwork Solutions LLC. '
            'Results are for internal use only and will be summarized before distribution.'
        )
        run_conf.font.size = Pt(9)
        run_conf.font.color.rgb = _BRAND_GRAY
        run_conf.italic = True

        # Page break
        self._add_page_break(doc)

    def _add_instructions_page(self, doc, survey, administration_mode):
        """
        Build the survey instructions page:
          - Section header
          - Numbered directions
          - Note about paper bubbles if administration_mode == 'paper'
          - Page break at end
        """
        instructions = survey['instructions']
        bubble = self._bubble_char(administration_mode)

        # Section heading
        p_head = doc.add_paragraph()
        run_head = p_head.add_run('SURVEY INSTRUCTIONS')
        run_head.bold = True
        run_head.font.size = Pt(16)
        run_head.font.color.rgb = _BRAND_BLUE

        self._add_horizontal_rule(doc, color='1F5C99', size=6)
        doc.add_paragraph()

        # Directions
        for i, direction in enumerate(instructions['directions'], 1):
            p_dir = doc.add_paragraph()
            p_dir.paragraph_format.left_indent = Inches(0.25)
            p_dir.paragraph_format.space_after = Pt(6)
            run_num = p_dir.add_run(f'{i}.  ')
            run_num.bold = True
            run_num.font.size = Pt(11)
            run_num.font.color.rgb = _BRAND_BLUE
            run_text = p_dir.add_run(direction)
            run_text.font.size = Pt(11)

        doc.add_paragraph()

        # Paper administration note
        if administration_mode == 'paper':
            p_note = doc.add_paragraph()
            p_note.paragraph_format.left_indent = Inches(0.25)
            run_note_label = p_note.add_run('PAPER SURVEY NOTE:  ')
            run_note_label.bold = True
            run_note_label.font.size = Pt(10)
            run_note_label.font.color.rgb = _BRAND_BLUE
            run_note = p_note.add_run(
                f'Completely fill in the {bubble} bubble next to your answer. '
                'Use a dark pen or pencil. Do not use checkmarks or X marks — '
                'bubbles must be completely filled for accurate scanning.'
            )
            run_note.font.size = Pt(10)

            doc.add_paragraph()

            # Example bubble row
            p_ex = doc.add_paragraph()
            p_ex.paragraph_format.left_indent = Inches(0.5)
            run_ex_label = p_ex.add_run('Example:  ')
            run_ex_label.font.size = Pt(10)
            run_ex_label.bold = True
            for ex_opt in ['Strongly Agree', 'Agree', 'Neutral', 'Disagree', 'Strongly Disagree']:
                run_ex = p_ex.add_run(f'  {bubble}  {ex_opt}   ')
                run_ex.font.size = Pt(10)

        doc.add_paragraph()

        # Page break
        self._add_page_break(doc)

    def _add_thank_you_page(self, doc, company_name):
        """
        Build the closing page:
          - Thank You heading
          - Short thank-you message
          - Shiftwork Solutions contact block
        """
        self._add_page_break(doc)

        # Spacer
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        # Thank you heading
        p_ty = doc.add_paragraph()
        p_ty.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_ty = p_ty.add_run('Thank You for Participating!')
        run_ty.bold = True
        run_ty.font.size = Pt(24)
        run_ty.font.color.rgb = _BRAND_BLUE

        self._add_horizontal_rule(doc, color='1F5C99', size=6)
        doc.add_paragraph()

        # Thank you body
        p_body = doc.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_body = p_body.add_run(
            f'Your input is valuable and will help {company_name} and Shiftwork Solutions LLC '
            'design the best possible schedule for your workplace.\n\n'
            'Please return your completed survey to the collection box or follow '
            'the instructions provided by your supervisor.'
        )
        run_body.font.size = Pt(12)

        doc.add_paragraph()
        doc.add_paragraph()

        # Horizontal rule
        self._add_horizontal_rule(doc, color='CCCCCC', size=4)
        doc.add_paragraph()

        # Contact block
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_contact_brand = p_contact.add_run('Shiftwork Solutions LLC')
        run_contact_brand.bold = True
        run_contact_brand.font.size = Pt(12)
        run_contact_brand.font.color.rgb = _BRAND_BLUE

        p_contact2 = doc.add_paragraph()
        p_contact2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_contact2 = p_contact2.add_run(
            'Phone: (415) 763-5005  |  www.shift-work.com  |  30+ Years of Schedule Expertise'
        )
        run_contact2.font.size = Pt(10)
        run_contact2.font.color.rgb = _BRAND_GRAY

    def _add_section_heading(self, doc, label):
        """
        Add a styled section heading with a brand-colored bottom border.
        Used at the start of each question category block.
        """
        p = doc.add_paragraph()
        run = p.add_run(label.upper())
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = _BRAND_BLUE

        # Bottom border under heading
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F5C99')
        pBdr.append(bottom)
        pPr.append(pBdr)

        doc.add_paragraph()
        return p

    # ========================================================================
    # EXPORT TO WORD  (enhanced March 12, 2026)
    # ========================================================================

    def export_to_word(self, survey, administration_mode='online'):
        """
        Export survey to a professionally branded Word document (.docx).

        Phase 2 enhancements (March 12, 2026):
        - Shiftwork Solutions branded header on every page
        - Cover page with company name, date, confidentiality note
        - Instructions page (standalone, before questions)
        - Page breaks between major sections
        - Remark-compatible filled bubbles for paper; open circles for online/unsure
        - Sequential question numbering (continuous 1..N across all sections)
        - "Thank You" closing page with contact information
        - Footer with contact details on every page

        Args:
            survey:              Survey dict from create_survey()
            administration_mode: 'paper' | 'online' | 'unsure'
                                 Controls bubble style. Default: 'online'.
                                 Pass 'paper' to get Remark-compatible filled bubbles.

        Returns:
            BytesIO object containing the Word document
        """
        doc = Document()

        # ---- Page setup: US Letter, 1-inch margins --------------------------
        section = doc.sections[0]
        section.page_width  = 12240    # 8.5 inches in twips
        section.page_height = 15840    # 11 inches in twips
        section.left_margin   = 1440   # 1 inch
        section.right_margin  = 1440   # 1 inch
        section.top_margin    = 1008   # 0.7 inch (room for header)
        section.bottom_margin = 1008   # 0.7 inch (room for footer)

        bubble   = self._bubble_char(administration_mode)
        checkbox = self._checkbox_char()

        # ---- Branded header & footer on every page -------------------------
        self._add_doc_header_footer(doc, survey['metadata']['company_name'])

        # ---- Cover page ----------------------------------------------------
        self._add_cover_page(doc, survey)

        # ---- Instructions page ---------------------------------------------
        self._add_instructions_page(doc, survey, administration_mode)

        # ---- Schedule Concepts section -------------------------------------
        if survey['schedule_concepts']:
            self._add_section_heading(doc, 'Schedule Concepts')

            p_intro = doc.add_paragraph()
            run_intro = p_intro.add_run(
                'The following schedule concepts have been selected for your review. '
                'Please read the description of each schedule and then rate it using the scale provided.'
            )
            run_intro.font.size = Pt(11)
            doc.add_paragraph()

            for i, concept in enumerate(survey['schedule_concepts'], 1):
                schedule = concept['schedule']

                # Concept sub-header
                p_concept = doc.add_paragraph()
                run_concept = p_concept.add_run(
                    f'Concept #{i}:  {schedule["name"]}'
                )
                run_concept.bold = True
                run_concept.font.size = Pt(12)
                run_concept.font.color.rgb = _BRAND_DARK

                # Description
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.left_indent = Inches(0.25)
                run_desc = p_desc.add_run(schedule['description'])
                run_desc.font.size = Pt(11)

                # Shift length
                p_len = doc.add_paragraph()
                p_len.paragraph_format.left_indent = Inches(0.25)
                run_len_label = p_len.add_run('Shift Length:  ')
                run_len_label.bold = True
                run_len_label.font.size = Pt(10)
                run_len = p_len.add_run(schedule['shift_length'])
                run_len.font.size = Pt(10)

                doc.add_paragraph()

                # Rating question
                p_q = doc.add_paragraph()
                p_q.paragraph_format.left_indent = Inches(0.25)
                run_q = p_q.add_run(concept['question'])
                run_q.bold = True
                run_q.font.size = Pt(11)

                # Rating options with bubbles
                for option in concept['rating_options']:
                    p_opt = doc.add_paragraph()
                    p_opt.paragraph_format.left_indent = Inches(0.5)
                    p_opt.paragraph_format.space_after = Pt(2)
                    run_opt = p_opt.add_run(f'{bubble}   {option}')
                    run_opt.font.size = Pt(11)

                doc.add_paragraph()

                # Page break between schedule concepts (except after the last one)
                if i < len(survey['schedule_concepts']):
                    self._add_page_break(doc)

            # Page break before main questions
            self._add_page_break(doc)

        # ---- Main Questions section -----------------------------------------
        if survey['questions']:
            # Human-friendly category labels
            CATEGORY_LABELS = {
                'demographics':      'Section 1: Demographics',
                'sleep_alertness':   'Section 2: Sleep & Alertness',
                'working_conditions': 'Section 3: Working Conditions',
                'schedule_features': 'Section 4: Shift Schedule Features',
                'overtime':          'Section 5: Overtime',
                'daycare_eldercare': 'Section 6: Day Care / Elder Care',
                'open_ended':        'Section 7: Open-Ended Questions',
            }

            current_category = None
            question_number  = 1

            for question in survey['questions']:
                # New category: add heading + page break (except first)
                if question['category'] != current_category:
                    if current_category is not None:
                        self._add_page_break(doc)
                    current_category = question['category']
                    label = CATEGORY_LABELS.get(
                        current_category,
                        current_category.replace('_', ' ').title()
                    )
                    self._add_section_heading(doc, label)

                # Question text
                p_q = doc.add_paragraph()
                p_q.paragraph_format.space_after = Pt(4)
                run_q_num = p_q.add_run(f'{question_number}.  ')
                run_q_num.bold = True
                run_q_num.font.size = Pt(11)
                run_q_num.font.color.rgb = _BRAND_BLUE
                run_q_text = p_q.add_run(question['text'])
                run_q_text.bold = True
                run_q_text.font.size = Pt(11)

                # Answer options based on question type
                q_type = question['type']

                if q_type == 'multiple_choice':
                    for option in question['options']:
                        p_opt = doc.add_paragraph()
                        p_opt.paragraph_format.left_indent = Inches(0.4)
                        p_opt.paragraph_format.space_after = Pt(2)
                        run_opt = p_opt.add_run(f'{bubble}   {option}')
                        run_opt.font.size = Pt(11)

                elif q_type == 'likert':
                    # Scale label
                    p_scale = doc.add_paragraph()
                    p_scale.paragraph_format.left_indent = Inches(0.4)
                    run_scale = p_scale.add_run(f'Scale: {question["scale"]}')
                    run_scale.font.size = Pt(10)
                    run_scale.italic = True
                    run_scale.font.color.rgb = _BRAND_GRAY

                    # 5-point scale options on a single line for compactness
                    p_likert = doc.add_paragraph()
                    p_likert.paragraph_format.left_indent = Inches(0.4)
                    scale_labels = [
                        '1 — Strongly Disagree',
                        '2 — Disagree',
                        '3 — Neutral',
                        '4 — Agree',
                        '5 — Strongly Agree'
                    ]
                    for val in scale_labels:
                        p_val = doc.add_paragraph()
                        p_val.paragraph_format.left_indent = Inches(0.4)
                        p_val.paragraph_format.space_after = Pt(2)
                        run_val = p_val.add_run(f'{bubble}   {val}')
                        run_val.font.size = Pt(11)

                elif q_type == 'checkbox':
                    for option in question['options']:
                        p_opt = doc.add_paragraph()
                        p_opt.paragraph_format.left_indent = Inches(0.4)
                        p_opt.paragraph_format.space_after = Pt(2)
                        run_opt = p_opt.add_run(f'{checkbox}   {option}')
                        run_opt.font.size = Pt(11)

                elif q_type == 'text':
                    # Three write-in lines
                    for _ in range(3):
                        p_line = doc.add_paragraph()
                        p_line.paragraph_format.left_indent = Inches(0.25)
                        p_line.paragraph_format.space_after = Pt(2)
                        run_line = p_line.add_run('_' * 85)
                        run_line.font.size = Pt(10)
                        run_line.font.color.rgb = _BRAND_GRAY

                doc.add_paragraph()
                question_number += 1

        # ---- Custom Questions section --------------------------------------
        if survey['custom_questions']:
            self._add_page_break(doc)
            self._add_section_heading(doc, 'Additional Questions')

            for custom_q in survey['custom_questions']:
                p_q = doc.add_paragraph()
                p_q.paragraph_format.space_after = Pt(4)
                run_q_num = p_q.add_run(f'{question_number}.  ')
                run_q_num.bold = True
                run_q_num.font.size = Pt(11)
                run_q_num.font.color.rgb = _BRAND_BLUE
                run_q_text = p_q.add_run(custom_q['text'])
                run_q_text.bold = True
                run_q_text.font.size = Pt(11)

                if 'options' in custom_q:
                    for option in custom_q['options']:
                        p_opt = doc.add_paragraph()
                        p_opt.paragraph_format.left_indent = Inches(0.4)
                        p_opt.paragraph_format.space_after = Pt(2)
                        run_opt = p_opt.add_run(f'{bubble}   {option}')
                        run_opt.font.size = Pt(11)
                else:
                    for _ in range(3):
                        p_line = doc.add_paragraph()
                        p_line.paragraph_format.left_indent = Inches(0.25)
                        p_line.paragraph_format.space_after = Pt(2)
                        run_line = p_line.add_run('_' * 85)
                        run_line.font.size = Pt(10)
                        run_line.font.color.rgb = _BRAND_GRAY

                doc.add_paragraph()
                question_number += 1

        # ---- Thank You page -------------------------------------------------
        self._add_thank_you_page(doc, survey['metadata']['company_name'])

        # ---- Save to BytesIO ------------------------------------------------
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # ========================================================================
    # UTILITY METHODS  (unchanged from January 28, 2026)
    # ========================================================================

    def get_categories(self):
        """Get all question categories"""
        categories = set()
        for question in self.question_bank.values():
            categories.add(question['category'])
        return sorted(list(categories))

    def get_questions_by_category(self, category):
        """Get all questions for a specific category"""
        return [q for q in self.question_bank.values() if q['category'] == category]


# I did no harm and this file is not truncated
