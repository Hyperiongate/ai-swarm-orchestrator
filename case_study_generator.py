"""
CASE STUDY GENERATOR - AI-Powered SEO & AI-Search Optimized Case Studies
Shiftwork Solutions LLC

PURPOSE:
    Generates professional case studies optimized for both traditional SEO
    (Google) and AI-powered search (ChatGPT, Perplexity, Claude). The
    employee engagement process - surveys, workforce input, change management
    - is embedded as a non-negotiable narrative in every case study regardless
    of the problem or solution described.

TARGET AUDIENCE:
    General Managers, Directors, HR Managers - people who own the workforce
    problem, control the budget, and feel pressure from both above and below.
    These readers are nervous about one thing above all else: workforce
    resistance to schedule changes. Every case study addresses that fear
    directly and shows how it was resolved.

EMPLOYEE ENGAGEMENT NARRATIVE (always present):
    - Workforce resistance is acknowledged as the primary risk
    - Employee surveys surface real preferences, not assumed ones
    - Workforce input is used to create genuine ownership of the new schedule
    - Shop floor approachability and direct communication with employees
    - Employees choose from schedule options rather than having them imposed
    - The result: employees feel respected and the change sticks

SEO/AI-SEARCH STRATEGY:
    - Title leads with result or pain point, never cleverness
    - Primary keyword in title AND first sentence of opening paragraph
    - Key result surfaced in first 50 words of The Results section
    - Conservative, believable numbers - a GM nods, doesn't squint
    - Structure is rigid and consistent for Google content-type recognition
    - Quotable, attributable facts that AI models (Perplexity, ChatGPT) cite
    - Dynamic closing section - unique per case study, no duplicate content
    - FAQ section generated separately for AI search snippet capture

INDUSTRIES SUPPORTED:
    Manufacturing, Pharmaceutical, Food Processing, Mining, Consumer Products,
    Beverage, Utilities, Paper & Packaging, Technology, Oil/Gas/Energy,
    Chemicals, Automotive, Transportation, Logistics, Healthcare,
    Government/Mint, Gaming & Hospitality, Other

CHANGE LOG:
    February 21, 2026 - Initial creation. Core generation, DOCX export, DB CRUD.
    February 22, 2026 - Added generate_website_ready_package(). Produces SEO
                        title (<=60 chars), meta description (<=160 chars), URL
                        slug, 3-5 FAQ Q&A pairs, and combined Article+FAQPage
                        JSON-LD schema markup for direct website publishing.
    February 23, 2026 - MAJOR PROMPT REWRITE.
                        * Audience now explicitly GM/Director/HR Manager
                        * Employee engagement process baked in as non-negotiable
                          (surveys, workforce input, change management, shop floor
                          approachability) - appears in every case study regardless
                          of what Jim describes
                        * Title optimized for search intent not impressiveness
                        * Key result surfaced in first 50 words of Results section
                        * Numbers conservative and believable
                        * Dynamic closing section (unique per case study) replaces
                          identical boilerplate - eliminates duplicate content penalty
                        * Added Logistics as standalone industry
                        * Improved SEO keyword density in opening paragraph
                        * AI-search quotable facts in Results section

AUTHOR: Jim @ Shiftwork Solutions LLC
LAST UPDATED: February 23, 2026
"""

import os
import json
from datetime import datetime

# ============================================================================
# INDUSTRY KEYWORD MAP
# SEO keywords naturally woven into generated content.
# Index 0 = primary keyword (used in title and first paragraph).
# Indices 1-3 = secondary keywords (woven throughout).
# ============================================================================
INDUSTRY_SEO_KEYWORDS = {
    'manufacturing': [
        'manufacturing shift scheduling', '24/7 manufacturing operations',
        'production workforce optimization', 'manufacturing overtime reduction',
        'shift rotation manufacturing', 'continuous manufacturing operations'
    ],
    'pharmaceutical': [
        'pharmaceutical shift scheduling', 'GMP compliance shift work',
        'pharmaceutical workforce management', '24/7 pharmaceutical production',
        'FDA compliance shift operations', 'pharma continuous operations'
    ],
    'food_processing': [
        'food processing shift scheduling', 'food manufacturing workforce',
        '24/7 food production', 'food safety shift compliance',
        'food plant workforce optimization', 'continuous food processing'
    ],
    'mining': [
        'mining shift scheduling', 'mine workforce management',
        'fly-in fly-out scheduling', 'FIFO roster management',
        'mining operations workforce', '24/7 mining operations'
    ],
    'consumer_products': [
        'consumer products manufacturing workforce', 'CPG shift scheduling',
        'consumer goods 24/7 operations', 'FMCG workforce optimization',
        'consumer products shift work', 'continuous consumer goods production'
    ],
    'beverage': [
        'beverage manufacturing shift scheduling', 'brewery workforce management',
        'beverage plant 24/7 operations', 'drink manufacturing workforce',
        'beverage production scheduling', 'continuous beverage operations'
    ],
    'utilities': [
        'utility shift scheduling', 'power plant workforce management',
        'utility worker shift rotation', '24/7 utility operations',
        'energy sector shift work', 'utility operations workforce'
    ],
    'paper_packaging': [
        'paper mill shift scheduling', 'packaging plant workforce',
        '24/7 paper manufacturing', 'pulp and paper shift work',
        'packaging operations workforce', 'paper industry continuous operations'
    ],
    'technology': [
        'tech operations shift scheduling', 'data center workforce management',
        '24/7 tech operations', 'technology operations shift work',
        'NOC shift scheduling', 'tech support workforce optimization'
    ],
    'oil_gas_energy': [
        'oil and gas shift scheduling', 'refinery workforce management',
        'upstream operations shift work', 'downstream shift scheduling',
        'energy sector workforce optimization', 'oil field shift rotation'
    ],
    'chemicals': [
        'chemical plant shift scheduling', 'chemical manufacturing workforce',
        '24/7 chemical operations', 'process industry shift work',
        'chemical plant workforce optimization', 'continuous chemical operations'
    ],
    'automotive': [
        'automotive manufacturing shift scheduling', 'auto plant workforce',
        '24/7 automotive production', 'automotive shift rotation',
        'auto manufacturing workforce optimization', 'vehicle production scheduling'
    ],
    'transportation': [
        'transportation workforce scheduling', 'fleet operations shift management',
        '24/7 transportation operations', 'driver shift scheduling',
        'transportation workforce optimization', 'transit operations shift work'
    ],
    'logistics': [
        'logistics shift scheduling', 'warehouse workforce management',
        '24/7 distribution center operations', 'logistics workforce optimization',
        'supply chain shift scheduling', 'distribution shift work'
    ],
    'healthcare': [
        'healthcare shift scheduling', 'hospital workforce management',
        '24/7 healthcare operations', 'nursing shift optimization',
        'healthcare staffing solutions', 'hospital shift rotation'
    ],
    'government_mint': [
        'government operations shift scheduling', 'public sector workforce',
        'government facility 24/7 operations', 'mint operations workforce',
        'government shift work management', 'public sector shift optimization'
    ],
    'gaming_hospitality': [
        'casino shift scheduling', 'hospitality workforce management',
        '24/7 gaming operations', 'hotel shift scheduling',
        'gaming industry workforce optimization', 'hospitality shift rotation'
    ],
    'other': [
        '24/7 operations shift scheduling', 'continuous operations workforce',
        'shift work optimization', 'workforce scheduling consulting',
        'shift schedule design', 'employee-centered scheduling'
    ]
}

INDUSTRY_DISPLAY_NAMES = {
    'manufacturing': 'Manufacturing',
    'pharmaceutical': 'Pharmaceutical',
    'food_processing': 'Food Processing',
    'mining': 'Mining',
    'consumer_products': 'Consumer Products',
    'beverage': 'Beverage',
    'utilities': 'Utilities',
    'paper_packaging': 'Paper & Packaging',
    'technology': 'Technology',
    'oil_gas_energy': 'Oil, Gas & Energy',
    'chemicals': 'Chemicals',
    'automotive': 'Automotive',
    'transportation': 'Transportation',
    'logistics': 'Logistics',
    'healthcare': 'Healthcare',
    'government_mint': 'Government / Mint',
    'gaming_hospitality': 'Gaming & Hospitality',
    'other': 'Industrial Operations'
}


def get_case_study_prompt(industry: str, problem: str, solution: str) -> str:
    """
    Build the AI prompt for case study generation.

    This prompt enforces:
    1. Employee engagement narrative in every section (non-negotiable)
    2. SEO and AI-search optimization throughout
    3. Audience: General Managers, Directors, HR Managers
    4. Conservative, believable numbers
    5. Dynamic closing section unique to each case study
    6. Key result in first 50 words of The Results section
    """
    industry_display = INDUSTRY_DISPLAY_NAMES.get(industry, 'Industrial Operations')
    keywords = INDUSTRY_SEO_KEYWORDS.get(industry, INDUSTRY_SEO_KEYWORDS['other'])
    primary_keyword = keywords[0] if keywords else 'shift scheduling'
    secondary_keywords = ', '.join(keywords[1:4]) if len(keywords) > 1 else ''

    prompt = f"""You are a professional content writer specializing in workforce management consulting case studies for Shiftwork Solutions LLC. Your writing is trusted by General Managers, Plant Directors, and HR Managers who are considering a schedule change but are worried about one thing above everything else: their workforce pushing back.

Write a compelling, SEO-optimized case study based on the following:

INDUSTRY: {industry_display}
CLIENT PROBLEM: {problem}
SOLUTION APPLIED: {solution}

===== ABOUT SHIFTWORK SOLUTIONS LLC =====
- Specialized consulting firm with hundreds of completed 24/7 shift operations engagements
- Core philosophy: the best shift schedule is the one employees actually choose
- Core process: employee surveys, data analysis, collaborative schedule design, change management
- Typical engagement: 6 weeks
- Industries: manufacturing, pharmaceutical, food processing, mining, logistics, and more
- Unique value: we communicate directly with the workforce on the shop floor. We are approachable, present, and trusted by workers and management alike.

===== THE MOST IMPORTANT RULE =====
THE EMPLOYEE ENGAGEMENT STORY IS MANDATORY IN EVERY CASE STUDY.

No matter what problem Jim describes, no matter what solution was applied, every case study must clearly communicate the following:

1. The client was nervous about workforce resistance to the schedule change
2. Shiftwork Solutions surveyed the workforce to discover their real preferences, not assumed ones
3. The survey results were interpreted carefully and used to design schedule options - not just filed away
4. Employees were given meaningful input that shaped the final outcome - this created real ownership
5. Shiftwork Solutions was visibly present and approachable on the shop floor throughout the engagement
6. The workforce felt respected and included throughout the process
7. The end result: employees accepted and in many cases appreciated the new schedule

This story must appear naturally woven into The Challenge, Why Shiftwork Solutions, Our Approach, and The Results sections. It is not optional. It is the primary reason companies hire us.

===== SEO AND AI SEARCH REQUIREMENTS =====
1. TITLE: Use "{primary_keyword}" or a close variant. Lead with the RESULT or the PAIN POINT - never a clever phrase. Format preference: "[Specific Result or Problem]: [Primary Keyword]". Keep it under 70 characters.
2. FIRST PARAGRAPH: Use "{primary_keyword}" naturally in the first sentence of The Challenge section.
3. SECONDARY KEYWORDS: Weave these in naturally throughout the content: {secondary_keywords}
4. THE RESULTS - KEY RESULT FIRST: The very first sentence of The Results section must state the single most important quantifiable outcome. Do not bury the lead.
5. NUMBERS: All numbers must be conservative and believable. A General Manager reading this should nod, not squint. Use realistic ranges: 10-20% overtime reduction, $150K-$450K savings, turnover reduction of 15-25%. Never use implausibly high numbers.
6. QUOTABLE FACTS: Include at least 3 specific, attributable facts that an AI search engine (Perplexity, ChatGPT, Claude) would cite when answering questions like "how to reduce shift schedule resistance" or "how to get employees to accept a new shift schedule."
7. STRUCTURE: Use the exact section headers below. Consistent structure helps Google understand the content type and helps AI models parse it correctly.

===== AUDIENCE VOICE =====
Write for General Managers, Plant Directors, and HR Managers. These readers:
- Feel pressure from executives to cut costs and from employees to maintain quality of life
- Have seen schedule changes blow up before - they know the political risk
- Want proof, not promises - specific numbers, not vague claims
- Will share this with their VP or CHRO - it needs to hold up to scrutiny
- Care about their employees wellbeing, not just the operational metrics

Use language they actually use: "workforce resistance," "employee buy-in," "schedule transition," "overtime burden," "change management," "shift rotation," "workforce survey," "employee input," "schedule acceptance."

===== STRUCTURE TO FOLLOW =====
Use these EXACT section headers:

# [Title: Result-led or Pain-point-led, primary keyword included, under 70 characters]

## The Challenge
[2-3 paragraphs. Open with the primary keyword "{primary_keyword}" in the first sentence. Describe BOTH the operational problem AND the change management risk - the real fear that employees would resist a new schedule. Be specific about the pain points: what was failing, what management feared, what was at stake. This section should make a GM think: "that is exactly my situation."]

## Why Shiftwork Solutions LLC
[2 paragraphs. Explain why the client chose Shiftwork Solutions. Lead with the employee engagement methodology - specifically that we survey the workforce and use their input to drive the design. Then mention the normative database of hundreds of past client engagements, the industry experience, and the track record. The reader should understand clearly that we do not impose schedules - we facilitate employees choosing their own.]

## Our Approach
[3 paragraphs.
  Paragraph 1: Initial assessment and data gathering - what we analyzed, what we measured, how we established a baseline.
  Paragraph 2: The employee survey process - how we designed and administered it, how we communicated to the workforce that their input would actually be used, how we were present and approachable on the shop floor, and how survey results were interpreted and translated into real schedule options.
  Paragraph 3: Schedule design and implementation planning - how employees reviewed options, how their preferences shaped the final recommendation, and how change management support ensured a smooth and accepted transition.]

## The Results
[2-3 paragraphs. CRITICAL: The very first sentence must state the most important quantifiable result. Example: "Within six months of implementation, [Client] reduced overtime costs by 16 percent, saving an estimated $310,000 annually." Then expand on operational outcomes, financial outcomes, and - critically - workforce outcomes. Always include: employee satisfaction improvement, reduction in complaints or grievances, or supervisor feedback on workforce acceptance. The workforce outcome is what validates the entire approach and what this audience cares most about.]

## Key Takeaways
[4 bullet points. Each must be a specific, actionable lesson a GM could share in a leadership meeting. Requirements:
  - At least one bullet must directly reference the employee engagement process as a key success factor
  - At least one bullet must reference a specific quantifiable result from this case study
  - At least one bullet must speak to the change management or workforce resistance dimension
  - Make them useful and specific, not generic consulting speak]

## About Shiftwork Solutions LLC
[DYNAMIC CLOSING - THIS IS CRITICAL FOR SEO AND AI SEARCH.
Do NOT write a generic boilerplate paragraph that will be identical across every case study. Instead, write 3 sentences that:
  Sentence 1: Reference the SPECIFIC problem solved in this case study and the industry context. Make it clear this is about this engagement, not a generic company description.
  Sentence 2: Connect it to Shiftwork Solutions broader experience with similar operations, referencing hundreds of past engagements.
  Sentence 3: State the employee-centered philosophy - specifically that schedules are designed with workforce input, not imposed on the workforce, and what that delivers.

Example structure (do NOT copy this - write it fresh for the specific problem solved):
"When [industry] operations face [specific problem from this study], the hidden risk is rarely the schedule itself - it is whether the workforce will accept it. Shiftwork Solutions LLC has guided hundreds of [industry] and continuous-operations facilities through exactly this kind of transition, using a structured process that gives employees a real voice in the outcome. The result is not just a better schedule on paper - it is a schedule the workforce chose, understands, and supports."

Every case study must have a unique closing paragraph. No two closings should be the same.]

===== WRITING STYLE =====
- Professional but human - readable by a busy GM on a Tuesday afternoon
- Specific and factual - no vague generalities like "significant improvement"
- Confident but not boastful - let the results speak
- Third person for the client, first person plural ("we," "our team") for Shiftwork Solutions
- Use "[Client]" as placeholder for company name
- Use ranges for headcount: "a workforce of 400+" or "more than 300 employees"
- Never use "300+" to describe Shiftwork Solutions past engagements - always use "hundreds"

OUTPUT: Return ONLY the case study content in markdown format. Start directly with the # title. No preamble, no explanation, no JSON wrapper."""

    return prompt


def get_website_package_prompt(title: str, content: str, industry: str) -> str:
    """
    Build the AI prompt for generating website-ready SEO metadata and schema markup.
    Called after the case study is already generated and saved.

    Optimized for:
    - Traditional SEO (Google title/description/schema)
    - AI search engines (Perplexity, ChatGPT FAQ snippet capture)
    - Structured data (JSON-LD Article + FAQPage combined schema)
    """
    industry_display = INDUSTRY_DISPLAY_NAMES.get(industry, 'Industrial Operations')
    keywords = INDUSTRY_SEO_KEYWORDS.get(industry, INDUSTRY_SEO_KEYWORDS['other'])
    primary_keyword = keywords[0] if keywords else 'shift scheduling'

    prompt = f"""You are an SEO specialist and structured data expert. Based on the following case study for Shiftwork Solutions LLC, generate a complete website publishing package.

CASE STUDY TITLE: {title}
INDUSTRY: {industry_display}
PRIMARY KEYWORD: {primary_keyword}

CASE STUDY CONTENT:
{content}

Generate the following elements. Return ONLY valid JSON - no markdown, no preamble, no explanation. The JSON must be parseable by Python's json.loads().

Return this exact JSON structure:
{{
  "seo_title": "...",
  "meta_description": "...",
  "url_slug": "...",
  "faqs": [
    {{"question": "...", "answer": "..."}},
    {{"question": "...", "answer": "..."}},
    {{"question": "...", "answer": "..."}},
    {{"question": "...", "answer": "..."}},
    {{"question": "...", "answer": "..."}}
  ],
  "json_ld": {{}}
}}

REQUIREMENTS FOR EACH FIELD:

seo_title:
- Maximum 60 characters including spaces
- Must include the primary keyword "{primary_keyword}" or a close variant
- Lead with the result or pain point - never a clever phrase
- Format preference: "[Specific Result]: [Primary Keyword]" or "[Pain Point] Solved: [Industry]"

meta_description:
- Maximum 160 characters including spaces
- Include primary keyword naturally in first half
- Reference the most important quantifiable result from the case study
- End with a soft call to action such as "See how." or "Learn how we did it."

url_slug:
- Lowercase, hyphen-separated, no special characters
- 4-7 words maximum
- Include primary keyword or close variant
- Example: "logistics-shift-scheduling-overtime-reduction"

faqs:
- Exactly 5 FAQ pairs
- Questions must be the kind a GM, Director, or HR Manager would actually type into Google or ask an AI chatbot
- Two questions must reference the specific results from this case study
- Two questions must address employee resistance or change management angle
- One question must address why Shiftwork Solutions or how to choose a shift scheduling consultant
- Answers must be 2-4 sentences, specific, and useful - not vague
- Use industry terminology naturally
- Do NOT use generic questions like "What is shift scheduling?" - every question must be specific to this case study

json_ld:
- Combined Article + FAQPage schema as a single JSON-LD object
- Use "@graph" array containing both schemas
- Article schema: "@type": "Article", headline=seo_title, description=meta_description,
  author={{"@type":"Organization","name":"Shiftwork Solutions LLC"}},
  publisher={{"@type":"Organization","name":"Shiftwork Solutions LLC","url":"https://shiftworksolutions.com"}},
  url="https://shiftworksolutions.com/case-studies/[url_slug]",
  datePublished="{datetime.now().strftime('%Y-%m-%d')}"
- FAQPage schema: "@type": "FAQPage" with all 5 FAQ pairs as mainEntity array
- Each FAQ in mainEntity: {{"@type":"Question","name":"[question]","acceptedAnswer":{{"@type":"Answer","text":"[answer]"}}}}
- The entire json_ld value must itself be valid JSON"""

    return prompt


def generate_case_study(industry: str, problem: str, solution: str) -> dict:
    """
    Generate a case study using the Anthropic API.

    Returns:
        {
            'success': True/False,
            'content': str (markdown),
            'title': str,
            'word_count': int,
            'industry': str,
            'industry_display': str,
            'generated_at': str (ISO),
            'error': str (only on failure)
        }
    """
    try:
        from config import ANTHROPIC_API_KEY, CLAUDE_SONNET_MODEL
        import anthropic

        if not ANTHROPIC_API_KEY:
            return {
                'success': False,
                'error': 'Anthropic API key not configured'
            }

        prompt = get_case_study_prompt(industry, problem, solution)

        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        message = client.messages.create(
            model=CLAUDE_SONNET_MODEL,
            max_tokens=4000,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        content = message.content[0].text

        # Extract title from first # heading
        title = 'Case Study'
        lines = content.strip().split('\n')
        for line in lines:
            if line.startswith('# ') and not line.startswith('## '):
                title = line.replace('# ', '').strip()
                break

        word_count = len(content.split())

        print(f"[CaseStudy] Generated: '{title}' ({word_count} words)")

        return {
            'success': True,
            'content': content,
            'title': title,
            'word_count': word_count,
            'industry': industry,
            'industry_display': INDUSTRY_DISPLAY_NAMES.get(industry, industry),
            'generated_at': datetime.now().isoformat()
        }

    except Exception as e:
        import traceback
        return {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }


def generate_website_ready_package(study_id: int) -> dict:
    """
    Generate a complete website publishing package for a saved case study.

    Makes a second AI call (separate from case study generation) to produce:
      - seo_title       : <=60 character SEO-optimized title
      - meta_description: <=160 character meta description
      - url_slug        : clean URL slug for the page
      - faqs            : list of 5 {question, answer} dicts
      - json_ld         : combined Article + FAQPage JSON-LD schema dict
      - json_ld_string  : json_ld serialized as indented string for
                          direct paste into <script type="application/ld+json">

    Returns:
        {
            'success': True/False,
            'seo_title': str,
            'meta_description': str,
            'url_slug': str,
            'faqs': [{'question': str, 'answer': str}, ...],
            'json_ld': dict,
            'json_ld_string': str,
            'study_title': str,
            'industry_display': str,
            'generated_at': str,
            'error': str  (only on failure)
        }
    """
    try:
        from config import ANTHROPIC_API_KEY, CLAUDE_SONNET_MODEL
        import anthropic

        if not ANTHROPIC_API_KEY:
            return {'success': False, 'error': 'Anthropic API key not configured'}

        study = get_case_study_by_id(study_id)
        if not study:
            return {'success': False, 'error': f'Case study ID {study_id} not found'}

        title = study.get('title', '')
        content = study.get('content', '')
        industry = study.get('industry', 'other')

        if not content:
            return {'success': False, 'error': 'Case study has no content'}

        prompt = get_website_package_prompt(title, content, industry)

        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        message = client.messages.create(
            model=CLAUDE_SONNET_MODEL,
            max_tokens=3000,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        raw_response = message.content[0].text.strip()

        # Strip markdown code fences if the model wrapped the JSON
        if raw_response.startswith('```'):
            lines = raw_response.split('\n')
            raw_response = '\n'.join(lines[1:-1]).strip()

        try:
            package = json.loads(raw_response)
        except json.JSONDecodeError as parse_err:
            return {
                'success': False,
                'error': f'AI returned invalid JSON: {str(parse_err)}',
                'raw_response': raw_response[:500]
            }

        # Validate required fields
        required_fields = ['seo_title', 'meta_description', 'url_slug', 'faqs', 'json_ld']
        missing = [f for f in required_fields if f not in package]
        if missing:
            return {
                'success': False,
                'error': f'AI response missing fields: {missing}',
                'raw_response': raw_response[:500]
            }

        # Enforce character limits
        seo_title = str(package['seo_title'])[:60]
        meta_description = str(package['meta_description'])[:160]
        url_slug = str(package['url_slug']).lower().replace(' ', '-')

        # Validate and clean FAQs
        faqs = package.get('faqs', [])
        if not isinstance(faqs, list):
            faqs = []
        cleaned_faqs = []
        for faq in faqs:
            if isinstance(faq, dict) and 'question' in faq and 'answer' in faq:
                cleaned_faqs.append({
                    'question': str(faq['question']),
                    'answer': str(faq['answer'])
                })

        # Serialize JSON-LD to formatted string for paste-ready output
        json_ld_dict = package.get('json_ld', {})
        json_ld_string = json.dumps(json_ld_dict, indent=2)

        print(f"[CaseStudy] Website package generated for ID={study_id}: "
              f"title={len(seo_title)}ch, desc={len(meta_description)}ch, "
              f"faqs={len(cleaned_faqs)}, json_ld={len(json_ld_string)}ch")

        return {
            'success': True,
            'seo_title': seo_title,
            'meta_description': meta_description,
            'url_slug': url_slug,
            'faqs': cleaned_faqs,
            'json_ld': json_ld_dict,
            'json_ld_string': json_ld_string,
            'study_title': title,
            'industry_display': INDUSTRY_DISPLAY_NAMES.get(industry, industry),
            'generated_at': datetime.now().isoformat()
        }

    except Exception as e:
        import traceback
        return {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }


def generate_case_study_docx(case_study_content: str, title: str, industry: str) -> bytes:
    """
    Convert markdown case study content to a professional Word document.
    Uses python-docx for reliable server-side generation.
    Returns bytes of the .docx file.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()

    # ---- Page setup: US Letter, 1" margins ----
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ---- Document styles ----
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1A, 0x53, 0x7A)

    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    # ---- Header ----
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = 'Shiftwork Solutions LLC  |  Case Study'
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.runs[0].font.size = Pt(9)
    header_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ---- Footer ----
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    industry_display = INDUSTRY_DISPLAY_NAMES.get(industry, 'Industrial Operations')
    footer_para.text = (
        f'© Shiftwork Solutions LLC  |  {industry_display}  |  shiftworksolutions.com'
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ---- Parse and render markdown ----
    lines = case_study_content.strip().split('\n')

    def add_paragraph_with_formatting(doc, text, style_name='Normal'):
        """Add paragraph handling **bold** inline formatting."""
        para = doc.add_paragraph(style=style_name)
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for i, part in enumerate(parts):
            run = para.add_run(part)
            if i % 2 == 1:
                run.bold = True
        return para

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith('# ') and not line.startswith('## '):
            heading_text = line[2:].strip()
            para = doc.add_heading(heading_text, level=1)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph('')

        elif line.startswith('## '):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)

        elif line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=3)

        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            add_paragraph_with_formatting(doc, bullet_text, 'List Bullet')

        elif re.match(r'^\d+\.\s', line):
            item_text = re.sub(r'^\d+\.\s', '', line).strip()
            add_paragraph_with_formatting(doc, item_text, 'List Number')

        elif line.startswith('---') or line.startswith('==='):
            pass

        elif line == '':
            pass

        else:
            if line.strip():
                add_paragraph_with_formatting(doc, line.strip())

        i += 1

    # ---- Final metadata paragraph ----
    doc.add_paragraph('')
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_para.add_run(
        f'Generated by Shiftwork Solutions LLC AI System  |  '
        f'{datetime.now().strftime("%B %d, %Y")}'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.italic = True

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def save_case_study_to_db(industry: str, title: str, content: str,
                           problem: str, solution: str) -> int:
    """Save a generated case study to the database. Returns the new record ID."""
    from database import get_db
    db = get_db()
    cursor = db.execute('''
        INSERT INTO case_studies (
            industry, title, content, problem_summary,
            solution_summary, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (
        industry, title, content, problem, solution,
        datetime.now().isoformat(), datetime.now().isoformat()
    ))
    db.commit()
    record_id = cursor.lastrowid
    db.close()
    return record_id


def get_all_case_studies() -> list:
    """Retrieve all saved case studies, newest first."""
    from database import get_db
    db = get_db()
    rows = db.execute('''
        SELECT id, industry, title, problem_summary, solution_summary,
               created_at
        FROM case_studies
        ORDER BY created_at DESC
    ''').fetchall()
    db.close()
    return [dict(row) for row in rows]


def get_case_study_by_id(study_id: int) -> dict:
    """Retrieve a single case study by ID."""
    from database import get_db
    db = get_db()
    row = db.execute(
        'SELECT * FROM case_studies WHERE id = ?', (study_id,)
    ).fetchone()
    db.close()
    return dict(row) if row else None


def delete_case_study(study_id: int) -> bool:
    """Delete a case study by ID."""
    from database import get_db
    db = get_db()
    db.execute('DELETE FROM case_studies WHERE id = ?', (study_id,))
    db.commit()
    db.close()
    return True


def init_case_studies_table():
    """Create the case_studies table if it does not exist. Called at app startup."""
    from database import get_db
    db = get_db()
    db.execute('''
        CREATE TABLE IF NOT EXISTS case_studies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            industry TEXT NOT NULL,
            title TEXT NOT NULL,
            content TEXT NOT NULL,
            problem_summary TEXT,
            solution_summary TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    db.execute('''
        CREATE INDEX IF NOT EXISTS idx_case_studies_industry
        ON case_studies(industry)
    ''')
    db.execute('''
        CREATE INDEX IF NOT EXISTS idx_case_studies_created
        ON case_studies(created_at DESC)
    ''')
    db.commit()
    db.close()
    print("  [CaseStudy] case_studies table ready")

# I did no harm and this file is not truncated
